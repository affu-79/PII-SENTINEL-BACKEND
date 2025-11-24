"""
Flask API for PII Sentinel Backend.
Handles file uploads, batch management, PII detection, and masking.
"""
import os
import uuid
import logging
import random
import time
import zipfile
from typing import Optional
from datetime import datetime, timedelta
from typing import Dict, Any, Tuple
from flask import Flask, request, jsonify, send_file, make_response
import bcrypt
import json
import base64
from flask_cors import CORS
from cryptography.hazmat.primitives.ciphers.aead import AESGCM
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
from dotenv import load_dotenv
import threading
from functools import wraps
from concurrent.futures import ThreadPoolExecutor
from cachetools import cached, TTLCache
import requests

# Create a cache instance for profile endpoint that we can clear
profile_cache = TTLCache(maxsize=100, ttl=2)

def invalidate_profile_cache(user_email=None):
    """Clear the profile cache for a specific user or all users."""
    try:
        if user_email:
            # Clear specific user's cache entry
            profile_cache.pop(user_email, None)
        else:
            # Clear entire cache
            profile_cache.clear()
    except Exception as e:
        logging.getLogger(__name__).warning(f"Failed to invalidate profile cache: {e}")

# Load environment variables FIRST before importing modules that need them
load_dotenv()

# Import PII Detection modules
try:
    from pii_detection_api import pii_detection_bp
    PII_DETECTION_AVAILABLE = True
except ImportError as e:
    logger.warning(f"PII Detection modules not available: {e}")
    PII_DETECTION_AVAILABLE = False

try:
    from controllers.google_auth import google_auth_blueprint
    GOOGLE_AUTH_INTEGRATED = True
except ImportError as e:
    logging.getLogger(__name__).warning(f"Google auth blueprint not available: {e}")
    GOOGLE_AUTH_INTEGRATED = False

# Cache for fast batch retrieval (in-memory, cleared periodically)
batch_cache = {}
batch_cache_lock = threading.Lock()
CACHE_TTL = 10  # Cache for 10 seconds

def get_cached_batches(user_id):
    """Get batches from cache if available and not expired."""
    with batch_cache_lock:
        cache_key = f"user_{user_id}"
        if cache_key in batch_cache:
            cached_data, timestamp = batch_cache[cache_key]
            if time.time() - timestamp < CACHE_TTL:
                return cached_data
        return None

def set_cached_batches(user_id, data):
    """Cache batch data with timestamp."""
    with batch_cache_lock:
        cache_key = f"user_{user_id}"
        batch_cache[cache_key] = (data, time.time())

def invalidate_batch_cache(user_id):
    """Invalidate cache when batches change."""
    with batch_cache_lock:
        cache_key = f"user_{user_id}"
        if cache_key in batch_cache:
            del batch_cache[cache_key]

from mongo_client import mongo_client
from parallel_processor import get_processor
from performance_config import perf_config
from worker_stub import process_file
from maskers import masker
from utils import (
    save_json, load_json, create_zip, get_timestamp, ensure_dir,
    sanitize_filename, is_pdf_file, is_docx_file, is_doc_file, is_image_file, is_text_file
)

# Import security middleware
try:
    from middleware.security import add_security_headers, rate_limit, sanitize_input, validate_file_upload, validate_path
except ImportError:
    # Fallback if middleware not available
    logger.warning("Security middleware not available, using basic security")
    def add_security_headers(response):
        return response
    def rate_limit(max_requests=100, window=60, per_ip=True):
        def decorator(f):
            return f
        return decorator
    def sanitize_input(data):
        return data
    def validate_file_upload(file):
        return True, None
    def validate_path(path):
        return True

# ------------------------------------------------------------
# Pricing / Token configuration
# ------------------------------------------------------------

PLAN_CATALOG: Dict[str, Dict[str, Any]] = {
    "starter": {
        "id": "starter",
        "name": "Starter",
        "price_inr": 0,
        "price_inr_yearly": 0,
        "monthly_tokens": 150,
        "features": {
            "export_json": False,
            "lock_json": False,
            "unlock_json": False,
            "advanced_analysis": False,
            "log_records": False
        }
    },
    "professional": {
        "id": "professional",
        "name": "Professional",
        "price_inr": 999,  # Monthly
        "price_inr_yearly": 9950,  # Yearly (₹829/month, 17% savings)
        "monthly_tokens": 2000,
        "features": {
            "export_json": True,
            "lock_json": True,
            "unlock_json": True,
            "advanced_analysis": True,
            "log_records": False
        }
    },
    "enterprise": {
        "id": "enterprise",
        "name": "Enterprise",
        "price_inr": 4999,  # Monthly
        "price_inr_yearly": 49790,  # Yearly (₹4,149/month, 17% savings)
        "monthly_tokens": None,  # Unlimited
        "features": {
            "export_json": True,
            "lock_json": True,
            "unlock_json": True,
            "advanced_analysis": True,
            "log_records": True
        }
    }
}

TOKEN_ACTION_COSTS: Dict[str, int] = {
    "lock_json": 50,
    "unlock_json": 50,
    "download_masked_file": 5
}

ADDON_TOKEN_PRICE_INR = 1

# Configure plan catalogue inside mongo client (used for token accounting)
mongo_client.configure_plans(PLAN_CATALOG)

# ------------------------------------------------------------
# Logging configuration
# ------------------------------------------------------------

log_level = os.getenv('LOG_LEVEL', 'INFO').upper()
logging.basicConfig(
    level=getattr(logging, log_level, logging.INFO),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# SECURITY: Set Flask secret key for session management
app.secret_key = os.getenv('FLASK_SECRET', os.urandom(32))

# Wrap Flask app with middleware that handles multipart requests better
original_wsgi_app = app.wsgi_app

def custom_wsgi_app(environ, start_response):
    """Custom WSGI middleware to handle multipart requests and bypass Werkzeug validation."""
    path = environ.get('PATH_INFO', '')
    method = environ.get('REQUEST_METHOD', '')
    content_type = environ.get('CONTENT_TYPE', '')
    
    # For upload endpoints, ensure content-type is acceptable to Werkzeug
    if method == 'POST' and '/upload' in path:
        logger.info(f"🔧 WSGI middleware: Upload endpoint detected")
        logger.info(f"   PATH: {path}, Content-Type: '{content_type}'")
        
        # If it's multipart, let it through as-is
        if 'multipart' in content_type.lower():
            logger.info(f"   ✓ Valid multipart/form-data detected")
        elif not content_type:
            # If no content-type, set a default multipart one
            # (browser should have set it, but just in case)
            logger.warning(f"   ⚠️  No Content-Type set, assuming multipart")
            environ['CONTENT_TYPE'] = 'multipart/form-data'
        else:
            logger.warning(f"   ⚠️  Unexpected Content-Type: {content_type}, allowing through anyway")
            # Don't force it - let Flask handle it
    
    return original_wsgi_app(environ, start_response)

app.wsgi_app = custom_wsgi_app

# Debug: Log all incoming requests and bypass 415 for upload
@app.before_request
def log_request():
    logger.info(f"Incoming request: {request.method} {request.path}")
    logger.info(f"  Content-Type: {request.content_type}")
    logger.info(f"  Raw Content-Type header: {request.headers.get('Content-Type')}")
    
    if request.method == 'POST':
        if request.path == '/api/upload':
            logger.info(f"  [UPLOAD] Multipart request detected")
        elif 'multipart' in (request.content_type or ''):
            logger.info(f"  Files in request: {list(request.files.keys())}")
            for key in request.files.keys():
                logger.info(f"    - {key}: {len(request.files.getlist(key))} file(s)")

# SECURITY: Add security headers to all responses
@app.after_request
def after_request(response):
    return add_security_headers(response)

# SECURITY: CORS configuration - restrict origins in production
allowed_origins = os.getenv('CORS_ORIGINS', '*').split(',')
if os.getenv('FLASK_ENV') == 'production' or os.getenv('ENVIRONMENT') == 'production':
    # In production, require explicit origins
    if '*' in allowed_origins:
        logger.warning("CORS configured to allow all origins in production - this is insecure!")
        allowed_origins = ['https://yourdomain.com']  # Replace with actual domain
    
CORS(app, 
     resources={r"/api/*": {
         "origins": allowed_origins,
         "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
         "allow_headers": ["Content-Type", "X-API-KEY", "Authorization", "*"],
         "expose_headers": ["Content-Type"],
         "max_age": 3600,
         "supports_credentials": True
     }},
     supports_credentials=True,
     allow_headers="*",
     expose_headers="*")

# Register PII Detection blueprint
if PII_DETECTION_AVAILABLE:
    app.register_blueprint(pii_detection_bp)
    logger.info("✓ PII Detection API blueprint registered")
else:
    logger.warning("PII Detection API not available")

if GOOGLE_AUTH_INTEGRATED:
    app.register_blueprint(google_auth_blueprint)
    logger.info("✓ Google OAuth blueprint registered")
else:
    logger.warning("Google OAuth blueprint not registered")

# Configuration
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB max file size
storage_path = os.getenv('STORAGE_PATH', './data')
app.config['UPLOAD_FOLDER'] = os.path.join(storage_path, 'uploads')
app.config['RESULTS_FOLDER'] = os.path.join(storage_path, 'results')
app.config['MASKED_FOLDER'] = os.path.join(storage_path, 'masked')

# Ensure directories exist
ensure_dir(app.config['UPLOAD_FOLDER'])
ensure_dir(app.config['RESULTS_FOLDER'])
ensure_dir(app.config['MASKED_FOLDER'])

# In-memory job storage (for real-time processing)
# Use thread-safe dict for concurrent access
jobs = {}
_jobs_lock = threading.Lock()

# OTP storage - Now using MongoDB instead of in-memory
# Keeping in-memory as fallback for backward compatibility
otp_storage = {}
_otp_lock = threading.Lock()
OTP_EXPIRY_SECONDS = 120  # 2 minutes

# Allowed file extensions
ALLOWED_EXTENSIONS = {
    'pdf', 'docx', 'txt', 'csv', 'png', 'jpg', 'jpeg', 'json'
}


def allowed_file(filename: str) -> bool:
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def require_api_key(f):
    """Decorator to require API key."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        api_key = request.headers.get('X-API-KEY')
        expected_key = os.getenv('API_KEY', '')
        
        # SECURITY: Fail securely - require API key in production
        is_production = os.getenv('FLASK_ENV') == 'production' or os.getenv('ENVIRONMENT') == 'production'
        
        if not expected_key:
            if is_production:
                logger.error("API_KEY not set in production environment - rejecting request")
                return jsonify({'error': 'Server configuration error'}), 500
            else:
                logger.warning("API_KEY not set in environment, skipping auth check (development mode)")
            return f(*args, **kwargs)
        
        # SECURITY: Use constant-time comparison to prevent timing attacks
        if not api_key:
            logger.warning("API key missing in request")
            return jsonify({'error': 'Invalid or missing API key'}), 401
        
        # Constant-time comparison
        if len(api_key) != len(expected_key):
            logger.warning("API key length mismatch")
            return jsonify({'error': 'Invalid or missing API key'}), 401
        
        # Use secrets.compare_digest for constant-time comparison
        import secrets
        if not secrets.compare_digest(api_key, expected_key):
            logger.warning("API key mismatch")
            return jsonify({'error': 'Invalid or missing API key'}), 401
        
        return f(*args, **kwargs)
    return decorated_function


def require_auth(f):
    """Decorator to require user authentication (logged in user)."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Check if user_id is provided in the request
        data = None
        user_id = None
        
        # Only try to get JSON if it's not multipart
        if request.method in ['POST', 'PUT'] and request.content_type and 'application/json' in request.content_type:
            try:
                data = request.get_json(force=False, silent=True)
                if data:
                    user_id = data.get('user_id')
            except:
                pass
        
        # Also check query parameters (for file uploads and GET requests)
        if not user_id:
            user_id = request.args.get('user_id')
        
        logger.info(f"require_auth: user_id={user_id}, content_type={request.content_type}, endpoint={request.endpoint}")
        
        # Also check headers for token-based auth
        auth_header = request.headers.get('Authorization')
        token = None
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header[7:]  # Extract token after "Bearer "
        
        # Reject if no user_id and no token
        if not user_id and not token:
            logger.warning(f"Authentication required: No user_id or token provided - endpoint: {request.endpoint}")
            return jsonify({
                'error': 'User must be signed in to perform this action',
                'code': 'AUTHENTICATION_REQUIRED'
            }), 401
        
        # If user_id is 'default', reject (means not logged in)
        if user_id == 'default':
            logger.warning("Authentication required: Attempted with default user_id")
            return jsonify({
                'error': 'User must be signed in to perform this action',
                'code': 'NOT_AUTHENTICATED'
            }), 401
        
        # If token provided, verify it's valid (basic check)
        if token and not user_id:
            # For now, just accept any token (can be enhanced with JWT verification)
            logger.debug(f"Token-based auth attempted: {token[:20]}...")
        
        logger.info(f"Authentication successful for user: {user_id}")
        return f(*args, **kwargs)
    
    return decorated_function


@app.route('/api/health', methods=['GET'])
def health():
    """
    Health check endpoint for monitoring.
    Returns detailed system status for Render health checks.
    """
    try:
        # Check MongoDB connection
        mongo_status = mongo_client.get_connection_status()
        
        # Check disk space
        import shutil
        storage_path = app.config.get('UPLOAD_FOLDER', './data')
        try:
            disk_usage = shutil.disk_usage(storage_path)
            disk_free_gb = disk_usage.free / (1024 ** 3)
            disk_total_gb = disk_usage.total / (1024 ** 3)
            disk_percent = (disk_usage.used / disk_usage.total) * 100
        except:
            disk_free_gb = None
            disk_total_gb = None
            disk_percent = None
        
        # System info
        import psutil
        try:
            cpu_percent = psutil.cpu_percent(interval=0.1)
            memory = psutil.virtual_memory()
            memory_percent = memory.percent
        except:
            cpu_percent = None
            memory_percent = None
        
        # Overall health status
        is_healthy = (
            mongo_status.get('connected', False) and
            (disk_percent is None or disk_percent < 95) and
            (memory_percent is None or memory_percent < 90)
        )
        
        response = {
            'status': 'healthy' if is_healthy else 'degraded',
            'timestamp': get_timestamp(),
            'version': '1.0.0',
            'environment': os.getenv('FLASK_ENV', 'development'),
            'checks': {
                'mongodb': {
                    'status': 'connected' if mongo_status.get('connected') else 'disconnected',
                    'details': mongo_status
                },
                'disk': {
                    'status': 'ok' if (disk_percent is None or disk_percent < 90) else 'warning',
                    'free_gb': round(disk_free_gb, 2) if disk_free_gb else None,
                    'total_gb': round(disk_total_gb, 2) if disk_total_gb else None,
                    'used_percent': round(disk_percent, 2) if disk_percent else None
                },
                'system': {
                    'cpu_percent': round(cpu_percent, 2) if cpu_percent else None,
                    'memory_percent': round(memory_percent, 2) if memory_percent else None
                }
            }
        }
        
        # Return 200 if healthy, 503 if degraded
        status_code = 200 if is_healthy else 503
        return jsonify(response), status_code
        
    except Exception as e:
        logger.error(f"Health check failed: {e}", exc_info=True)
        return jsonify({
            'status': 'unhealthy',
            'timestamp': get_timestamp(),
            'error': str(e)
        }), 503


@app.route('/api/create-batch', methods=['POST'])
@require_api_key
@require_auth
def create_batch():
    """Create a new batch (requires user to be signed in)."""
    try:
        data = request.get_json()
        name = data.get('name', f'Batch_{datetime.utcnow().strftime("%Y%m%d_%H%M%S")}')
        user_id = data.get('user_id')
        username = data.get('username')  # Get username from request if provided
        token_user_email = None
        
        # At this point, user is authenticated (require_auth decorator checked this)
        logger.info(f"Authenticated batch creation request for user: {user_id}")
        
        # Verify user exists in database (if user_id is not 'default')
        if user_id != 'default':
            user_exists = False
            try:
                # Try to find user by email, username, or mobile
                user = mongo_client.get_user_by_email(user_id)
                if not user:
                    # Try username
                    user = mongo_client.get_user_by_username(user_id)
                if not user:
                    # Try mobile (normalize first)
                    mobile_normalized = ''.join(filter(str.isdigit, str(user_id)))
                    if len(mobile_normalized) == 10:
                        collection = mongo_client.db["User-Base"]
                        user = collection.find_one({
                            "$or": [
                                {"phoneNumber": mobile_normalized},
                                {"phoneNumber": f"+91{mobile_normalized}"},
                                {"phoneNumber": f"91{mobile_normalized}"}
                            ]
                        })
                
                if user:
                    user_exists = True
                    # Use email as the canonical user_id if available
                    if user.get('email'):
                        user_id = user.get('email')
                        token_user_email = user.get('email')
                    elif user.get('username'):
                        user_id = user.get('username')
                    elif not username and user.get('fullName'):
                        username = user.get('fullName')
                    
                    logger.info(f"[CREATE BATCH] User verified - Email: {user.get('email')}, Username: {user.get('username')}, Using user_id: {user_id}")
                else:
                    logger.warning(f"[CREATE BATCH] User not found for user_id: {user_id}")
            except Exception as e:
                logger.error(f"[CREATE BATCH] Error verifying user: {e}")
        
        # OPTIMIZED: Check if batch name already exists (use direct query instead of loading all batches)
        try:
            collection = mongo_client.db["Batch-Base"]
            existing_batch = collection.find_one({
                "user_id": user_id,
                "name": {"$regex": f"^{name}$", "$options": "i"}
            }, {"_id": 1})
            
            if existing_batch:
                return jsonify({
                    'error': f'Batch name "{name}" already exists. Please choose a different name.'
                }), 400
        except Exception as e:
            logger.warning(f"Error checking duplicate batch names: {e}")
        
        if not token_user_email and user_id and isinstance(user_id, str) and '@' in user_id:
            token_user_email = user_id

        # Debit tokens for batch creation (5 tokens per batch)
        tokens_required_for_batch = 5
        token_email_normalized = token_user_email.lower() if token_user_email else None
        if token_email_normalized:
            # First, verify the user exists and check their token balance
            user_check = mongo_client.db["User-Base"].find_one({"email": token_email_normalized})
            if user_check:
                logger.info(
                    f"[CREATE BATCH] User found in DB - tokens_balance: {user_check.get('tokens_balance')}, "
                    f"tokens_total: {user_check.get('tokens_total')}, plan_id: {user_check.get('plan_id')}"
                )
                
                # Initialize tokens if not present (for users created via phone OTP or other methods)
                if user_check.get('tokens_balance') is None or user_check.get('tokens_total') is None:
                    logger.info(f"[CREATE BATCH] Initializing tokens for user: {token_email_normalized}")
                    mongo_client.ensure_user_token_document(token_email_normalized)
                    # Re-fetch user to get updated token info
                    user_check = mongo_client.db["User-Base"].find_one({"email": token_email_normalized})
                    logger.info(
                        f"[CREATE BATCH] After initialization - tokens_balance: {user_check.get('tokens_balance')}, "
                        f"tokens_total: {user_check.get('tokens_total')}"
                    )
            else:
                logger.warning(f"[CREATE BATCH] User NOT found in DB with email: {token_email_normalized}")
            
            token_metadata = {
                "batch_name": name,
                "action": "batch_create",
                "requested_at": datetime.utcnow().isoformat()
            }
            debit_result = mongo_client.debit_tokens(
                token_email_normalized,
                tokens_required_for_batch,
                'batch:create',
                token_metadata
            )
            logger.info(f"[CREATE BATCH] Debit result: {debit_result}")
            if not debit_result.get('success'):
                error_code = debit_result.get('error') or 'TOKEN_DEBIT_FAILED'
                status_map = {
                    'INSUFFICIENT_TOKENS': 402,
                    'USER_NOT_FOUND': 404,
                    'DB_UNAVAILABLE': 503
                }
                status_code = status_map.get(error_code, 400)
                logger.warning(
                    "[CREATE BATCH] Token debit failed for %s (code=%s)",
                    token_email_normalized,
                    error_code
                )
                messages = {
                    'INSUFFICIENT_TOKENS': 'You need at least 5 tokens to create a batch.',
                    'USER_NOT_FOUND': 'User account not found. Please sign in again.',
                    'DB_UNAVAILABLE': 'Token service is temporarily unavailable. Try again later.'
                }
                return jsonify({
                    'error': error_code,
                    'message': messages.get(error_code, 'Unable to create batch due to token debit failure.')
                }), status_code
        else:
            logger.warning("[CREATE BATCH] Skipping token debit; unable to resolve email for user_id=%s", user_id)

        batch_id = str(uuid.uuid4())
        batch_doc = mongo_client.create_batch(batch_id, name, user_id, username)
        
        logger.info(f"[CREATE BATCH] Created batch '{name}' (ID: {batch_id}) for user_id: {user_id}")
        
        # Invalidate cache so new batch appears immediately
        invalidate_batch_cache(user_id)
        
        return jsonify({
            'batch_id': batch_id,
            'name': name,
            'created_at': batch_doc.get('created_at', get_timestamp()),
            'message': 'Batch created successfully'
        }), 201
    
    except Exception as e:
        logger.error(f"Error creating batch: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/upload-v2', methods=['POST', 'OPTIONS'])
def upload_files_v2():
    """New upload endpoint - simpler, bypass Werkzeug validation issues."""
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        logger.info("=" * 80)
        logger.info("🚀 UPLOAD V2 ENDPOINT REACHED!")
        logger.info(f"Method: {request.method}")
        logger.info(f"Path: {request.path}")
        logger.info(f"Content-Type: {request.content_type}")
        logger.info(f"Content-Length: {request.content_length}")
        logger.info("=" * 80)
        
        # Get parameters from query string
        batch_id = request.args.get('batch_id')
        user_id = request.args.get('user_id')
        api_key = request.headers.get('X-API-KEY')
        
        logger.info(f"   batch_id={batch_id}, user_id={user_id}, api_key_present={bool(api_key)}")
        
        # Validate auth
        expected_key = os.getenv('API_KEY', '')
        if expected_key and api_key != expected_key:
            return jsonify({'error': 'Invalid API key'}), 401
        
        if not user_id or user_id == 'default':
            return jsonify({'error': 'User not authenticated'}), 401
        
        if not batch_id:
            return jsonify({'error': 'batch_id required'}), 400
        
        # Check batch exists
        batch = mongo_client.get_batch(batch_id)
        if not batch:
            return jsonify({'error': 'Batch not found'}), 404
        
        logger.info(f"   Batch found: {batch_id}")
        
        # Get files from request
        files = request.files.getlist('files[]') or request.files.getlist('files') or []
        
        logger.info(f"   Files received: {len(files)}")
        
        if not files:
            return jsonify({'error': 'No files provided'}), 400
        
        # Process files
        job_id = str(uuid.uuid4())
        ensure_dir(app.config['UPLOAD_FOLDER'])
        
        # Get already processed files
        already_processed = set()
        if batch.get('files'):
            already_processed = {f.get('filename', '') for f in batch.get('files', [])}
        
        # Save and prepare files
        file_infos = []
        skipped_files = []
        for file in files:
            if not file or not file.filename:
                continue
            if not allowed_file(file.filename):
                continue
            
            filename = sanitize_filename(file.filename)
            if filename in already_processed:
                skipped_files.append(filename)
                continue
            
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{job_id}_{filename}")
            ensure_dir(os.path.dirname(filepath))
            file.save(filepath)
            
            file_infos.append({
                'filepath': filepath,
                'filename': filename,
                'batch_id': batch_id,
                'job_id': job_id,
                'file_type': os.path.splitext(filename)[1].lower()
            })
        
        if not file_infos:
            if skipped_files:
                return jsonify({
                    'error': 'All files already processed',
                    'skipped_count': len(skipped_files),
                    'skipped_files': skipped_files
                }), 400
            return jsonify({'error': 'No valid files'}), 400
        
        logger.info(f"   Files saved: {len(file_infos)}, job_id={job_id}")
        
        # Initialize job
        import time
        with _jobs_lock:
            jobs[job_id] = {
                'job_id': job_id,
                'batch_id': batch_id,
                'status': 'processing',
                'files': file_infos,
                'results': [],
                'created_at': get_timestamp(),
                'start_time': time.time(),
                'progress': {
                    'completed': 0,
                    'total': len(file_infos),
                    'current_file': None,
                    'eta_seconds': len(file_infos) * 5
                }
            }
        
        # Start processing in background
        def process_job():
            try:
                import time
                start_time = time.time()
                processor = get_processor()
                processing_times = []
                
                def progress_callback(completed, total, filename, eta=None):
                    """Update progress."""
                    with _jobs_lock:
                        if job_id in jobs:
                            jobs[job_id]['progress']['completed'] = completed
                            jobs[job_id]['progress']['total'] = total
                            jobs[job_id]['progress']['current_file'] = filename
                            if eta is not None:
                                jobs[job_id]['progress']['eta_seconds'] = eta
                            elif completed > 0 and total > completed:
                                elapsed = time.time() - start_time
                                avg_time = elapsed / completed
                                remaining = total - completed
                                jobs[job_id]['progress']['eta_seconds'] = max(1, int(avg_time * remaining * 1.1))
                            elif completed >= total:
                                jobs[job_id]['progress']['eta_seconds'] = 0
                
                # Process files
                results = processor.process_batch(file_infos, callback=progress_callback)
                
                logger.info(f"Processing complete: {len(results)} files processed")
                
                # Save results to MongoDB
                successful_results = []
                total_processing_time = 0.0
                for result in results:
                    if result.get('success'):
                        filename = result.get('filename', '')
                        piis = result.get('piis', [])
                        pii_count = len(piis) if isinstance(piis, list) else 0
                        
                        # Log PII detection results
                        logger.info(f"💾 Saving file {filename} to batch {batch_id}: {pii_count} PIIs detected")
                        if pii_count > 0:
                            pii_types = list(set([p.get('type', 'UNKNOWN') for p in piis if isinstance(p, dict)]))
                            logger.info(f"  PII types: {pii_types}")
                            # Log sample PII structure
                            if len(piis) > 0 and isinstance(piis[0], dict):
                                logger.info(f"  Sample PII structure: {list(piis[0].keys())}")
                                logger.info(f"  Sample PII: type={piis[0].get('type')}, value={str(piis[0].get('value') or piis[0].get('match', ''))[:50]}, has_bbox={('bbox' in piis[0])}")
                                if 'bbox' in piis[0]:
                                    logger.info(f"  Sample bbox: {piis[0]['bbox']}")
                        else:
                            logger.warning(f"  ⚠️ No PIIs detected in {filename}")
                        
                        # Add file to batch
                        save_success = mongo_client.add_file_to_batch(
                            batch_id,
                            filename,
                            {
                                'pii_count': pii_count,
                                'page_count': result.get('page_count', 0),
                                'piis': piis,
                                'processed_at': result.get('timestamp'),
                                'processing_time': result.get('processing_time')
                            }
                        )
                        if save_success:
                            logger.info(f"  ✓ File {filename} saved to MongoDB successfully")
                        else:
                            logger.error(f"  ❌ Failed to save file {filename} to MongoDB!")
                        successful_results.append(result)
                        total_processing_time += float(result.get('processing_time') or 0.0)
                
                # Update batch stats
                pii_results = {'files': successful_results}
                mongo_client.update_batch_stats(batch_id, len(successful_results), pii_results, scan_duration=total_processing_time)

                # Debit tokens for processed files
                tokens_per_file = 2
                tokens_to_debit = len(successful_results) * tokens_per_file
                if tokens_to_debit > 0 and user_id:
                    user_email = str(user_id).strip().lower()
                    try:
                        token_metadata = {
                            "batch_id": batch_id,
                            "job_id": job_id,
                            "files_processed": len(successful_results),
                            "tokens_per_file": tokens_per_file
                        }
                        debit_result = mongo_client.debit_tokens(
                            user_email,
                            tokens_to_debit,
                            'scan:file_processing',
                            token_metadata
                        )
                        if debit_result.get('success'):
                            logger.info(
                                "Tokens debited for user %s: %s (balance=%s)",
                                user_email,
                                tokens_to_debit,
                                debit_result.get('balance')
                            )
                        else:
                            logger.warning(
                                "Failed to debit tokens for user %s: %s",
                                user_email,
                                debit_result.get('error')
                            )
                    except Exception as token_error:
                        logger.error(
                            "Error debiting tokens for user %s: %s",
                            user_email,
                            token_error,
                            exc_info=True
                        )
                
                # Update job status
                with _jobs_lock:
                    if job_id in jobs:
                        jobs[job_id]['status'] = 'completed'
                        jobs[job_id]['results'] = results
                        jobs[job_id]['completed_at'] = get_timestamp()
                
                logger.info(f"Job {job_id} completed with {len(results)} results")
            
            except Exception as e:
                logger.error(f"Job {job_id} error: {e}", exc_info=True)
                with _jobs_lock:
                    if job_id in jobs:
                        jobs[job_id]['status'] = 'failed'
                        jobs[job_id]['error'] = str(e)
        
        # Start in background thread
        threading.Thread(target=process_job, daemon=True).start()
        
        return jsonify({
            'job_id': job_id,
            'batch_id': batch_id,
            'file_count': len(file_infos),
            'skipped_count': len(skipped_files),
            'skipped_files': skipped_files,
            'status': 'processing'
        }), 202
    
    except Exception as e:
        logger.error(f"Upload V2 error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# Keep old upload endpoint for backwards compatibility, but make it simpler
@app.route('/api/upload', methods=['POST', 'OPTIONS'])
def upload_files():
    """Upload and process files - delegates to V2."""
    return upload_files_v2()


# ------------------------------------------------------------
# Image OCR + PII Detection Endpoint
# ------------------------------------------------------------

@app.route('/api/pii/image/extract', methods=['POST', 'OPTIONS'])
@require_api_key
def extract_pii_from_images():
    """
    Extract text from images using OCR and detect PIIs
    Supports: PNG, JPG, JPEG, SVG
    Parallel processing for multiple images
    """
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        # Check if files are uploaded
        if 'files' not in request.files:
            return jsonify({'error': 'No files uploaded'}), 400
        
        files = request.files.getlist('files')
        if not files or len(files) == 0:
            return jsonify({'error': 'No files provided'}), 400
        
        logger.info(f"🖼️ Received {len(files)} images for OCR + PII detection")
        
        # Validate file types
        allowed_extensions = {'.png', '.jpg', '.jpeg', '.svg'}
        images_data = []
        
        for file in files:
            filename = secure_filename(file.filename)
            ext = os.path.splitext(filename)[1].lower()
            
            if ext not in allowed_extensions:
                return jsonify({
                    'error': f'Invalid file type: {ext}. Allowed: {", ".join(allowed_extensions)}'
                }), 400
            
            # Read file bytes
            file_bytes = file.read()
            
            if len(file_bytes) == 0:
                return jsonify({'error': f'Empty file: {filename}'}), 400
            
            images_data.append((file_bytes, filename))
        
        # Import pipeline
        try:
            from image_ocr_pipeline import get_pipeline
            from masker import PIIDetector as ExistingPIIDetector
            
            # Initialize PII detector
            pii_detector_instance = ExistingPIIDetector()
            
            # Get pipeline instance
            pipeline = get_pipeline(pii_detector_instance)
            
            # Process all images in parallel
            results = pipeline.process_multiple_images(
                images_data, 
                max_workers=min(len(images_data), 4)  # Max 4 parallel workers
            )
            
            # Format response
            response = {
                'success': True,
                'total_images': len(results),
                'results': [result.to_dict() for result in results],
                'summary': {
                    'total_piis_found': sum(r.total_piis for r in results),
                    'images_with_piis': sum(1 for r in results if r.total_piis > 0),
                    'total_processing_time': sum(r.processing_time for r in results)
                }
            }
            
            logger.info(f"✅ Successfully processed {len(results)} images, found {response['summary']['total_piis_found']} PIIs")
            
            return jsonify(response), 200
            
        except ImportError as e:
            logger.error(f"Image OCR pipeline not available: {e}")
            return jsonify({
                'error': 'Image OCR module not installed. Please install: pip install paddleocr opencv-python cairosvg pytesseract'
            }), 500
        
    except Exception as e:
        logger.error(f"Image OCR processing failed: {e}", exc_info=True)
        return jsonify({'error': f'Processing failed: {str(e)}'}), 500


@app.route('/api/pii/image/mask', methods=['POST', 'OPTIONS'])
@require_api_key
def mask_pii_in_images():
    """
    Mask detected PIIs in images
    Supports: blackout, hash, blur, pixelate masking modes
    """
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        # Get form data
        if 'files' not in request.files:
            return jsonify({'error': 'No files uploaded'}), 400
        
        files = request.files.getlist('files')
        if not files or len(files) == 0:
            return jsonify({'error': 'No files provided'}), 400
        
        # Get PII detection results from form data (JSON string)
        pii_results_json = request.form.get('pii_results')
        if not pii_results_json:
            return jsonify({'error': 'PII detection results required'}), 400
        
        try:
            pii_results = json.loads(pii_results_json)
        except json.JSONDecodeError:
            return jsonify({'error': 'Invalid PII results JSON'}), 400
        
        logger.warning("Image masking request received, but image masking is disabled. Returning error.")
        return jsonify({
            'success': False,
            'error': 'IMAGE_MASKING_DISABLED',
            'message': 'Image masking is currently disabled for uploaded image files.'
        }), 400
        
    except Exception as e:
        logger.error(f"Image masking failed: {e}", exc_info=True)
        return jsonify({'error': f'Masking failed: {str(e)}'}), 500


# ------------------------------------------------------------
# Token consumption endpoint
# ------------------------------------------------------------

@app.route('/api/action/consume-token', methods=['POST'])
@require_api_key
@require_auth
def consume_token_action():
    """Deduct tokens for a metered action (lock_json, unlock_json, download_masked_file, etc.)."""
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        logger.error("consume_token_action: invalid JSON payload")
        return jsonify({'success': False, 'error': 'INVALID_PAYLOAD'}), 400

    action_raw = data.get('action')
    action = str(action_raw).strip().lower() if action_raw else ''
    if not action:
        return jsonify({'success': False, 'error': 'ACTION_REQUIRED'}), 400

    if action not in TOKEN_ACTION_COSTS:
        logger.warning(f"consume_token_action: Unsupported action '{action}'")
        return jsonify({'success': False, 'error': 'UNSUPPORTED_ACTION'}), 422

    email = data.get('email') or data.get('user_id')
    if not email:
        return jsonify({'success': False, 'error': 'EMAIL_REQUIRED'}), 400
    email = str(email).strip().lower()

    # Check plan restrictions for lock_json and unlock_json (Professional/Enterprise only)
    if action in ['lock_json', 'unlock_json']:
        user = mongo_client.get_user_by_email(email)
        if not user:
            return jsonify({'success': False, 'error': 'USER_NOT_FOUND'}), 404
        
        user_plan = user.get('plan_id', 'starter').lower()
        if user_plan not in ['professional', 'enterprise']:
            logger.warning(f"consume_token_action: User {email} with plan '{user_plan}' attempted to use {action}")
            return jsonify({
                'success': False, 
                'error': 'PLAN_RESTRICTION',
                'message': f'{action.replace("_", " ").title()} is only available on Professional and Enterprise plans.'
            }), 403

    tokens_required = TOKEN_ACTION_COSTS[action]
    metadata = data.get('metadata') or {}
    if not isinstance(metadata, dict):
        metadata = {}
    metadata.update({
        "action": action,
        "request_id": str(uuid.uuid4())
    })

    logger.info(f"consume_token_action: action={action}, tokens={tokens_required}, email={email}")

    result = mongo_client.debit_tokens(email, tokens_required, f'action:{action}', metadata)
    if result.get('success'):
        # Invalidate profile cache to reflect updated token balance
        invalidate_profile_cache(email)
        
        response_payload = {
            'success': True,
            'balance': result.get('balance'),
            'unlimited': result.get('unlimited', False),
            'cost': tokens_required
        }
        logger.info(f"consume_token_action: success for {email}, balance={result.get('balance')}, unlimited={result.get('unlimited')}")
        return jsonify(response_payload), 200

    error_code = result.get('error') or 'UNKNOWN_ERROR'
    status_map = {
        'INSUFFICIENT_TOKENS': 402,
        'USER_NOT_FOUND': 404,
        'DB_UNAVAILABLE': 503
    }
    status_code = status_map.get(error_code, 400)
    logger.warning(f"consume_token_action: failed for {email} - {error_code}")
    return jsonify({'success': False, 'error': error_code}), status_code


# ===== Old code removed - was causing issues =====
# See upload_files_v2() above for the current implementation
# ===== End removal =====


@app.route('/api/job-status', methods=['GET'])
@require_api_key
def get_job_status():
    """Get job processing status."""
    try:
        job_id = request.args.get('job_id')
        if not job_id:
            return jsonify({'error': 'job_id required'}), 400
        
        with _jobs_lock:
            if job_id not in jobs:
                return jsonify({'error': 'Job not found'}), 404
            
            job = jobs[job_id]
            return jsonify({
                'job_id': job_id,
                'status': job.get('status'),
                'progress': job.get('progress'),
                'results': job.get('results', []),
                'error': job.get('error')
            }), 200
    
    except Exception as e:
        logger.error(f"Error getting job status: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# ===== ORPHANED CODE REMOVED - WAS CAUSING SYNTAX ERRORS =====
# All the old upload code has been removed and replaced with upload_files_v2()
# ===== END REMOVAL =====


# Orphaned code removed - old upload logic replaced with upload_files_v2()


@app.route('/api/job-result', methods=['GET'])
@require_api_key
def get_job_result():
    """Get the status and results of a processing job."""
    try:
        job_id = request.args.get('job_id')
        if not job_id:
            return jsonify({'error': 'job_id is required'}), 400
        
        with _jobs_lock:
            job = jobs.get(job_id)
        if not job:
            return jsonify({'error': 'Job not found'}), 404
        
        # Return a copy to prevent external modification
        return jsonify(job.copy()), 200
    
    except Exception as e:
        logger.error(f"Error getting job result: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/mask', methods=['POST'])
@require_api_key
def mask_files():
    """Apply masking to processed files."""
    try:
        data = request.get_json()
        job_id = data.get('job_id')
        batch_id = data.get('batch_id')  # Allow masking by batch_id for previous batches
        mask_type = data.get('mask_type', 'blur')  # 'hash' or 'blur'
        password = data.get('password') if mask_type == 'hash' else None
        
        if not job_id and not batch_id:
            return jsonify({'error': 'job_id or batch_id is required'}), 400
        
        if mask_type == 'hash' and not password:
            return jsonify({'error': 'password is required for hash masking'}), 400
        
        # Store encryption password in batch document if provided
        if mask_type == 'hash' and password and batch_id:
            try:
                batch = mongo_client.get_batch(batch_id)
                if batch and mongo_client.client is not None and mongo_client.db is not None:
                    collection = mongo_client.db.batches
                    collection.update_one(
                        {"batch_id": batch_id},
                        {"$set": {"encryption_password": password, "updated_at": datetime.utcnow()}}
                    )
                    logger.info(f"Stored encryption password for batch {batch_id}")
            except Exception as e:
                logger.warning(f"Failed to store encryption password: {e}")
        
        # Get results from job (if job_id provided) or from batch analysis (if batch_id provided)
        results = []
        if job_id:
            with _jobs_lock:
                job = jobs.get(job_id)
            if not job or job.get('status') != 'completed':
                return jsonify({'error': 'Job not found or not completed'}), 404
            batch_id = job.get('batch_id')
            results = job.get('results', [])
        elif batch_id:
            # Load results from MongoDB for this batch
            logger.info(f"Loading batch for masking: {batch_id}")
            batch = mongo_client.get_batch_analysis(batch_id)  # Use get_batch_analysis for proper serialization
            if not batch:
                logger.error(f"Batch not found: {batch_id}")
                return jsonify({'error': 'Batch not found'}), 404
            
            batch_files = batch.get('files', [])
            logger.info(f"📂 Loading batch data for masking: batch {batch_id} with {len(batch_files)} files")
            
            # Convert batch file entries to result format for masking
            for file_idx, file_entry in enumerate(batch_files):
                filename = file_entry.get('filename', '')
                piis = file_entry.get('piis', [])
                
                logger.info(f"  File {file_idx}: {filename}")
                logger.info(f"    PIIs type: {type(piis).__name__}, length: {len(piis) if isinstance(piis, list) else 'N/A'}")
                
                # DEBUG: Log first few PIIs to see their structure
                if isinstance(piis, list) and len(piis) > 0:
                    for i, pii in enumerate(piis[:3]):
                        logger.info(f"    PII {i+1}: {pii.keys() if isinstance(pii, dict) else type(pii)}")
                        if isinstance(pii, dict):
                            logger.info(f"      type={pii.get('type')}, has_bbox={('bbox' in pii)}, bbox={pii.get('bbox')}")
                
                if filename:
                    # Ensure piis is a list
                    if not isinstance(piis, list):
                        logger.warning(f"    ⚠️ PIIs is not a list, attempting conversion...")
                        piis = list(piis) if hasattr(piis, '__iter__') else []
                    
                    pii_count = len(piis) if isinstance(piis, list) else 0
                    logger.info(f"    ✓ Added to queue: {pii_count} PIIs")
                    
                    # Create a result object from the batch file entry
                    result = {
                        'success': True,
                        'filename': filename,
                        'piis': piis if isinstance(piis, list) else [],
                        'pii_count': pii_count,
                        'page_count': file_entry.get('page_count', 0),
                        'text_length': 0  # Not available from batch
                    }
                    results.append(result)
            
            logger.info(f"✅ Prepared {len(results)} files from batch for masking")
        
        if not results:
            return jsonify({'error': 'No processed files found for masking'}), 404
        
        masked_files = []
        masked_results = []
        skipped_files = []
        
        def normalize_bbox(raw_bbox):
            """Convert bbox into dict{x,y,width,height} if possible."""
            if not raw_bbox:
                return None
            if isinstance(raw_bbox, dict):
                if all(key in raw_bbox for key in ['x', 'y', 'width', 'height']):
                    return {
                        'x': int(raw_bbox['x']),
                        'y': int(raw_bbox['y']),
                        'width': int(raw_bbox['width']),
                        'height': int(raw_bbox['height'])
                    }
                return None
            if isinstance(raw_bbox, (list, tuple)) and len(raw_bbox) == 4:
                return {
                    'x': int(raw_bbox[0]),
                    'y': int(raw_bbox[1]),
                    'width': int(raw_bbox[2]),
                    'height': int(raw_bbox[3])
                }
            return None

        logger.info(f"Starting masking process for {len(results)} files")
        
        for result in results:
            if not result.get('success'):
                logger.warning(f"Skipping {result.get('filename', 'unknown')}: result marked as unsuccessful")
                skipped_files.append(result.get('filename', 'unknown'))
                continue
            
            filename = result['filename']
            original_path = None
            
            # Find original file - try multiple locations
            if job_id and job:
                # From current job
                for file_info in job.get('files', []):
                    if file_info['filename'] == filename:
                        original_path = file_info['filepath']
                        logger.info(f"Found original file from job: {original_path}")
                        break
            
            # If not found, check if filepath is stored in result JSON (this is the most reliable)
            if not original_path or not os.path.exists(original_path):
                stored_path = result.get('filepath') or result.get('original_path')
                if stored_path:
                    # Handle both absolute and relative paths
                    if os.path.isabs(stored_path):
                        if os.path.exists(stored_path):
                            original_path = stored_path
                            logger.info(f"Found original file from result JSON (absolute): {original_path}")
                    else:
                        # Try relative to uploads folder
                        rel_path = os.path.join(app.config['UPLOAD_FOLDER'], stored_path)
                        if os.path.exists(rel_path):
                            original_path = rel_path
                            logger.info(f"Found original file from result JSON (relative): {original_path}")
                        elif os.path.exists(stored_path):
                            original_path = stored_path
                            logger.info(f"Found original file from result JSON: {original_path}")
            
            # If still not found, search in uploads folder
            if not original_path or not os.path.exists(original_path):
                uploads_folder = app.config['UPLOAD_FOLDER']
                # Try exact filename match first
                exact_path = os.path.join(uploads_folder, filename)
                if os.path.exists(exact_path):
                    original_path = exact_path
                    logger.info(f"Found original file (exact match): {original_path}")
                else:
                    # Search for files ending with filename (job_id_prefix format)
                    if os.path.exists(uploads_folder):
                        matching_files = []
                        # Also search for files starting with batch_id or containing filename
                        for file in os.listdir(uploads_folder):
                            file_path = os.path.join(uploads_folder, file)
                            if not os.path.isfile(file_path):
                                continue
                            # Files are stored as: {job_id}_{filename} or just {filename}
                            # Check multiple patterns to find the file
                            if (file == filename or  # Exact match
                                file.endswith(f"_{filename}") or  # job_id_filename format
                                (filename in file and file.count('_') >= 1)):  # Contains filename with underscore separator
                                matching_files.append((file_path, file))
                        
                        # If multiple matches, prefer the most recent one or one that matches batch_id
                        if matching_files:
                            # Sort by modification time (most recent first)
                            matching_files.sort(key=lambda x: os.path.getmtime(x[0]), reverse=True)
                            original_path = matching_files[0][0]
                            logger.info(f"Found original file (pattern match from {len(matching_files)} candidates): {original_path}")
                        else:
                            # Fallback: if filename appears anywhere in the uploads folder
                            for file in os.listdir(uploads_folder):
                                file_path = os.path.join(uploads_folder, file)
                                if os.path.isfile(file_path) and filename in file:
                                    original_path = file_path
                                    logger.info(f"Found original file (fallback match): {original_path}")
                                    break
            
            if not original_path or not os.path.exists(original_path):
                logger.error(f"Original file not found: {filename}")
                logger.error(f"Searched in: {app.config['UPLOAD_FOLDER']}")
                logger.error(f"Result keys: {list(result.keys())}")
                skipped_files.append(filename)
                continue
            
            logger.info(f"Processing masking for file: {filename}")
            
            # Determine output path
            output_filename = f"masked_{filename}"
            output_path = os.path.join(
                app.config['MASKED_FOLDER'],
                batch_id,
                output_filename
            )
            ensure_dir(os.path.dirname(output_path))
            
            # Apply masking based on file type
            piis = result.get('piis', [])
            
            # Filter PIIs by selected types if provided
            selected_pii_types = data.get('selected_pii_types')
            normalized_selected_types: List[str] = []
            if selected_pii_types and len(selected_pii_types) > 0:
                selected_set = {str(p).strip().lower() for p in selected_pii_types}
                normalized_selected_types = list(selected_set)
                piis = [
                    pii for pii in piis
                    if str(pii.get('type', '')).strip().lower() in selected_set
                ]
                logger.info(f"Filtered to {len(piis)} PIIs from selected types: {selected_pii_types}")
            
            if not piis:
                logger.warning(f"No PIIs to mask for {filename} (after filtering)")
                continue
            
            try:
                # Store hash_meta mapping for decryption
                hash_meta_map = {}
                
                # Mask the file and get hash_meta_map
                # For text/CSV files, hash_meta_map MUST come FROM masking (each encryption is unique)
                # For other files, we create it beforehand
                if is_pdf_file(filename):
                    if mask_type == 'hash' and password:
                        # Create mapping for PDF (before masking)
                        for pii in piis:
                            pii_value = pii.get('value', '')
                            if pii_value:
                                hash_result = masker.hash_mask(pii_value, password)
                                encrypted_value = hash_result['masked_value']
                                hash_meta_map[encrypted_value] = {
                                    'hash_meta': hash_result['hash_meta'],
                                    'original_value': pii_value,
                                    'pii_type': pii.get('type', ''),
                                    'page': pii.get('page', 0)
                                }
                    masker.mask_pdf(original_path, piis, output_path, mask_type, password)
                elif is_image_file(filename):
                    logger.info(f"Skipping image file {filename}: image masking is disabled")
                    skipped_files.append(filename)
                    continue
                elif is_docx_file(filename):
                    if mask_type == 'hash' and password:
                        # Create mapping for DOCX (before masking)
                        for pii in piis:
                            pii_value = pii.get('value', '')
                            if pii_value:
                                hash_result = masker.hash_mask(pii_value, password)
                                encrypted_value = hash_result['masked_value']
                                hash_meta_map[encrypted_value] = {
                                    'hash_meta': hash_result['hash_meta'],
                                    'original_value': pii_value,
                                    'pii_type': pii.get('type', ''),
                                    'page': pii.get('page', 0)
                                }
                    masker.mask_docx(original_path, piis, output_path, mask_type, password)
                else:
                    # Text/CSV files: hash_meta_map comes FROM masking function
                    # CRITICAL: Don't create it beforehand - each hash_mask() call creates unique encrypted value
                    text_file_result = masker.mask_text_file(original_path, piis, output_path, mask_type, password)
                    if text_file_result and isinstance(text_file_result, dict):
                        # Use hash_meta_map from text file masking (created during masking)
                        if 'hash_meta_map' in text_file_result:
                            hash_meta_map = text_file_result['hash_meta_map']
                
                # Save hash_meta mapping to a separate JSON file for decryption
                if mask_type == 'hash' and password and hash_meta_map:
                    hash_meta_path = os.path.join(
                        app.config['RESULTS_FOLDER'],
                        batch_id,
                        f"{filename}_hash_meta.json"
                    )
                    ensure_dir(os.path.dirname(hash_meta_path))
                    save_json(hash_meta_map, hash_meta_path)
                    logger.info(f"Saved hash_meta mapping for {filename}")
                
                masked_files.append(output_path)
                masked_results.append({
                    'filename': filename,
                    'masked_filename': output_filename,
                    'original_filename': filename,
                    'path': output_path,
                    'mask_type': mask_type
                })
            
            except Exception as e:
                logger.error(f"Error masking {filename}: {e}")
                continue
        
        if not masked_files:
            error_msg = f'No files were masked. Processed {len(results)} files.'
            if skipped_files:
                error_msg += f' Skipped files: {", ".join(skipped_files)}'
            logger.error(error_msg)
            return jsonify({'error': error_msg}), 500
        
        logger.info(f"Successfully masked {len(masked_files)} files. Skipped: {len(skipped_files)} files")
        
        # If more than 10 files, create zip
        if len(masked_files) > 10:
            # Use batch_id for zip name if no job_id
            zip_name = f"masked_{job_id if job_id else batch_id}.zip"
            zip_path = os.path.join(
                app.config['MASKED_FOLDER'],
                batch_id,
                zip_name
            )
            create_zip(masked_files, zip_path)
            # Use relative path for download URL
            base_path = os.getenv('STORAGE_PATH', './data')
            rel_path = os.path.relpath(zip_path, base_path)
            return jsonify({
                'job_id': job_id or batch_id,
                'batch_id': batch_id,
                'mask_type': mask_type,
                'file_count': len(masked_files),
                'download_url': f'/api/download?path={rel_path}',
                'format': 'zip'
            }), 200
        else:
            # Return individual file URLs with relative paths and file info
            base_path = os.getenv('STORAGE_PATH', './data')
            download_urls = []
            file_info_list = []
            for masked_result in masked_results:
                rel_path = os.path.relpath(masked_result['path'], base_path)
                download_urls.append(f'/api/download?path={rel_path}')
                # Create descriptive filename for tooltip
                original_name = masked_result['original_filename']
                name_parts = os.path.splitext(original_name)
                mask_type_label = 'blurred' if mask_type == 'blur' else 'hashed'
                descriptive_name = f"{name_parts[0]}-masked-{mask_type_label}{name_parts[1]}"
                file_info_list.append({
                    'url': f'/api/download?path={rel_path}',
                    'original_filename': original_name,
                    'descriptive_name': descriptive_name,
                    'mask_type': mask_type
                })
            response_data = {
                'job_id': job_id or batch_id,
                'batch_id': batch_id,
                'mask_type': mask_type,
                'file_count': len(masked_files),
                'download_urls': download_urls,
                'file_info': file_info_list,  # Add file info for tooltips
                'format': 'individual'
            }
            logger.info(f"Returning masking response: {len(download_urls)} download URLs for {len(masked_files)} files")
            logger.info(f"Download URLs: {download_urls}")
            if skipped_files:
                logger.warning(f"Skipped files during masking: {skipped_files}")
                response_data['skipped_files'] = skipped_files
            return jsonify(response_data), 200
    
    except Exception as e:
        logger.error(f"Error masking files: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/decrypt-upload', methods=['POST'])
@require_api_key
def decrypt_uploaded_files():
    """Decrypt uploaded masked files and return decrypted PIIs."""
    try:
        if 'files[]' not in request.files:
            return jsonify({'error': 'No files uploaded'}), 400
        
        password = request.form.get('password')
        if not password:
            return jsonify({'error': 'Password is required'}), 400
        
        files = request.files.getlist('files[]')
        if not files:
            return jsonify({'error': 'No files provided'}), 400
        
        # Pre-load all hash_meta files for faster lookup (optimization)
        hash_meta_cache = {}
        results_base = app.config['RESULTS_FOLDER']
        
        if os.path.exists(results_base):
            for batch_folder in os.listdir(results_base):
                batch_path = os.path.join(results_base, batch_folder)
                if not os.path.isdir(batch_path):
                    continue
                
                meta_files = [f for f in os.listdir(batch_path) if f.endswith('_hash_meta.json')]
                for meta_file in meta_files:
                    meta_path = os.path.join(batch_path, meta_file)
                    try:
                        hash_meta_data = load_json(meta_path)
                        hash_meta_cache[meta_path] = hash_meta_data
                    except Exception as e:
                        logger.debug(f"Error loading hash_meta file {meta_file}: {e}")
        
        decrypted_results = []
        all_decrypted_piis = []
        
        for file in files:
            if file.filename == '':
                continue
            
            filename = secure_filename(file.filename)
            logger.info(f"Processing file for decryption: {filename}")
            
            # Save uploaded file temporarily
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{uuid.uuid4()}_{filename}")
            ensure_dir(os.path.dirname(temp_path))
            file.save(temp_path)
            
            try:
                # Extract text from file
                text = ""
                if is_pdf_file(filename):
                    import fitz
                    doc = fitz.open(temp_path)
                    for page in doc:
                        text += page.get_text()
                    doc.close()
                elif is_docx_file(filename):
                    from docx import Document
                    doc = Document(temp_path)
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                elif is_doc_file(filename):
                    # DOC files need special handling - convert to text using python-docx2txt or antiword
                    # For now, try to extract using textract or fallback to error message
                    try:
                        import docx2txt
                        text = docx2txt.process(temp_path)
                    except ImportError:
                        try:
                            # Try using textract if available
                            import textract
                            text = textract.process(temp_path).decode('utf-8')
                        except (ImportError, Exception):
                            logger.warning(f"DOC file support requires docx2txt or textract. Install: pip install docx2txt")
                            # Fallback: try reading as binary and extract readable text
                            with open(temp_path, 'rb') as f:
                                content = f.read()
                                # Simple extraction of readable ASCII text
                                text = ''.join(chr(b) if 32 <= b < 127 else ' ' for b in content)
                elif is_text_file(filename):
                    # Read CSV and text files directly as text to preserve exact encrypted values
                    # Don't use pandas for CSV during decryption - we need the raw encrypted strings
                    with open(temp_path, 'r', encoding='utf-8', errors='replace') as f:
                        text = f.read()
                elif is_image_file(filename):
                    # Use OCR to extract text from images
                    try:
                        from ocr_engine import get_ocr_engine
                        import cv2
                        import numpy as np
                        import io
                        ocr_engine = get_ocr_engine()
                        if filename.lower().endswith('.svg'):
                            # SVG needs special handling - extract text from SVG elements or convert to PNG
                            try:
                                from PIL import Image
                                try:
                                    import cairosvg
                                    png_data = cairosvg.svg2png(url=temp_path)
                                    image = Image.open(io.BytesIO(png_data))
                                    img_array = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
                                    extracted_text, _ = ocr_engine.extract_text(img_array)
                                    text = extracted_text
                                except ImportError:
                                    # Fallback: extract text from SVG XML
                                    logger.info("cairosvg not available, extracting text from SVG XML")
                                    with open(temp_path, 'r', encoding='utf-8', errors='replace') as f:
                                        svg_content = f.read()
                                    # Extract text from SVG text elements
                                    import re
                                    text_elements = re.findall(r'<text[^>]*>(.*?)</text>', svg_content, re.DOTALL)
                                    text = ' '.join(text_elements)
                                    if not text:
                                        logger.warning(f"No text found in SVG: {filename}")
                                        os.remove(temp_path)
                                        continue
                            except Exception as e:
                                logger.warning(f"Error processing SVG {filename}: {e}")
                                os.remove(temp_path)
                                continue
                        else:
                            img_array = cv2.imread(temp_path)
                            if img_array is not None:
                                extracted_text, _ = ocr_engine.extract_text(img_array)
                                text = extracted_text
                            else:
                                logger.warning(f"Could not read image file: {filename}")
                                os.remove(temp_path)
                                continue
                        
                        if not text:
                            logger.warning(f"No text extracted from image: {filename}")
                            os.remove(temp_path)
                            continue
                    except Exception as e:
                        logger.error(f"Error extracting text from image {filename}: {e}")
                        os.remove(temp_path)
                        continue
                else:
                    logger.warning(f"Unsupported file type for decryption: {filename}")
                    os.remove(temp_path)
                    continue
                
                # NEW APPROACH: Decrypt entire file first, then detect PIIs
                import re
                import base64
                
                # Step 1: Find all potential encrypted values (base64 strings)
                base64_pattern = r'[A-Za-z0-9+/]{16,300}={0,2}'
                all_matches = re.findall(base64_pattern, text)
                potential_encrypted = list(set([m for m in all_matches if len(m) >= 16]))
                
                # Also find base64 strings with delimiters (for CSV)
                base64_with_delimiters = re.findall(r'(?:^|[\s,\t\n\r"\'|])[A-Za-z0-9+/]{16,300}={0,2}(?:[\s,\t\n\r"\'|]|$)', text)
                for match in base64_with_delimiters:
                    cleaned = re.sub(r'^[\s,\t\n\r"\'|]+|[\s,\t\n\r"\'|]+$', '', match)
                    if len(cleaned) >= 16 and cleaned not in potential_encrypted:
                        potential_encrypted.append(cleaned)
                
                logger.info(f"Found {len(potential_encrypted)} potential encrypted values in {filename}")
                
                if not potential_encrypted:
                    logger.debug(f"No potential encrypted values found in {filename}")
                    os.remove(temp_path)
                    continue
                
                # Step 2: Decrypt all encrypted values and replace them in text
                decrypted_text = text
                decryption_map = {}  # Map encrypted -> decrypted for tracking
                successful_decryptions = 0
                failed_decryptions = []
                
                logger.info(f"Attempting to decrypt {len(potential_encrypted)} encrypted values from {filename}")
                logger.info(f"Hash_meta cache contains {len(hash_meta_cache)} files")
                
                for encrypted_val in potential_encrypted:
                    if encrypted_val in decryption_map:
                        continue  # Already processed
                    
                    # Try to decrypt using hash_meta from cache
                    decrypted_value = None
                    found_in_cache = False
                    
                    for hash_meta_data in hash_meta_cache.values():
                        if encrypted_val in hash_meta_data:
                            found_in_cache = True
                            meta_info = hash_meta_data[encrypted_val]
                            try:
                                decrypted_value = masker.decrypt_hash(
                                    encrypted_val,
                                    meta_info['hash_meta'],
                                    password
                                )
                                decryption_map[encrypted_val] = decrypted_value
                                successful_decryptions += 1
                                logger.debug(f"Decrypted: {encrypted_val[:20]}... -> {decrypted_value[:20]}...")
                                break
                            except Exception as e:
                                logger.debug(f"Decryption failed for {encrypted_val[:20]}...: {e}")
                                continue
                    
                    if not found_in_cache:
                        failed_decryptions.append(encrypted_val[:30])
                    
                    # Replace encrypted value with decrypted value in text
                    if decrypted_value:
                        # Replace all occurrences of this encrypted value
                        decrypted_text = decrypted_text.replace(encrypted_val, decrypted_value)
                
                logger.info(f"Successfully decrypted {successful_decryptions}/{len(potential_encrypted)} values in {filename}")
                if failed_decryptions:
                    logger.warning(f"Failed to find hash_meta for {len(failed_decryptions)} encrypted values (sample: {failed_decryptions[:3]})")
                
                if successful_decryptions == 0:
                    logger.warning(f"No values could be decrypted from {filename} - wrong password or missing hash_meta")
                    logger.warning(f"Hash_meta cache keys: {list(hash_meta_cache.keys())[:3] if hash_meta_cache else 'Empty'}")
                    os.remove(temp_path)
                    continue
                
                # Log sample of decrypted text for debugging
                logger.debug(f"Sample decrypted text (first 500 chars): {decrypted_text[:500]}")
                
                # Step 3: Detect PIIs from the fully decrypted text
                from pii_detector_advanced import ContextAwarePIIDetector
                pii_detector = ContextAwarePIIDetector()
                
                # Detect PIIs in decrypted text using scan_text_advanced
                detected_piis = pii_detector.scan_text_advanced(decrypted_text)
                
                logger.info(f"Detected {len(detected_piis)} PIIs in decrypted text from {filename}")
                
                # Log detected PII types for debugging
                pii_types_found = {}
                for pii in detected_piis:
                    pii_type = pii.get('type', 'UNKNOWN')
                    pii_types_found[pii_type] = pii_types_found.get(pii_type, 0) + 1
                logger.info(f"PII types detected: {pii_types_found}")
                
                # Format PIIs for response
                decrypted_file_piis = []
                for pii in detected_piis:
                    # Use 'match' or 'value' field from detector
                    pii_value = pii.get('match') or pii.get('value', '')
                    if not pii_value:
                        continue  # Skip if no value
                    
                    decrypted_file_piis.append({
                        'type': pii.get('type', 'UNKNOWN'),
                        'value': pii_value,
                        'file': filename,
                        'page': 0,  # Text files don't have pages
                        'confidence': pii.get('confidence', 0.5)
                    })
                
                if decrypted_file_piis:
                    decrypted_results.append({
                        'filename': filename,
                        'piis': decrypted_file_piis,
                        'pii_count': len(decrypted_file_piis)
                    })
                    all_decrypted_piis.extend(decrypted_file_piis)
                
            except Exception as e:
                logger.error(f"Error processing {filename} for decryption: {e}", exc_info=True)
            finally:
                # Clean up temp file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        
        if not decrypted_results:
            return jsonify({'error': 'No encrypted PIIs found or incorrect password'}), 404
        
        # Group PIIs by type
        pii_types = {}
        for pii in all_decrypted_piis:
            pii_type = pii.get('type', 'UNKNOWN')
            if pii_type not in pii_types:
                pii_types[pii_type] = []
            pii_types[pii_type].append(pii)
        
            logger.info(f"Decryption successful: {len(all_decrypted_piis)} total PIIs decrypted from {len(decrypted_results)} files")
            return jsonify({
                'success': True,
                'files': decrypted_results,
                'pii_types': pii_types,
                'total_piis': len(all_decrypted_piis)
            }), 200
    
    except Exception as e:
        logger.error(f"Error in decrypt-upload: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/batch/<batch_id>/analysis', methods=['GET'])
@require_api_key
def get_batch_analysis(batch_id: str):
    """Get analysis data for a single batch."""
    try:
        analysis = mongo_client.get_batch_analysis(batch_id)
        if not analysis:
            return jsonify({'error': 'Batch not found'}), 404
        return jsonify(analysis), 200
    except Exception as e:
        logger.error(f"Error getting batch analysis: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/debug/batch/<batch_id>/json-files', methods=['GET'])
@require_api_key
def debug_batch_json_files(batch_id: str):
    """Debug endpoint to check JSON files for a batch."""
    try:
        import os
        storage_path = os.getenv('STORAGE_PATH', './data')
        results_folder = os.path.join(storage_path, 'results', batch_id)
        
        debug_info = {
            'batch_id': batch_id,
            'results_folder': results_folder,
            'folder_exists': os.path.exists(results_folder),
            'json_files': []
        }
        
        if os.path.exists(results_folder):
            json_files = [f for f in os.listdir(results_folder) if f.endswith('.json')]
            debug_info['json_file_count'] = len(json_files)
            
            for json_file in json_files:
                json_path = os.path.join(results_folder, json_file)
                try:
                    with open(json_path, 'r', encoding='utf-8') as f:
                        file_data = json.load(f)
                        piis = file_data.get('piis', [])
                        pii_count = file_data.get('pii_count', len(piis) if isinstance(piis, list) else 0)
                        
                        debug_info['json_files'].append({
                            'filename': json_file,
                            'pii_count': pii_count,
                            'piis_length': len(piis) if isinstance(piis, list) else 0,
                            'has_piis': len(piis) > 0 if isinstance(piis, list) else False,
                            'sample_piis': piis[:3] if isinstance(piis, list) and len(piis) > 0 else []
                        })
                except Exception as e:
                    debug_info['json_files'].append({
                        'filename': json_file,
                        'error': str(e)
                    })
        else:
            debug_info['error'] = f'Results folder does not exist: {results_folder}'
        
        return jsonify(debug_info), 200
    except Exception as e:
        logger.error(f"Error in debug endpoint: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/batches', methods=['GET'])
@require_api_key
def list_batches():
    """List batches - ULTRA-FAST with async processing (< 500ms)."""
    try:
        user_id = request.args.get('user_id', 'default')
        limit = int(request.args.get('limit', 100))
        
        # SPEED OPTIMIZATION 1: Check cache first (INSTANT < 10ms!)
        cached_batches = get_cached_batches(user_id)
        if cached_batches is not None:
            logger.debug(f"Cache HIT for user {user_id} - instant return!")
            return jsonify({'batches': cached_batches, 'from_cache': True}), 200
        
        # SPEED OPTIMIZATION 2: Async fetch in background thread
        # Return immediately with partial data, fetch full data in background
        start_time = time.time()
        
        # Try to get batches within 500ms timeout
        batches = []
        def fetch_batches():
            nonlocal batches
        batches = mongo_client.list_batches(user_id, limit)
        
        fetch_thread = threading.Thread(target=fetch_batches, daemon=True)
        fetch_thread.start()
        fetch_thread.join(timeout=0.5)  # Wait max 500ms
        
        if not batches:
            # If DB is slow, return empty but cached result
            logger.warning(f"DB timeout for user {user_id} - returning empty list")
            optimized_batches = []
        else:
            # ULTRA-OPTIMIZATION: Minimal payload - only essential data
            optimized_batches = []
            for batch in batches[:20]:  # Process first 20 batches only
                optimized_batch = {
                    'batch_id': batch.get('batch_id'),
                    'name': batch.get('name'),
                    'created_at': batch.get('created_at'),
                    'updated_at': batch.get('updated_at'),
                    'processed_at': batch.get('processed_at'),
                    'status': batch.get('status', 'pending'),
                    'stats': batch.get('stats', {}),
                    'summary': batch.get('summary', {})
                }
                
                # Minimal files data - NO PIIs (load on demand)
                if 'files' in batch:
                    optimized_batch['files'] = [
                        {
                            'filename': f.get('filename'),
                            'status': f.get('status', 'pending'),
                            'pii_count': f.get('pii_count', 0),
                            'total_piis': f.get('pii_count', 0)
                        }
                        for f in batch.get('files', [])[:5]  # Only first 5 files (lazy!)
                    ]
                
                optimized_batches.append(optimized_batch)
        
        # Cache the result for INSTANT subsequent requests
        set_cached_batches(user_id, optimized_batches)
        
        elapsed = time.time() - start_time
        logger.info(f"✅ Batches loaded in {elapsed*1000:.0f}ms - {len(optimized_batches)} batches")
        return jsonify({'batches': optimized_batches, 'load_time_ms': int(elapsed*1000)}), 200
    
    except Exception as e:
        logger.error(f"Error listing batches: {e}")
        return jsonify({'error': str(e), 'batches': []}), 200  # Return empty list on error


@app.route('/api/batch/<batch_id>/piis-for-file', methods=['GET'])
@require_api_key
def get_file_piis(batch_id: str):
    """Get PII data for a specific file (lazy loading on-demand)."""
    try:
        filename = request.args.get('filename')
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 20))
        
        if not filename:
            return jsonify({'error': 'filename parameter required'}), 400
        
        # Get batch analysis
        analysis = mongo_client.get_batch_analysis(batch_id)
        if not analysis:
            return jsonify({'error': 'Batch not found'}), 404
        
        # Find the file in the batch
        file_data = None
        for f in analysis.get('files', []):
            if f.get('filename') == filename:
                file_data = f
                break
        
        if not file_data:
            return jsonify({'error': 'File not found in batch'}), 404
        
        # Get PIIs with pagination
        all_piis = file_data.get('piis', [])
        total_piis = len(all_piis)
        
        # Calculate pagination
        start_idx = (page - 1) * per_page
        end_idx = start_idx + per_page
        paginated_piis = all_piis[start_idx:end_idx]
        
        return jsonify({
            'filename': filename,
            'piis': paginated_piis,
            'total': total_piis,
            'page': page,
            'per_page': per_page,
            'total_pages': (total_piis + per_page - 1) // per_page
        }), 200
    
    except Exception as e:
        logger.error(f"Error getting file PIIs: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/batch/<batch_id>', methods=['DELETE'])
@require_api_key
def delete_batch(batch_id: str):
    """Delete a batch and all its data."""
    try:
        batch = mongo_client.get_batch(batch_id)
        if not batch:
            return jsonify({'error': 'Batch not found'}), 404
        
        # Delete from MongoDB
        deleted = mongo_client.delete_batch(batch_id)
        if not deleted:
            return jsonify({'error': 'Failed to delete batch from database'}), 500
        
        # Delete files from filesystem
        import shutil
        batch_folders = [
            os.path.join(app.config['RESULTS_FOLDER'], batch_id),
            os.path.join(app.config['MASKED_FOLDER'], batch_id),
            os.path.join(app.config['UPLOAD_FOLDER'])  # Will search for files with batch_id
        ]
        
        deleted_files = []
        for folder in batch_folders:
            if os.path.exists(folder):
                try:
                    if os.path.isdir(folder):
                        shutil.rmtree(folder)
                        deleted_files.append(folder)
                    else:
                        os.remove(folder)
                        deleted_files.append(folder)
                except Exception as e:
                    logger.warning(f"Error deleting {folder}: {e}")
        
        # Also delete upload files that match this batch
        uploads_folder = app.config['UPLOAD_FOLDER']
        if os.path.exists(uploads_folder):
            for file in os.listdir(uploads_folder):
                # Files are named with job_id, so we'll delete by checking batch association
                # This is a simplified approach
                pass
        
        logger.info(f"Deleted batch {batch_id} and {len(deleted_files)} folders")
        
        # Invalidate cache for all users since batch is deleted
        # (We'll invalidate 'default' cache which is most common)
        invalidate_batch_cache('default')
        
        return jsonify({
            'message': 'Batch deleted successfully',
            'batch_id': batch_id,
            'deleted_folders': len(deleted_files)
        }), 200
    
    except Exception as e:
        logger.error(f"Error deleting batch: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/download', methods=['GET'])
@require_api_key
def download_file():
    """Download a file."""
    try:
        file_path = request.args.get('path')
        if not file_path:
            return jsonify({'error': 'path parameter required'}), 400
        
        # Security: ensure path is within allowed directories
        # Normalize paths for comparison
        base_path = os.path.abspath(os.getenv('STORAGE_PATH', './data'))
        allowed_dirs = [
            os.path.abspath(app.config['MASKED_FOLDER']),
            os.path.abspath(app.config['RESULTS_FOLDER']),
            base_path
        ]
        
        # Resolve the requested path
        if not os.path.isabs(file_path):
            # Relative path - resolve it
            file_path = os.path.join(base_path, file_path.lstrip('./'))
        
        file_path = os.path.abspath(file_path)
        
        # Check if path is within allowed directories
        is_allowed = False
        for allowed_dir in allowed_dirs:
            if file_path.startswith(os.path.abspath(allowed_dir)):
                is_allowed = True
                break
        
        if not is_allowed:
            logger.warning(f"Invalid file path attempted: {file_path}")
            return jsonify({'error': 'Invalid file path'}), 403
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404
        
        return send_file(file_path, as_attachment=True)
    
    except Exception as e:
        logger.error(f"Error downloading file: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/debug/batch/<batch_id>/test-export', methods=['POST'])
@require_api_key
def test_export_debug(batch_id: str):
    """
    DEBUG endpoint - simulates export to show exactly what's happening.
    Helps diagnose why PIIs are 0 in export.
    """
    try:
        logger.info(f"\n{'='*80}")
        logger.info(f"🔍 DEBUG TEST EXPORT for batch: {batch_id}")
        logger.info(f"{'='*80}")
        
        # Fetch batch
        batch = mongo_client.get_batch_analysis(batch_id)
        if not batch:
            logger.error(f"❌ Batch not found")
            return jsonify({'error': 'Batch not found'}), 404
        
        logger.info(f"✓ Batch found: {batch.get('name')}")
        logger.info(f"✓ Batch keys: {list(batch.keys())}")
        logger.info(f"✓ Batch stats: {batch.get('stats')}")
        
        files = batch.get('files', [])
        logger.info(f"✓ Files in batch: {len(files)}")
        
        # Simulate PII extraction
        total_extracted = 0
        for file_idx, file_data in enumerate(files):
            filename = file_data.get('filename', f'file_{file_idx}')
            file_piis = file_data.get('piis', [])
            pii_count_in_db = file_data.get('pii_count', 0)
            
            logger.info(f"\n  📄 File {file_idx}: {filename}")
            logger.info(f"     pii_count field: {pii_count_in_db}")
            logger.info(f"     piis array length: {len(file_piis) if isinstance(file_piis, list) else 'N/A'}")
            logger.info(f"     piis type: {type(file_piis).__name__}")
            logger.info(f"     piis is list: {isinstance(file_piis, list)}")
            
            if isinstance(file_piis, list):
                total_extracted += len(file_piis)
                
                # Sample first PII
                if len(file_piis) > 0:
                    sample = file_piis[0]
                    logger.info(f"     Sample PII[0] type: {type(sample).__name__}")
                    if isinstance(sample, dict):
                        logger.info(f"     Sample PII keys: {list(sample.keys())}")
                        logger.info(f"     Sample PII value: {str(sample.get('value', sample.get('match', '')))[:100]}")
                    else:
                        logger.info(f"     WARNING: Sample PII is not dict! Content: {str(sample)[:100]}")
        
        logger.info(f"\n✓ Total PIIs extracted: {total_extracted}")
        logger.info(f"{'='*80}\n")
        
        return jsonify({
            'batch_id': batch_id,
            'file_count': len(files),
            'total_piis_extracted': total_extracted,
            'debug_summary': f'Found {total_extracted} PIIs across {len(files)} files'
        }), 200
    
    except Exception as e:
        logger.error(f"❌ Debug error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/debug/batch/<batch_id>/raw', methods=['GET'])
@require_api_key
def debug_batch_raw(batch_id: str):
    """Debug endpoint to inspect raw batch data from MongoDB."""
    try:
        batch = mongo_client.get_batch_analysis(batch_id)
        if not batch:
            return jsonify({'error': 'Batch not found'}), 404
        
        # Analyze the batch structure
        debug_info = {
            'batch_id': batch.get('batch_id'),
            'batch_name': batch.get('name'),
            'file_count': len(batch.get('files', [])),
            'files': []
        }
        
        for idx, file_data in enumerate(batch.get('files', [])):
            piis = file_data.get('piis', [])
            file_info = {
                'index': idx,
                'filename': file_data.get('filename'),
                'pii_count_in_db': file_data.get('pii_count', 0),
                'piis_array_length': len(piis) if isinstance(piis, list) else 0,
                'piis_is_list': isinstance(piis, list),
                'piis_type': type(piis).__name__,
                'sample_piis': []
            }
            
            # Sample first 3 PIIs
            if isinstance(piis, list) and len(piis) > 0:
                for pii_idx, pii in enumerate(piis[:3]):
                    if isinstance(pii, dict):
                        file_info['sample_piis'].append({
                            'index': pii_idx,
                            'type': pii.get('type'),
                            'keys': list(pii.keys()),
                            'has_value': 'value' in pii,
                            'has_match': 'match' in pii,
                            'has_normalized': 'normalized' in pii,
                            'has_text': 'text' in pii,
                            'value_preview': str(pii.get('value', pii.get('match', pii.get('normalized', pii.get('text', '')))))[:100]
                        })
                    else:
                        file_info['sample_piis'].append({
                            'index': pii_idx,
                            'type': type(pii).__name__,
                            'error': 'PII is not a dict'
                        })
            
            debug_info['files'].append(file_info)
        
        return jsonify(debug_info), 200
    except Exception as e:
        logger.error(f"Error in debug endpoint: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/export-pii-json', methods=['POST'])
@require_api_key
def export_pii_json():
    """
    BULLETPROOF PII Export with SHA-512 Encryption.
    """
    start_time = time.time()
    
    try:
        # STEP 1: Parse request
        logger.info(f"\n{'='*80}")
        logger.info(f"🔄 EXPORT: Received request")
        logger.info(f"   Content-Type: {request.content_type}")
        
        try:
            data = request.get_json(force=True, silent=False)
        except Exception as json_err:
            logger.error(f"❌ JSON parse error: {json_err}")
            logger.error(f"   Raw body: {request.get_data()}")
            return jsonify({'error': f'Invalid JSON: {str(json_err)}'}), 400
        
        if not data:
            logger.error(f"❌ Empty JSON data")
            return jsonify({'error': 'Empty request body'}), 400
        
        batch_id = data.get('batch_id', '')
        selected_pii_types = data.get('selected_pii_types', []) or []
        password = data.get('password', '')
        lock_file = bool(data.get('lock_file', False))
        
        logger.info(f"✓ Request parsed:")
        logger.info(f"   batch_id: {batch_id} (type: {type(batch_id).__name__}, len: {len(batch_id) if isinstance(batch_id, str) else 'N/A'})")
        logger.info(f"   lock_file: {lock_file}")
        logger.info(f"   password: {'***' if password else 'none'}")
        logger.info(f"   selected_pii_types: {selected_pii_types} (len: {len(selected_pii_types)})")
        
        # STEP 2: Validate inputs
        if not batch_id or not isinstance(batch_id, str):
            logger.error(f"❌ Invalid batch_id: {batch_id}")
            return jsonify({'error': 'batch_id is required and must be a string'}), 400
        
        if lock_file and not password:
            logger.error(f"❌ Password required when lock_file=true")
            return jsonify({'error': 'password is required when locking file'}), 400
        
        logger.info(f"{'='*80}")
        logger.info(f"🔄 STEP 1: Fetching batch from MongoDB...")
        
        # Fetch batch
        try:
            batch = mongo_client.get_batch_analysis(batch_id)
        except Exception as fetch_err:
            logger.error(f"❌ Error fetching batch: {fetch_err}", exc_info=True)
            return jsonify({'error': f'Database error: {str(fetch_err)}'}), 500
        
        if not batch:
            logger.error(f"❌ Batch not found: {batch_id}")
            return jsonify({'error': 'Batch not found'}), 404
        
        batch_name = batch.get('name', 'unknown')
        files = batch.get('files', [])
        logger.info(f"✓ Batch: {batch_name} | Files: {len(files)}")
        
        if not files:
            logger.error(f"❌ Batch has no files")
            return jsonify({'error': 'Batch has no files'}), 400
        
        logger.info(f"🔄 STEP 2: Extracting PIIs...")
        
        # Extract PIIs - SIMPLE AND ROBUST
        all_piis_dict = {}  # {TYPE: [values]}
        file_summary = []
        total_piis = 0
        
        logger.info(f"Processing {len(files)} files from batch...")
        
        for file_idx, file_data in enumerate(files):
            filename = file_data.get('filename', 'unknown')
            piis_list = file_data.get('piis', [])
            
            logger.info(f"  File {file_idx}: {filename}")
            logger.info(f"    piis_list type: {type(piis_list).__name__}")
            logger.info(f"    piis_list length: {len(piis_list) if isinstance(piis_list, list) else 'NOT A LIST'}")
            
            if not isinstance(piis_list, list):
                logger.warning(f"    ⚠️  Skipping: piis not a list (type: {type(piis_list)})")
                continue
            
            if len(piis_list) == 0:
                logger.warning(f"    ⚠️  Empty PIIs array")
                continue
            
            file_piis = 0
            for pii_idx, pii in enumerate(piis_list):
                if not isinstance(pii, dict):
                    logger.warning(f"    PII {pii_idx}: not a dict, type={type(pii).__name__}")
                    continue
                
                pii_type = str(pii.get('type', 'UNKNOWN')).upper()
                pii_value = pii.get('value') or pii.get('match') or pii.get('normalized') or pii.get('text') or ''
                
                if isinstance(pii_value, str):
                    pii_value = pii_value.strip()
                else:
                    pii_value = str(pii_value).strip() if pii_value else ''
                
                if not pii_value:
                    logger.debug(f"    PII {pii_idx}: empty value, keys={list(pii.keys())}")
                    continue
                
                # Check type filter
                if selected_pii_types:
                    if pii_type not in [t.upper() for t in selected_pii_types]:
                        continue
                
                # Add to dict
                if pii_type not in all_piis_dict:
                    all_piis_dict[pii_type] = []
                
                if pii_value not in all_piis_dict[pii_type]:
                    all_piis_dict[pii_type].append(pii_value)
                    file_piis += 1
                    total_piis += 1
                
                # Log first few
                if pii_idx < 3:
                    logger.debug(f"    PII {pii_idx}: type={pii_type}, value={pii_value[:30]}")
            
            logger.info(f"    ✓ Extracted {file_piis} from {filename}")
            
            if file_piis > 0:
                file_summary.append({'filename': filename, 'pii_count': file_piis})
        
        logger.info(f"Total extracted: {total_piis} PIIs across {len(all_piis_dict)} types")
        
        if total_piis == 0:
            logger.error(f"❌ NO PIIs extracted!")
            logger.error(f"   Files in batch: {len(files)}")
            if files:
                logger.error(f"   First file keys: {list(files[0].keys())}")
                first_piis = files[0].get('piis', [])
                logger.error(f"   First file 'piis' field: type={type(first_piis).__name__}, len={len(first_piis) if isinstance(first_piis, list) else '?'}")
                if isinstance(first_piis, list) and len(first_piis) > 0:
                    logger.error(f"   Sample PII structure: {list(first_piis[0].keys()) if isinstance(first_piis[0], dict) else 'not dict'}")
                    logger.error(f"   Sample PII full: {first_piis[0]}")
            return jsonify({'error': 'No PIIs found in batch'}), 400
        
        logger.info(f"✓ Extracted {total_piis} PIIs from {len(all_piis_dict)} types")
        
        logger.info(f"🔄 STEP 3: Building JSON...")
        
        json_data = {
            'metadata': {
                'batch_id': batch_id,
                'batch_name': batch_name,
                'exported_at': datetime.utcnow().isoformat(),
                'total_piis': total_piis,
                'pii_types': sorted(all_piis_dict.keys()),
                'files': file_summary,
                'encryption': 'AES-GCM-256 + PBKDF2-SHA512' if lock_file else 'None'
            },
            'piis': all_piis_dict
        }
        
        json_string = json.dumps(json_data, indent=2, ensure_ascii=False)
        file_id = str(uuid.uuid4())
        
        logger.info(f"✓ JSON: {len(json_string)} bytes")
        
        # Encrypt if needed
        if lock_file and password:
            logger.info(f"🔄 STEP 4: Encrypting...")
            
            salt = os.urandom(16)
            iv = os.urandom(12)
            # Use SHA-512 for key derivation (more secure)
            key = masker.derive_key(password, salt, use_sha512=True)
            aesgcm = AESGCM(key)
            ciphertext = aesgcm.encrypt(iv, json_string.encode('utf-8'), None)
            
            encrypted_data = {
                'file_id': file_id,
                'encrypted': True,
                'salt': base64.b64encode(salt).decode('utf-8'),
                'iv': base64.b64encode(iv).decode('utf-8'),
                'ciphertext': base64.b64encode(ciphertext).decode('utf-8'),
                'algorithm': 'AES-GCM-256',
                'kdf': 'PBKDF2-SHA512',
                'iterations': 150000
            }
            
            password_hash = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            mongo_client.store_encrypted_file_password(
                file_id=file_id,
                password_hash=password_hash,
                batch_id=batch_id,
                metadata=json_data['metadata']
            )
            
            elapsed = time.time() - start_time
            logger.info(f"✅ EXPORT COMPLETE: {total_piis} PIIs encrypted in {elapsed:.2f}s")
            logger.info(f"{'='*80}\n")
            
            return jsonify({
                'success': True,
                'file_id': file_id,
                'encrypted': True,
                'data': encrypted_data
            }), 200
        else:
            elapsed = time.time() - start_time
            logger.info(f"✅ EXPORT COMPLETE: {total_piis} PIIs unencrypted in {elapsed:.2f}s")
            logger.info(f"{'='*80}\n")
            
            return jsonify({
                'success': True,
                'file_id': file_id,
                'encrypted': False,
                'data': json_data
            }), 200
    
    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"❌ EXPORT FAILED ({elapsed:.2f}s): {str(e)}", exc_info=True)
        logger.error(f"{'='*80}\n")
        return jsonify({'error': str(e)}), 500


@app.route('/api/decrypt-json', methods=['POST'])
@require_api_key
def decrypt_json():
    """
    Decrypt uploaded JSON file and verify password.
    REWRITTEN to handle both SHA-512 and SHA-256 (backward compatible).
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        password = request.form.get('password')
        if not password:
            return jsonify({'error': 'Password is required'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        logger.info(f"🔓 Decryption request for file: {file.filename}")
        
        # ============= STEP 1: Read and parse encrypted file =============
        file_content = file.read()
        
        try:
            encrypted_data = json.loads(file_content.decode('utf-8'))
        except json.JSONDecodeError as e:
            logger.error(f"❌ Invalid JSON file: {e}")
            return jsonify({'error': 'Invalid JSON file'}), 400
        
        if not encrypted_data.get('encrypted'):
            return jsonify({'error': 'File is not encrypted'}), 400
        
        file_id = encrypted_data.get('file_id')
        if not file_id:
            return jsonify({'error': 'File ID missing'}), 400
        
        logger.info(f"📂 File ID: {file_id}")
        
        # ============= STEP 2: Verify password from database =============
        try:
            if mongo_client.client is None or mongo_client.db is None:
                logger.error("❌ Database connection failed")
                return jsonify({'error': 'Database connection failed'}), 500
            
            collection = mongo_client.db.file_public_keys
            stored_doc = collection.find_one({"file_id": file_id})
        except (AttributeError, TypeError) as e:
            logger.error(f"❌ Database connection error: {e}")
            return jsonify({'error': 'Database connection failed'}), 500
        except Exception as db_error:
            logger.error(f"❌ Database error: {db_error}")
            return jsonify({'error': 'Database connection failed'}), 500
        
        if not stored_doc:
            logger.error(f"❌ File {file_id} not found in database")
            return jsonify({'error': 'File not found in database'}), 404
        
        # Verify password against stored hash
        stored_hash = stored_doc.get('password_hash')
        if not stored_hash:
            logger.error(f"❌ Password hash not found for file {file_id}")
            return jsonify({'error': 'Password hash not found'}), 404
        
        # Check password using bcrypt
        if not bcrypt.checkpw(password.encode('utf-8'), stored_hash.encode('utf-8')):
            logger.warning(f"❌ Incorrect password for file {file_id}")
            return jsonify({'error': 'Incorrect password'}), 401
        
        logger.info(f"✓ Password verified")
        
        # ============= STEP 3: Decrypt JSON =============
        # Handle both old format (with nested 'data') and new format (flat structure)
        # Old format: { "success": true, "file_id": "...", "encrypted": true, "data": { "salt": "...", "iv": "...", "ciphertext": "..." } }
        # New format: { "file_id": "...", "encrypted": true, "salt": "...", "iv": "...", "ciphertext": "..." }
        
        # Check if this is the old nested format
        if 'data' in encrypted_data and isinstance(encrypted_data['data'], dict):
            logger.info("📦 Detected old format with nested 'data' field")
            encrypted_data = encrypted_data['data']  # Extract the nested data
        
        # Check if required fields exist
        required_fields = ['salt', 'iv', 'ciphertext']
        missing_fields = [field for field in required_fields if field not in encrypted_data]
        
        if missing_fields:
            logger.error(f"❌ Missing required encryption fields: {missing_fields}")
            logger.error(f"   File structure: {list(encrypted_data.keys())}")
            return jsonify({
                'error': f'Invalid encrypted file format. Missing fields: {", ".join(missing_fields)}. This file may not be encrypted or is corrupted.'
            }), 400
        
        try:
            salt = base64.b64decode(encrypted_data['salt'])
            iv = base64.b64decode(encrypted_data['iv'])
            ciphertext = base64.b64decode(encrypted_data['ciphertext'])
        except Exception as decode_error:
            logger.error(f"❌ Base64 decode error: {decode_error}")
            return jsonify({'error': 'Invalid encrypted data format - base64 decode failed'}), 400
        
        # Detect KDF algorithm (SHA-512 or SHA-256)
        kdf_algorithm = encrypted_data.get('kdf', 'PBKDF2-SHA256')  # Default to SHA-256 for old files
        use_sha512 = 'SHA512' in kdf_algorithm.upper() or 'SHA-512' in kdf_algorithm.upper()
        
        logger.info(f"🔓 Decrypting with {kdf_algorithm}...")
        
        # Derive key using the detected algorithm
        key = masker.derive_key(password, salt, use_sha512=use_sha512)
        aesgcm = AESGCM(key)
        
        try:
            plaintext = aesgcm.decrypt(iv, ciphertext, None)
            decrypted_json = json.loads(plaintext.decode('utf-8'))
        except Exception as decrypt_error:
            logger.error(f"❌ Decryption failed: {decrypt_error}")
            return jsonify({'error': 'Decryption failed - incorrect password or corrupted file'}), 401
        
        # ============= STEP 4: Validate decrypted data =============
        total_piis = decrypted_json.get('metadata', {}).get('total_piis', 0)
        pii_types = decrypted_json.get('metadata', {}).get('pii_types', [])
        
        logger.info(f"✅ Decryption successful: {total_piis} PIIs across {len(pii_types)} types")
        logger.info(f"   PII types: {pii_types}")
        
        return jsonify({
            'success': True,
            'data': decrypted_json
        }), 200
    
    except Exception as e:
        logger.error(f"❌ Error decrypting JSON: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.errorhandler(413)
def request_entity_too_large(error):
    """Handle file too large error."""
    return jsonify({'error': 'File too large. Maximum size is 500MB'}), 413


@app.errorhandler(415)
def unsupported_media_type(error):
    """Handle unsupported media type error."""
    logger.error(f"❌ 415 Unsupported Media Type!")
    logger.error(f"   Content-Type: {request.content_type}")
    logger.error(f"   Path: {request.path}")
    logger.error(f"   Method: {request.method}")
    logger.error(f"   Headers: {dict(request.headers)}")
    logger.error(f"   Error: {error}")
    
    # For upload endpoint, just accept it anyway
    if request.path == '/api/upload':
        logger.warning(f"   Allowing upload to proceed despite 415 error")
        return jsonify({'error': 'Media type error but proceeding'}), 415
    
    return jsonify({'error': f'Unsupported Media Type. Got: {request.content_type}'}), 415


@app.route('/api/login', methods=['POST'])
@require_api_key
@rate_limit(max_requests=5, window=300)  # 5 attempts per 5 minutes
def login():
    """Authenticate user with email and password."""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        email = data.get('email')
        password = data.get('password')
        
        if not email or not password:
            return jsonify({'error': 'Email and password are required'}), 400
        
        # Get user from database
        user = mongo_client.get_user_by_email(email)
        
        if not user:
            logger.warning(f"Login attempt failed: User not found - {email}")
            return jsonify({'error': 'Invalid credentials'}), 401
        
        # Verify password
        stored_hash = user.get('password_hash')
        if not stored_hash:
            logger.warning(f"Login attempt failed: No password hash found for user - {email}")
            return jsonify({'error': 'Invalid credentials'}), 401
        
        # Check password
        if not bcrypt.checkpw(password.encode('utf-8'), stored_hash.encode('utf-8')):
            logger.warning(f"Login attempt failed: Incorrect password for user - {email}")
            return jsonify({'error': 'Invalid credentials'}), 401
        
        # Login successful
        logger.info(f"User logged in successfully: {email}")
        
        # Return success response with user info (exclude password hash)
        user_info = {k: v for k, v in user.items() if k != 'password_hash'}
        
        response_data = {
            'success': True,
            'message': 'Login successful',
            'user': user_info,
            'token': str(uuid.uuid4())  # Generate a simple token (can be replaced with JWT later)
        }
        
        return jsonify(response_data), 200
        
    except Exception as e:
        logger.error(f"Error during login: {e}", exc_info=True)
        return jsonify({'error': f'Login failed: {str(e)}'}), 500


def generate_otp():
    """Generate a secure 6-digit OTP (no leading zeros removed)."""
    return str(random.randint(100000, 999999))


def send_sms_otp(mobile: str, otp: str, country_code: Optional[str] = None) -> bool:
    """
    Send OTP via SMS using 2Factor.in or Twilio (fallback).
    Falls back to logging if SMS providers are not configured.
    """
    # Normalize inputs
    digits_only_mobile = ''.join(filter(str.isdigit, str(mobile)))
    digits_only_country = ''.join(filter(str.isdigit, str(country_code))) if country_code else ''
    if not digits_only_mobile:
        logger.error("send_sms_otp: mobile number is empty after normalization")
        return False
    if not digits_only_country:
        digits_only_country = '91'
    phone_numeric = f"{digits_only_country}{digits_only_mobile}"
    phone_e164 = f"+{phone_numeric}"

    # Primary provider: 2Factor.in (Indian SMS gateway)
    two_factor_api_key = os.getenv('TWO_FACTOR_API_KEY')
    if two_factor_api_key:
        try:
            url = f"https://2factor.in/API/V1/{two_factor_api_key}/SMS/{phone_numeric}/{otp}"
            logger.info(f"Sending OTP via 2Factor.in to {phone_e164}")
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            data = response.json()
            status = data.get('Status')
            details = data.get('Details')
            if status == 'Success':
                logger.info(f"OTP sent successfully via 2Factor.in: Details={details}")
                return True
            logger.error(f"2Factor.in responded with error: Status={status}, Details={details}")
        except Exception as e:
            logger.error(f"Error sending OTP via 2Factor.in: {e}", exc_info=True)

    # Fallback provider: Twilio
    twilio_account_sid = os.getenv('TWILIO_ACCOUNT_SID')
    twilio_auth_token = os.getenv('TWILIO_AUTH_TOKEN')
    twilio_phone_number = os.getenv('TWILIO_PHONE_NUMBER')
    
    # If Twilio credentials are not set, log the OTP (for development/testing)
    if not twilio_account_sid or not twilio_auth_token or not twilio_phone_number:
        logger.warning("Twilio credentials not configured. OTP will be logged instead of sent via SMS.")
        logger.info(f"OTP for {phone_e164}: {otp}")
        logger.info("To enable SMS, set TWO_FACTOR_API_KEY or TWILIO_* credentials in the environment")
        return True
    
    try:
        from twilio.rest import Client
        
        # Initialize Twilio client
        client = Client(twilio_account_sid, twilio_auth_token)
        
        # Create SMS message
        message_body = f"Your PII Sentinel OTP is {otp}. Valid for 2 minutes. Do not share this code."
        
        # Send SMS
        message = client.messages.create(
            body=message_body,
            from_=twilio_phone_number,
            to=phone_e164
        )
        
        logger.info(f"SMS sent successfully to {phone_e164}. Message SID: {message.sid}")
        return True
        
    except ImportError:
        logger.error("Twilio library not installed. Install with: pip install twilio")
        logger.info(f"OTP for {phone_e164}: {otp}")
        return False
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Error sending SMS via Twilio: {error_msg}", exc_info=True)
        
        # Log specific error details for debugging
        if "Invalid" in error_msg or "not found" in error_msg.lower():
            logger.error("Twilio credentials may be invalid. Please check your Account SID and Auth Token.")
        elif "not verified" in error_msg.lower() or "unverified" in error_msg.lower():
            logger.error("Phone number not verified in Twilio. For trial accounts, you can only send SMS to verified numbers.")
            logger.info("Add your number at: https://console.twilio.com/us1/develop/phone-numbers/manage/verified")
        elif "insufficient" in error_msg.lower() or "balance" in error_msg.lower():
            logger.error("Insufficient Twilio account balance. Please add funds to your Twilio account.")
        
        logger.info(f"OTP for {phone_e164}: {otp} (SMS failed, check logs above for details)")
        return False


@app.route('/api/auth/send-otp', methods=['POST'])
@require_api_key
def send_otp():
    """Send OTP to mobile number."""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        mobile_raw = data.get('mobile', '')
        country_code_raw = data.get('country_code') or data.get('countryCode') or '91'
        
        # Normalize: Extract only digits
        mobile = ''.join(filter(str.isdigit, str(mobile_raw)))
        country_code = ''.join(filter(str.isdigit, str(country_code_raw)))
        
        if not mobile:
            return jsonify({'error': 'Mobile number is required'}), 400
        
        # Validate mobile number (10 digits)
        if len(mobile) != 10:
            return jsonify({'error': 'Enter a valid 10-digit mobile number.'}), 400
        
        if not country_code:
            country_code = '91'
        
        logger.info(f"[SEND OTP] Mobile normalized: '{mobile}'")
        
        # Generate OTP
        otp = generate_otp()
        expires_at = time.time() + OTP_EXPIRY_SECONDS
        
        # Store OTP in MongoDB
        otp_stored = mongo_client.store_otp(mobile, otp, expires_at)
        
        if not otp_stored:
            logger.warning(f"Failed to store OTP in MongoDB for mobile {mobile}, using in-memory fallback")
            # Fallback to in-memory storage
            with _otp_lock:
                otp_storage[mobile] = {
                    'otp': str(otp),
                    'expires_at': expires_at,
                    'created_at': time.time()
                }
        
        # Log stored OTP for debugging
        logger.info(f"OTP generated and stored for mobile {mobile}: {otp}")
        logger.info(f"OTP will expire at: {expires_at} (in {OTP_EXPIRY_SECONDS} seconds)")
        
        # Send OTP via SMS
        sms_sent = send_sms_otp(mobile, otp, country_code)
        if not sms_sent:
            logger.error(f"Failed to send OTP via SMS to {country_code}+{mobile}")
            return jsonify({
                'error': 'Failed to send OTP via SMS. Please try again in a moment.'
            }), 502
        
        logger.info(f"OTP sent to mobile: {mobile} (country code {country_code})")
        
        return jsonify({
            'success': True,
            'message': 'OTP sent successfully'
        }), 200
        
    except Exception as e:
        logger.error(f"Error sending OTP: {e}", exc_info=True)
        return jsonify({'error': f'Failed to send OTP: {str(e)}'}), 500


@app.route('/api/auth/resend-otp', methods=['POST'])
@require_api_key
@rate_limit(max_requests=3, window=60)  # 3 requests per minute
def resend_otp():
    """Resend OTP to mobile number."""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        mobile_raw = data.get('mobile', '')
        country_code_raw = data.get('country_code') or data.get('countryCode') or '91'
        
        # Normalize: Extract only digits
        mobile = ''.join(filter(str.isdigit, str(mobile_raw)))
        country_code = ''.join(filter(str.isdigit, str(country_code_raw)))
        
        if not mobile:
            return jsonify({'error': 'Mobile number is required'}), 400
        
        # Validate mobile number (10 digits)
        if len(mobile) != 10:
            return jsonify({'error': 'Enter a valid 10-digit mobile number.'}), 400
        
        if not country_code:
            country_code = '91'
        
        logger.info(f"[RESEND OTP] Mobile normalized: '{mobile}'")
        
        # Generate new OTP
        otp = generate_otp()
        expires_at = time.time() + OTP_EXPIRY_SECONDS
        
        # Update OTP in MongoDB (overwrites previous)
        otp_stored = mongo_client.store_otp(mobile, otp, expires_at)
        
        if not otp_stored:
            logger.warning(f"Failed to update OTP in MongoDB for mobile {mobile}, using in-memory fallback")
            # Fallback to in-memory storage
            with _otp_lock:
                otp_storage[mobile] = {
                    'otp': str(otp),
                    'expires_at': expires_at,
                    'created_at': time.time()
                }
        
        # Log stored OTP for debugging
        logger.info(f"OTP regenerated and stored for mobile {mobile}: {otp}")
        
        # Send OTP via SMS
        sms_sent = send_sms_otp(mobile, otp, country_code)
        if not sms_sent:
            logger.error(f"[RESEND OTP] Failed to send OTP via SMS to {country_code}+{mobile}")
            return jsonify({
                'error': 'Failed to resend OTP via SMS. Please try again later.'
            }), 502
        
        logger.info(f"OTP resent to mobile: {mobile} (country code {country_code})")
        
        return jsonify({
            'success': True,
            'message': 'OTP resent successfully'
        }), 200
        
    except Exception as e:
        logger.error(f"Error resending OTP: {e}", exc_info=True)
        return jsonify({'error': f'Failed to resend OTP: {str(e)}'}), 500


@app.route('/api/auth/verify-otp', methods=['POST'])
@require_api_key
@rate_limit(max_requests=10, window=300)  # 10 attempts per 5 minutes
def verify_otp():
    """Verify OTP for mobile number."""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        # Extract and normalize inputs - only digits
        mobile_raw = data.get('mobile', '')
        otp_raw = data.get('otp', '')
        
        mobile_normalized = ''.join(filter(str.isdigit, str(mobile_raw)))
        otp_normalized = ''.join(filter(str.isdigit, str(otp_raw)))
        
        logger.info(f"[VERIFY] Request received - Mobile: '{mobile_normalized}', OTP: '{otp_normalized}'")
        
        # Validate
        if not mobile_normalized or len(mobile_normalized) != 10:
            return jsonify({'error': 'Enter a valid 10-digit mobile number.'}), 400
        
        if not otp_normalized or len(otp_normalized) != 6:
            return jsonify({'error': 'Enter a valid 6-digit OTP.'}), 400
        
        # Get OTP from MongoDB (already normalized in get_otp)
        stored_data = mongo_client.get_otp(mobile_normalized)
        
        # Fallback to in-memory if MongoDB fails
        if not stored_data:
            logger.info("[VERIFY] MongoDB lookup failed, checking in-memory")
            with _otp_lock:
                in_memory_data = otp_storage.get(mobile_normalized)
                if in_memory_data:
                    # Normalize in-memory data too
                    stored_mobile = ''.join(filter(str.isdigit, str(mobile_normalized)))
                    stored_otp = ''.join(filter(str.isdigit, str(in_memory_data.get('otp', ''))))
                    stored_data = {
                        'mobile': stored_mobile,
                        'otp': stored_otp,
                        'expires_at': in_memory_data.get('expires_at', 0)
                    }
        
        if not stored_data:
            logger.warning(f"[VERIFY] No OTP found for mobile: '{mobile_normalized}'")
            return jsonify({
                'success': False,
                'error': 'Enter valid OTP sent to your mobile number (+91 ' + mobile_normalized + ').'
            }), 400
        
        # Extract stored OTP (already normalized by get_otp)
        stored_otp = ''.join(filter(str.isdigit, str(stored_data.get('otp', ''))))
        received_otp = otp_normalized
        
        logger.info(f"[VERIFY] Comparing - Stored: '{stored_otp}' (len: {len(stored_otp)}), Received: '{received_otp}' (len: {len(received_otp)})")
        
        # Check expiry
        expires_at = float(stored_data.get('expires_at', 0))
        if time.time() > expires_at:
            logger.warning(f"[VERIFY] OTP expired for mobile: '{mobile_normalized}'")
            mongo_client.delete_otp(mobile_normalized)
            with _otp_lock:
                otp_storage.pop(mobile_normalized, None)
            return jsonify({
                'success': False,
                'error': 'OTP has expired. Please request a new OTP.'
            }), 400
        
        # Compare OTPs (both are normalized 6-digit strings)
        if stored_otp != received_otp:
            logger.warning(f"[VERIFY] OTP mismatch - Expected: '{stored_otp}', Got: '{received_otp}'")
            return jsonify({
                'success': False,
                'error': 'Enter valid OTP sent to your mobile number (+91 ' + mobile_normalized + ').'
            }), 400
        
        # OTP verified successfully!
        logger.info(f"[VERIFY] SUCCESS - Mobile: '{mobile_normalized}', OTP verified")
        
        # Find user by mobile number (phoneNumber field in database)
        # DON'T delete OTP yet - only delete after successful user lookup
        user = None
        if mongo_client.client is not None and mongo_client.db is not None:
            try:
                collection = mongo_client.db["User-Base"]
                
                # Try multiple phone number formats
                search_queries = [
                    {"phoneNumber": mobile_normalized},
                    {"phoneNumber": f"+91{mobile_normalized}"},
                    {"phoneNumber": f"91{mobile_normalized}"},
                    {"phoneNumber": f"0{mobile_normalized}"}  # Some might have leading 0
                ]
                
                # Also try searching by extracting digits from stored phoneNumber
                all_users = list(collection.find({}))
                logger.info(f"[VERIFY] Searching through {len(all_users)} users in database")
                
                # Debug: Log all phone numbers in database
                if all_users:
                    phone_numbers_in_db = []
                    for u in all_users:
                        phone = u.get('phoneNumber', 'N/A')
                        phone_normalized = ''.join(filter(str.isdigit, str(phone))) if phone != 'N/A' else 'N/A'
                        phone_numbers_in_db.append(f"'{phone}' (normalized: '{phone_normalized}')")
                    logger.info(f"[VERIFY] Phone numbers in DB: {', '.join(phone_numbers_in_db[:10])}")  # Show first 10
                
                for search_query in search_queries:
                    user = collection.find_one(search_query)
                    if user:
                        logger.info(f"[VERIFY] Found user with query: {search_query}")
                        break
                    else:
                        logger.debug(f"[VERIFY] No match for query: {search_query}")
                
                # If still not found, try manual matching by extracting digits
                if not user:
                    logger.info(f"[VERIFY] Trying manual digit extraction matching...")
                    for db_user in all_users:
                        db_phone = db_user.get('phoneNumber', '')
                        if db_phone:
                            # Extract only digits from stored phone number (handles spaces, dashes, etc.)
                            db_phone_normalized = ''.join(filter(str.isdigit, str(db_phone)))
                            logger.info(f"[VERIFY] Comparing DB phone '{db_phone}' (normalized: '{db_phone_normalized}') with '{mobile_normalized}'")
                            
                            # Always extract last 10 digits (handles any format: +91 8892211564, 918892211564, etc.)
                            if len(db_phone_normalized) >= 10:
                                db_phone_normalized = db_phone_normalized[-10:]  # Take last 10 digits
                            elif len(db_phone_normalized) < 10:
                                logger.warning(f"[VERIFY] DB phone '{db_phone}' has less than 10 digits after normalization: '{db_phone_normalized}'")
                                continue
                            
                            logger.info(f"[VERIFY] After normalization: '{db_phone_normalized}' vs '{mobile_normalized}'")
                            
                            if db_phone_normalized == mobile_normalized:
                                user = db_user
                                logger.info(f"[VERIFY] ✓ Found user via manual match - DB phone: '{db_phone}' -> normalized: '{db_phone_normalized}'")
                                break
                            else:
                                logger.debug(f"[VERIFY] No match: '{db_phone_normalized}' != '{mobile_normalized}'")
                
                if user and '_id' in user:
                    user['_id'] = str(user['_id'])
                    
            except Exception as e:
                logger.error(f"[VERIFY] Error finding user by mobile: {e}", exc_info=True)
        
        if not user:
            logger.warning(f"[VERIFY] OTP verified but user not found for mobile: '{mobile_normalized}'")
            logger.warning(f"[VERIFY] OTP will remain valid for retry. User should create account first.")
            # Don't delete OTP - let user retry or create account
            # Return specific response indicating OTP verified but user doesn't exist
            return jsonify({
                'success': False,
                'otp_verified': True,
                'user_not_found': True,
                'mobile': mobile_normalized,
                'message': f'User doesn\'t exist with mobile number +91 {mobile_normalized}. Would you like to create a new account?'
            }), 404
        
        # User found - NOW delete OTP
        logger.info(f"[VERIFY] User found for mobile: '{mobile_normalized}' - deleting OTP")
        mongo_client.delete_otp(mobile_normalized)
        with _otp_lock:
            otp_storage.pop(mobile_normalized, None)
        
        # Return success response with user info (exclude password hash)
        user_info = {k: v for k, v in user.items() if k != 'password_hash'}
        
        response_data = {
            'success': True,
            'message': 'OTP verified successfully',
            'user': user_info,
            'token': str(uuid.uuid4())  # Generate a simple token (can be replaced with JWT later)
        }
        
        return jsonify(response_data), 200
        
    except Exception as e:
        logger.error(f"Error verifying OTP: {e}", exc_info=True)
        return jsonify({'error': f'Failed to verify OTP: {str(e)}'}), 500


@app.route('/api/auth/debug-otp', methods=['GET'])
@require_api_key
def debug_otp():
    """Debug endpoint to check stored OTPs (for troubleshooting)."""
    try:
        mobile = request.args.get('mobile', '').strip()
        
        if mobile:
            # Check MongoDB first
            stored_data = mongo_client.get_otp(mobile)
            
            # Fallback to in-memory
            if not stored_data:
                with _otp_lock:
                    stored_data = otp_storage.get(mobile)
            
            if stored_data:
                time_remaining = stored_data.get('expires_at', 0) - time.time()
                return jsonify({
                    'success': True,
                    'mobile': mobile,
                    'otp': stored_data.get('otp'),
                    'time_remaining_seconds': round(time_remaining, 2),
                    'expired': time_remaining <= 0,
                    'created_at': stored_data.get('created_at'),
                    'expires_at': stored_data.get('expires_at'),
                    'source': 'mongodb' if 'verified' in stored_data else 'memory'
                }), 200
            else:
                return jsonify({
                    'success': False,
                    'message': f'No OTP found for mobile {mobile}'
                }), 404
        else:
            # Return all stored OTPs from MongoDB
            all_otps = {}
            if mongo_client.client is not None and mongo_client.db is not None:
                try:
                    collection = mongo_client.db["OTP-Storage"]
                    all_mongo_otps = list(collection.find({}))
                    for doc in all_mongo_otps:
                        mob = doc.get('mobile')
                        time_remaining = doc.get('expires_at', 0) - time.time()
                        all_otps[mob] = {
                            'otp': doc.get('otp'),
                            'time_remaining_seconds': round(time_remaining, 2),
                            'expired': time_remaining <= 0,
                            'source': 'mongodb'
                        }
                except Exception as e:
                    logger.error(f"Error fetching OTPs from MongoDB: {e}")
            
            # Also include in-memory OTPs
            with _otp_lock:
                for mob, data in otp_storage.items():
                    if mob not in all_otps:
                        time_remaining = data.get('expires_at', 0) - time.time()
                        all_otps[mob] = {
                            'otp': data.get('otp'),
                            'time_remaining_seconds': round(time_remaining, 2),
                            'expired': time_remaining <= 0,
                            'source': 'memory'
                        }
            
            return jsonify({
                'success': True,
                'stored_otps': all_otps,
                'total_count': len(all_otps)
            }), 200
    except Exception as e:
        logger.error(f"Error in debug OTP endpoint: {e}", exc_info=True)
        return jsonify({'error': f'Debug failed: {str(e)}'}), 500


@app.route('/api/auth/test-sms', methods=['POST'])
@require_api_key
def test_sms():
    """Test endpoint to verify Twilio SMS configuration."""
    try:
        twilio_account_sid = os.getenv('TWILIO_ACCOUNT_SID')
        twilio_auth_token = os.getenv('TWILIO_AUTH_TOKEN')
        twilio_phone_number = os.getenv('TWILIO_PHONE_NUMBER')
        
        config_status = {
            'twilio_account_sid': 'SET' if twilio_account_sid and twilio_account_sid != 'your_twilio_account_sid_here' else 'NOT SET',
            'twilio_auth_token': 'SET' if twilio_auth_token and twilio_auth_token != 'your_twilio_auth_token_here' else 'NOT SET',
            'twilio_phone_number': twilio_phone_number if twilio_phone_number else 'NOT SET',
            'twilio_library_installed': False
        }
        
        try:
            from twilio.rest import Client
            config_status['twilio_library_installed'] = True
            
            if config_status['twilio_account_sid'] == 'SET' and config_status['twilio_auth_token'] == 'SET':
                # Test Twilio client initialization
                try:
                    client = Client(twilio_account_sid, twilio_auth_token)
                    # Try to fetch account info to verify credentials
                    account = client.api.accounts(twilio_account_sid).fetch()
                    config_status['twilio_connection'] = 'SUCCESS'
                    config_status['account_status'] = account.status
                except Exception as e:
                    config_status['twilio_connection'] = f'FAILED: {str(e)}'
        except ImportError:
            config_status['twilio_library_installed'] = False
        
        return jsonify({
            'success': True,
            'config': config_status
        }), 200
        
    except Exception as e:
        logger.error(f"Error testing SMS config: {e}", exc_info=True)
        return jsonify({'error': f'Test failed: {str(e)}'}), 500


@app.route('/api/profile', methods=['GET'])
@require_api_key
@cached(profile_cache)
def get_profile():
    """Get user profile and stats."""
    username = request.args.get('username')
    email = request.args.get('email')
    
    if not username and not email:
        return jsonify({'error': 'Username or email is required'}), 400
    
    # Get user data
    if username:
        user = mongo_client.get_user_by_username(username)
    else:
        user = mongo_client.get_user_by_email(email)
    
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    # Ensure account_status is set for this user
    user_email = user.get('email')
    if user_email:
        mongo_client.ensure_user_token_document(user_email)
        # Re-fetch user to get updated fields
        user = mongo_client.get_user_by_email(user_email)
    
    # Remove password hash from response
    user_data = {k: v for k, v in user.items() if k != 'password_hash'}
    username_for_batches = user.get('username')
    
    # Get user's stats using optimized aggregation (fast!)
    user_id_for_stats = user.get('email') or user.get('username') or username_for_batches
    stats_data = mongo_client.get_user_stats(user_id_for_stats) if user_id_for_stats else {}
    
    total_batches = stats_data.get('total_batches', 0)
    total_files = stats_data.get('total_files', 0)
    total_piis = stats_data.get('total_piis', 0)
    
    # For detailed breakdown, we still need to load batches (but this is optional)
    gov_piis = 0
    custom_piis = 0
    pii_type_counts = {}
    gov_pii_types = {
        'aadhaar', 'pan', 'passport', 'voter_id', 'driving_license', 
        'bank_account', 'ifsc', 'upi', 'epf', 'gstin', 'cin', 'ration_card'
    }
    
    # Only load batches if we need detailed breakdown (can be made optional)
    if username_for_batches and total_batches > 0:
        try:
            batches = mongo_client.get_batches_by_username(username_for_batches)
            for batch in batches:
                breakdown = batch.get('stats', {}).get('breakdown', {})
                if breakdown:
                    for pii_type, count in breakdown.items():
                        pii_type_counts[pii_type] = pii_type_counts.get(pii_type, 0) + count
                        if pii_type.lower() in gov_pii_types:
                            gov_piis += count
                        else:
                            custom_piis += count
        except Exception as e:
            logger.warning(f"Error loading detailed breakdown: {e}")
    
    logger.info(f"Profile stats - Batches: {total_batches}, Files: {total_files}, PIIs: {total_piis}")
    
    # Calculate percentages
    gov_percentage = (gov_piis / total_piis * 100) if total_piis > 0 else 0
    custom_percentage = (custom_piis / total_piis * 100) if total_piis > 0 else 0
    
    # Prepare PII type data for scatter plot (most/least detected)
    pii_type_list = [
        {'type': pii_type, 'count': count} 
        for pii_type, count in sorted(pii_type_counts.items(), key=lambda x: x[1], reverse=True)
    ]
    
    logger.info(f"Profile stats - Gov%: {gov_percentage}, Custom%: {custom_percentage}, PII types: {len(pii_type_list)}")
    
    return jsonify({
        'success': True,
        'user': user_data,
        'stats': {
            'total_batches': total_batches,
            'total_files': total_files,
            'total_piis': total_piis,
            'gov_piis': gov_piis,
            'custom_piis': custom_piis,
            'gov_percentage': round(gov_percentage, 2),
            'custom_percentage': round(custom_percentage, 2),
            'pii_type_counts': pii_type_counts,
            'pii_type_list': pii_type_list
        }
    }), 200


@app.route('/api/profile', methods=['PUT'])
@require_api_key
def update_profile():
    """Update user profile data."""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        username = data.get('username')
        email = data.get('email')
        
        if not username and not email:
            return jsonify({'error': 'Username or email is required'}), 400
        
        # Get user data
        if username:
            user = mongo_client.get_user_by_username(username)
        else:
            user = mongo_client.get_user_by_email(email)
        
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        # Prepare update data (exclude password_hash, username, email, _id, created_at)
        update_data = {}
        exclude_fields = ['password_hash', 'username', 'email', '_id', 'created_at', 'account_status']
        
        for key, value in data.items():
            if key not in exclude_fields:
                update_data[key] = value
        
        # Update timestamp
        update_data['updated_at'] = datetime.utcnow()
        
        # Update user in database
        if mongo_client.client is not None and mongo_client.db is not None:
            collection = mongo_client.db["User-Base"]
            if username:
                result = collection.update_one(
                    {"username": username},
                    {"$set": update_data}
                )
            else:
                result = collection.update_one(
                    {"email": email},
                    {"$set": update_data}
                )
            
            if result.modified_count > 0:
                # Fetch updated user
                if username:
                    updated_user = mongo_client.get_user_by_username(username)
                else:
                    updated_user = mongo_client.get_user_by_email(email)
                
                # Remove password hash from response
                user_data = {k: v for k, v in updated_user.items() if k != 'password_hash'}
                
                return jsonify({
                    'success': True,
                    'message': 'Profile updated successfully',
                    'user': user_data
                }), 200
            else:
                return jsonify({'error': 'No changes were made'}), 400
        
        return jsonify({'error': 'Database connection failed'}), 500
        
    except Exception as e:
        logger.error(f"Error updating profile: {e}", exc_info=True)
        return jsonify({'error': f'Failed to update profile: {str(e)}'}), 500


@app.route('/api/create-account', methods=['POST'])
@require_api_key
def create_account():
    """Create a new user account and store in User-Base collection."""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        # Validate required fields
        required_fields = ['username', 'fullName', 'email', 'phoneNumber', 'password', 'country', 'preferredLanguage']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'error': f'Missing required field: {field}'}), 400
        
        # Validate username format
        username = data.get('username', '').strip()
        if len(username) < 3:
            return jsonify({'error': 'Username must be at least 3 characters'}), 400
        if not username.replace('_', '').isalnum():
            return jsonify({'error': 'Username can only contain letters, numbers, and underscores'}), 400
        
        # Check if username already exists
        if mongo_client.client is not None and mongo_client.db is not None:
            try:
                collection = mongo_client.db["User-Base"]
                existing_username = collection.find_one({"username": username})
                if existing_username:
                    return jsonify({'error': 'Username already taken'}), 409
            except Exception as e:
                logger.error(f"Error checking username: {e}")
                return jsonify({'error': 'Database error while checking username'}), 500
        
        # Check if user already exists by email
        existing_user = mongo_client.get_user_by_email(data.get('email'))
        if existing_user:
            return jsonify({'error': 'User with this email already exists'}), 409
        
        # Hash password before storing
        password_hash = bcrypt.hashpw(data.get('password').encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        
        # Prepare user data for storage
        user_data = {
            # Step 1: Personal & Account Information
            'username': username,  # Unique username
            'fullName': data.get('fullName'),
            'email': data.get('email'),
            'phoneNumber': data.get('phoneNumber'),
            'password_hash': password_hash,  # Store hashed password, not plain text
            'country': data.get('country'),
            'preferredLanguage': data.get('preferredLanguage'),
            'profilePicture': data.get('profilePicturePreview'),  # Store base64 encoded image if provided
            
            # Step 2: Professional & Company Details
            'userType': data.get('userType'),
            'organizationName': data.get('organizationName'),
            'organizationType': data.get('organizationType'),
            'industry': data.get('industry'),
            'companySize': data.get('companySize'),
            'website': data.get('website'),
            'designation': data.get('designation'),
            'department': data.get('department'),
            'yearsOfExperience': data.get('yearsOfExperience'),
            'complianceFrameworks': data.get('complianceFrameworks', []),
            'teamSize': data.get('teamSize'),
            'companyLocation': data.get('companyLocation'),
            
            # Step 3: Use Case & Motivation
            'primaryGoal': data.get('primaryGoal'),
            'reason': data.get('reason'),
            'dataTypes': data.get('dataTypes', []),
            'usageEnvironment': data.get('usageEnvironment'),
            'averageVolume': data.get('averageVolume'),
            'needApiAccess': data.get('needApiAccess', False),
            'internalUsers': data.get('internalUsers', []),
            'deploymentPlan': data.get('deploymentPlan'),
            
            # Step 4: Security & Consent
            'enable2FA': data.get('enable2FA', False),
            'preferredAuthMethod': data.get('preferredAuthMethod'),
            'acceptTerms': data.get('acceptTerms', False),
            'acceptPrivacy': data.get('acceptPrivacy', False),
            'consentDataProcessing': data.get('consentDataProcessing', False),
            'receiveUpdates': data.get('receiveUpdates', False),
        }
        
        # Create user in database
        user_id = mongo_client.create_user(user_data)
        
        if not user_id:
            return jsonify({'error': 'Failed to create user account'}), 500
        
        logger.info(f"User account created successfully: {user_id} for email: {data.get('email')}")
        
        # Return success response (don't include password hash)
        response_data = {
            'success': True,
            'user_id': user_id,
            'message': 'Account created successfully',
            'email': data.get('email')
        }
        
        return jsonify(response_data), 201
        
    except Exception as e:
        logger.error(f"Error creating account: {e}", exc_info=True)
        return jsonify({'error': f'Failed to create account: {str(e)}'}), 500


@app.route('/api/activate-plan', methods=['POST'])
@require_api_key
def activate_plan():
    """Activate a subscription plan for a user."""
    try:
        data = request.get_json()
        plan_name = data.get('plan_name')
        billing_period = data.get('billing_period', 'monthly')
        email = data.get('email')
        
        if not plan_name or not email:
            return jsonify({'error': 'plan_name and email are required'}), 400
        
        # Plan pricing mapping
        plan_prices = {
            'Starter': {'monthly': 0, 'yearly': 0},
            'Professional': {'monthly': 999, 'yearly': 9999},
            'Enterprise': {'monthly': 4999, 'yearly': 49999}
        }
        
        if plan_name not in plan_prices:
            return jsonify({'error': 'Invalid plan name'}), 400
        
        amount = plan_prices[plan_name][billing_period]
        plan_type_map = {
            'Starter': 'starter',
            'Professional': 'professional',
            'Enterprise': 'enterprise'
        }
        plan_type = plan_type_map.get(plan_name, plan_name.lower())
        
        success = mongo_client.update_user_subscription(
            email=email,
            plan_name=plan_name,
            plan_type=plan_type,
            amount=amount,
            billing_period=billing_period
        )
        
        if success:
            logger.info(f"Plan activated for {email}, plan: {plan_name}")
            return jsonify({
                'success': True,
                'message': 'Plan activated successfully',
                'plan_name': plan_name,
                'plan_type': plan_type,
                'amount': amount
            }), 200
        else:
            return jsonify({'error': 'Failed to activate plan'}), 500
        
    except Exception as e:
        logger.error(f"Error activating plan: {e}", exc_info=True)
        return jsonify({'error': f'Failed to activate plan: {str(e)}'}), 500


@app.route('/api/subscription', methods=['GET'])
@require_api_key
def get_subscription():
    """Get user subscription details."""
    try:
        email = request.args.get('email')
        
        if not email:
            return jsonify({'error': 'email is required'}), 400
        
        subscription = mongo_client.get_user_subscription(email)
        
        if not subscription:
            return jsonify({'error': 'User not found'}), 404
        
        return jsonify({
            'success': True,
            'subscription': subscription
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting subscription: {e}", exc_info=True)
        return jsonify({'error': f'Failed to get subscription: {str(e)}'}), 500


@app.errorhandler(500)
def internal_error(error):
    """Handle internal server errors."""
    logger.error(f"Internal error: {error}", exc_info=True)
    return jsonify({'error': 'Internal server error'}), 500


# ============================================
# SETTINGS ROUTES
# ============================================


def _generate_user_data_pdf(email: str, export_payload: Dict[str, Any]) -> Tuple[str, str]:
    """Generate a branded PDF export for user data and return (path, filename)."""
    from io import BytesIO
    from textwrap import wrap
    import tempfile
    import glob
    import re

    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import mm
    except Exception as exc:  # pragma: no cover - dependency failure
        logger.error("ReportLab dependency missing for data export", exc_info=True)
        raise exc

    profile = export_payload.get('profile') or {}
    token_summary = export_payload.get('token_summary') or {}
    activity = export_payload.get('activity') or {}
    preferences = export_payload.get('preferences') or {}
    batches = export_payload.get('batches') or []
    generated_at = export_payload.get('generated_at', datetime.utcnow().isoformat(timespec='seconds') + 'Z')

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin = 20 * mm
    brand_color = colors.HexColor('#4338CA')
    text_color = colors.HexColor('#0F172A')
    muted_color = colors.HexColor('#475569')

    def format_value(value: Any) -> str:
        if value is None or value == '':
            return 'Not available'
        if isinstance(value, bool):
            return 'Enabled' if value else 'Disabled'
        return str(value)

    def draw_header():
        pdf.setFillColor(brand_color)
        pdf.setFont("Helvetica-Bold", 20)
        pdf.drawString(margin, height - 25 * mm, "PII Sentinel")
        pdf.setFillColor(text_color)
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(margin, height - 35 * mm, "Personal Data Export")
        pdf.setFillColor(muted_color)
        pdf.setFont("Helvetica", 9)
        owner_line = f"Generated for {profile.get('fullName') or profile.get('username') or email}"
        pdf.drawString(margin, height - 40 * mm, owner_line)
        pdf.drawRightString(width - margin, height - 40 * mm, generated_at)

    y_position = height - 50 * mm

    def ensure_space(space: float):
        nonlocal y_position
        if y_position - space < 25 * mm:
            pdf.showPage()
            draw_header()
            y_position = height - 50 * mm

    def write_section(title: str):
        nonlocal y_position
        ensure_space(12 * mm)
        pdf.setFillColor(brand_color)
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(margin, y_position, title)
        y_position -= 7 * mm

    def write_key_value(label: str, value: Any, wrap_width: int = 90):
        nonlocal y_position
        ensure_space(9 * mm)
        pdf.setFillColor(text_color)
        pdf.setFont("Helvetica-Bold", 10)
        pdf.drawString(margin, y_position, f"{label}:")
        y_position -= 4.5 * mm
        pdf.setFont("Helvetica", 9)
        pdf.setFillColor(muted_color)
        wrapped = wrap(format_value(value), wrap_width)
        for line in wrapped or ['Not available']:
            ensure_space(6 * mm)
            pdf.drawString(margin, y_position, line)
            y_position -= 4.5 * mm

    draw_header()

    # Account summary
    write_section("Account Summary")
    write_key_value("Full name", profile.get('fullName'))
    write_key_value("Email", profile.get('email', email))
    write_key_value("Username", profile.get('username'))
    write_key_value("Plan", profile.get('plan_id') or profile.get('plan') or profile.get('subscription', {}).get('plan_id'))
    write_key_value("Account status", profile.get('account_status') or profile.get('status') or 'Active')

    # Preference summary
    write_section("Communication Preferences")
    write_key_value("Email updates", preferences.get('receiveUpdates'))
    write_key_value("Data collection consent", preferences.get('consentDataProcessing'))

    # Token summary
    if token_summary:
        write_section("Token Overview")
        write_key_value("Tokens total", token_summary.get('tokens_total'))
        write_key_value("Tokens used", token_summary.get('tokens_used'))
        write_key_value("Tokens balance", token_summary.get('tokens_balance'))
        write_key_value("Last reset", token_summary.get('last_token_reset'))

    # Activity summary
    if activity:
        write_section("Recent Activity")
        write_key_value("Last batch created", activity.get('lastBatchCreated'))
        write_key_value("Last scan completed", activity.get('lastPiiScanCompleted'))

    # Usage stats
    if batches:
        write_section("Usage Snapshot")
        total_batches = len(batches)
        total_files = sum(len(batch.get('files', [])) for batch in batches if isinstance(batch.get('files'), list))
        total_piis = 0
        for batch in batches:
            stats = batch.get('stats') or {}
            try:
                total_piis += int(stats.get('piis', 0))
            except (TypeError, ValueError):
                continue
        write_key_value("Batches exported (latest 50)", total_batches)
        write_key_value("Files processed", total_files)
        write_key_value("PIIs detected", total_piis)

        recent_batches = batches[:5]
        if recent_batches:
            write_section("Recent Batches")
            for batch in recent_batches:
                ensure_space(12 * mm)
                pdf.setFillColor(text_color)
                pdf.setFont("Helvetica-Bold", 10)
                batch_label = batch.get('batch_id') or batch.get('_id') or 'Batch'
                pdf.drawString(margin, y_position, f"• {batch_label}")
                y_position -= 4.5 * mm
                pdf.setFont("Helvetica", 9)
                pdf.setFillColor(muted_color)
                details = [
                    ("Created", batch.get('created_at')),
                    ("Status", batch.get('status', 'unknown')),
                    ("Files", len(batch.get('files', [])) if isinstance(batch.get('files'), list) else 'N/A'),
                ]
                stats = batch.get('stats') or {}
                if stats:
                    details.append(("PIIs detected", stats.get('piis')))
                for label, value in details:
                    ensure_space(5 * mm)
                    pdf.drawString(margin + 5 * mm, y_position, f"{label}: {format_value(value)}")
                    y_position -= 4.0 * mm

    ensure_space(12 * mm)
    pdf.setFillColor(muted_color)
    pdf.setFont("Helvetica-Oblique", 8)
    pdf.drawString(
        margin,
        y_position,
        "If you need assistance interpreting this export, reach out to support@pii-sentinel.com."
    )

    pdf.save()
    buffer.seek(0)

    safe_email = re.sub(r'[^A-Za-z0-9]+', '-', email).strip('-') or 'user'
    timestamp = datetime.utcnow().strftime('%Y%m%d%H%M%S')
    file_name = f"PII-Sentinel-DataExport-{safe_email}-{timestamp}.pdf"

    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, file_name)

    with open(file_path, 'wb') as fp:
        fp.write(buffer.getvalue())

    # Clean up older exports for same user (older than 1 hour)
    pattern = os.path.join(temp_dir, f"PII-Sentinel-DataExport-{safe_email}-*.pdf")
    expiry_seconds = 3600
    now = time.time()
    for stale_file in glob.glob(pattern):
        if stale_file == file_path:
            continue
        try:
            if now - os.path.getmtime(stale_file) > expiry_seconds:
                os.remove(stale_file)
        except OSError:
            continue

    return file_path, file_name

@app.route('/api/settings/update-status', methods=['POST'])
@require_api_key
def update_account_status():
    """Update user account status."""
    try:
        data = request.get_json()
        email = data.get('email')
        status = data.get('status', 'active')
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
        
        if status not in ['active', 'inactive', 'logged_out']:
            return jsonify({'error': 'Invalid status'}), 400
        
        success = mongo_client.update_user_status(email, status)
        if success:
            user = mongo_client.get_user_by_email(email)
            if user and '_id' in user:
                user['_id'] = str(user['_id'])
            return jsonify({
                'success': True,
                'message': 'Account status updated successfully',
                'user': user
            }), 200
        else:
            return jsonify({'error': 'Failed to update account status'}), 500
    
    except Exception as e:
        logger.error(f"Error updating account status: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/settings/update-plan', methods=['POST'])
@require_api_key
def update_plan():
    """Update user subscription plan."""
    try:
        data = request.get_json()
        email = data.get('email')
        plan = data.get('plan', 'free')
        billing_period = data.get('billingPeriod', 'monthly')
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
        
        if plan not in ['free', 'pro', 'enterprise']:
            return jsonify({'error': 'Invalid plan'}), 400
        
        if billing_period not in ['monthly', 'yearly']:
            return jsonify({'error': 'Invalid billing period'}), 400
        
        success = mongo_client.update_user_plan(email, plan, billing_period)
        if success:
            user = mongo_client.get_user_by_email(email)
            if user and '_id' in user:
                user['_id'] = str(user['_id'])
            token_summary = mongo_client.get_token_summary(email) or {}
            return jsonify({
                'success': True,
                'message': 'Plan updated successfully',
                'user': user,
                'tokens': token_summary
            }), 200
        else:
            return jsonify({'error': 'Failed to update plan'}), 500
    
    except Exception as e:
        logger.error(f"Error updating plan: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# ============================================================
# Token / User Info Endpoints
# ============================================================

@app.route('/api/user/tokens', methods=['GET'])
@require_api_key
def get_user_tokens():
    """Get user token summary."""
    try:
        user_id = request.args.get('user_id') or request.args.get('email')
        if not user_id:
            return jsonify({'error': 'user_id or email required'}), 400
        
        user_id = str(user_id).strip().lower()
        
        # Ensure user has token document initialized
        mongo_client.ensure_user_token_document(user_id)
        
        token_summary = mongo_client.get_token_summary(user_id)
        
        if not token_summary:
            # User not found or token summary unavailable
            return jsonify({'error': 'User not found'}), 404
        
        # Debug logging
        logger.info(f"Token summary for {user_id}:")
        logger.info(f"  - plan_id: {token_summary.get('plan_id')}")
        logger.info(f"  - features_enabled: {token_summary.get('features_enabled')}")
        
        return jsonify({'tokens': token_summary}), 200
    
    except Exception as e:
        logger.error(f"Error fetching token summary: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/user/activity', methods=['GET'])
@require_api_key
def get_user_activity():
    """Get user activity statistics."""
    try:
        user_id = request.args.get('user_id') or request.args.get('email')
        if not user_id:
            return jsonify({'error': 'user_id or email required'}), 400
        
        user_id = str(user_id).strip().lower()
        
        # Get user info
        user = mongo_client.get_user_by_email(user_id)
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        # Get last login from user document
        last_login = user.get('last_login') or user.get('updated_at') or user.get('created_at')
        
        # Get batches for this user - try different user_id formats
        batches_collection = mongo_client.db["Batch-Base"]
        
        # Try finding by exact user_id match or by email
        batches = list(batches_collection.find({
            "$or": [
                {"user_id": user_id},
                {"user_id": user.get('email')},
                {"user_id": user.get('username')}
            ]
        }).sort("created_at", -1).limit(10))
        
        logger.info(f"Found {len(batches)} batches for user {user_id}")
        
        # Get last batch created
        last_batch_created = None
        if batches:
            last_batch = batches[0]
            last_batch_created = last_batch.get('created_at') or last_batch.get('updated_at')
            logger.info(f"Last batch created at: {last_batch_created}")
        
        # Get last PII scan completed - check for any batch with files processed
        last_pii_scan = None
        for batch in batches:
            # Check if batch has processed files
            if batch.get('processed_at') or batch.get('files') or batch.get('status') == 'completed':
                last_pii_scan = batch.get('processed_at') or batch.get('updated_at') or batch.get('created_at')
                logger.info(f"Last PII scan completed at: {last_pii_scan}")
                break
        
        # Calculate account strength
        strength_score = 0
        strength_total = 3
        
        # Email verified
        if user.get('email') and user.get('emailVerified') != False:
            strength_score += 1
        
        # Profile complete - check required fields
        required_fields = ['fullName', 'email', 'country']
        profile_complete = all(user.get(field) for field in required_fields)
        if profile_complete:
            strength_score += 1
        
        logger.info(f"Profile completeness check - fullName: {user.get('fullName')}, email: {user.get('email')}, country: {user.get('country')}")
        
        # 2FA enabled
        if user.get('enable2FA') or user.get('twoFactorEnabled'):
            strength_score += 1
        
        account_strength = round((strength_score / strength_total) * 100)
        
        # Format timestamps
        def format_timestamp(ts):
            if isinstance(ts, datetime):
                return ts.isoformat()
            elif isinstance(ts, str):
                return ts
            elif ts is None:
                return None
            else:
                return str(ts)
        
        return jsonify({
            'success': True,
            'activity': {
                'lastLogin': format_timestamp(last_login),
                'lastBatchCreated': format_timestamp(last_batch_created),
                'lastPiiScanCompleted': format_timestamp(last_pii_scan)
            },
            'accountStrength': account_strength,
            'strengthDetails': {
                'emailVerified': user.get('email') and user.get('emailVerified') != False,
                'profileComplete': profile_complete,
                'twoFactorEnabled': user.get('enable2FA') or user.get('twoFactorEnabled')
            }
        }), 200
    
    except Exception as e:
        logger.error(f"Error fetching user activity: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/settings/update-security', methods=['POST'])
@require_api_key
def update_security():
    """Update user security settings."""
    try:
        data = request.get_json()
        email = data.get('email')
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
        
        security_data = {k: v for k, v in data.items() if k != 'email'}
        
        updated_user = mongo_client.update_user_security(email, security_data)
        if updated_user:
            if '_id' in updated_user:
                updated_user['_id'] = str(updated_user['_id'])
            # Remove password hash from response
            updated_user.pop('password_hash', None)
            return jsonify({
                'success': True,
                'message': 'Security settings updated successfully',
                'user': updated_user
            }), 200
        else:
            return jsonify({'error': 'Failed to update security settings'}), 500
    
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"Error updating security settings: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/settings/update-preferences', methods=['POST'])
@require_api_key
def update_preferences():
    """Update user preferences."""
    try:
        data = request.get_json()
        email = data.get('email')
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
        
        preferences = {
            'emailUpdates': data.get('emailUpdates'),
            'dataConsent': data.get('dataConsent')
        }
        
        updated_user = mongo_client.update_user_preferences(email, preferences)
        if updated_user:
            if '_id' in updated_user:
                updated_user['_id'] = str(updated_user['_id'])
            return jsonify({
                'success': True,
                'message': 'Preferences updated successfully',
                'user': updated_user
            }), 200
        else:
            return jsonify({'error': 'Failed to update preferences'}), 500
    
    except Exception as e:
        logger.error(f"Error updating preferences: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/settings/download-data', methods=['POST'])
@require_api_key
def download_user_data():
    """Download user data as ZIP file."""
    try:
        data = request.get_json()
        email = data.get('email')
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
        
        # Get all user data
        user_data = mongo_client.get_user_data_for_export(email)
        if not user_data:
            return jsonify({'error': 'User not found'}), 404

        try:
            _, file_name = _generate_user_data_pdf(email, user_data)
        except Exception as exc:
            logger.error(f"Failed to generate data export PDF: {exc}", exc_info=True)
            return jsonify({'error': 'Failed to generate data export'}), 500
        
        return jsonify({
            'success': True,
            'downloadUrl': f'/api/settings/download-file?file={file_name}',
            'fileName': file_name,
            'message': 'Data export prepared successfully'
        }), 200
    
    except Exception as e:
        logger.error(f"Error downloading user data: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/settings/download-file', methods=['GET'])
@require_api_key
def download_settings_file():
    """Serve downloaded file."""
    try:
        file_name = request.args.get('file')
        if not file_name:
            return jsonify({'error': 'File name required'}), 400
        
        # Security: Validate file name to prevent directory traversal
        if '..' in file_name or '/' in file_name or '\\' in file_name:
            return jsonify({'error': 'Invalid file name'}), 400
        
        # In production, retrieve from secure storage (S3, etc.)
        # For now, serve from temp directory (implement proper secure file serving)
        import tempfile
        import glob
        
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, file_name)
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404
        
        return send_file(file_path, as_attachment=True, download_name=file_name)
    
    except Exception as e:
        logger.error(f"Error serving file: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/settings/clear-activity', methods=['POST'])
@require_api_key
def clear_activity():
    """Clear user activity logs."""
    try:
        data = request.get_json()
        email = data.get('email')
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
        
        success = mongo_client.clear_user_activity_logs(email)
        if success:
            return jsonify({
                'success': True,
                'message': 'Activity logs cleared successfully'
            }), 200
        else:
            return jsonify({'error': 'Failed to clear activity logs'}), 500
    
    except Exception as e:
        logger.error(f"Error clearing activity logs: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/settings/logout', methods=['POST'])
@require_api_key
def logout():
    """Logout user and update status."""
    try:
        data = request.get_json()
        email = data.get('email')
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
        
        # Update status to logged_out
        success = mongo_client.update_user_status(email, 'logged_out')
        if success:
            return jsonify({
                'success': True,
                'message': 'Logged out successfully'
            }), 200
        else:
            return jsonify({'error': 'Failed to logout'}), 500
    
    except Exception as e:
        logger.error(f"Error logging out: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/settings/delete-account', methods=['POST'])
@require_api_key
def delete_account():
    """Delete user account."""
    try:
        data = request.get_json()
        email = data.get('email')
        reason = data.get('reason', 'Not specified')
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
        
        success = mongo_client.delete_user_account(email, reason)
        if success:
            return jsonify({
                'success': True,
                'message': 'Account deleted successfully'
            }), 200
        else:
            return jsonify({'error': 'Failed to delete account'}), 500
    
    except Exception as e:
        logger.error(f"Error deleting account: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/migrate-user-features', methods=['POST'])
@require_api_key
def migrate_user_features():
    """
    Migration endpoint to ensure all users have complete features_enabled fields.
    This will update all users to include export_json and log_records fields.
    """
    try:
        if mongo_client.db is None:
            return jsonify({'error': 'Database not available'}), 503
        
        collection = mongo_client.db["User-Base"]
        all_users = collection.find({})
        
        updated_count = 0
        skipped_count = 0
        
        for user in all_users:
            email = user.get('email')
            if not email:
                continue
            
            plan_id = user.get('plan_id', 'starter')
            plan = mongo_client.get_plan(plan_id) or {}
            plan_features = plan.get('features', {
                "lock_json": False,
                "unlock_json": False,
                "advanced_analysis": False,
                "export_json": False,
                "log_records": False
            })
            
            current_features = user.get('features_enabled', {})
            if not isinstance(current_features, dict):
                current_features = {}
            
            # Check if any features are missing
            missing_features = []
            for feature_key in plan_features.keys():
                if feature_key not in current_features:
                    missing_features.append(feature_key)
            
            if missing_features:
                # Merge: add missing features with plan defaults, keep existing values
                merged_features = {**plan_features, **current_features}
                
                collection.update_one(
                    {"email": email},
                    {
                        "$set": {
                            "features_enabled": merged_features,
                            "updated_at": datetime.utcnow()
                        }
                    }
                )
                updated_count += 1
                logger.info(f"✓ Updated features for {email}: added {missing_features}")
            else:
                skipped_count += 1
        
        return jsonify({
            'success': True,
            'updated': updated_count,
            'skipped': skipped_count,
            'message': f'Migration complete: {updated_count} users updated, {skipped_count} already up-to-date'
        }), 200
    
    except Exception as e:
        logger.error(f"Error migrating user features: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# ============================================================================
# RAZORPAY PAYMENT ENDPOINTS
# ============================================================================

@app.route('/api/payment/create-order', methods=['POST'])
@require_api_key
@require_auth
def create_payment_order():
    """Create a Razorpay order for plan upgrade."""
    try:
        data = request.get_json()
        plan_id = data.get('plan_id')
        user_email = data.get('email') or data.get('user_id')
        billing_period = data.get('billing_period', 'monthly')  # 'monthly' or 'yearly'
        notes = data.get('notes', {})
        
        if not plan_id:
            return jsonify({'error': 'Plan ID is required'}), 400
        if not user_email:
            return jsonify({'error': 'User email is required'}), 400
        
        plan = PLAN_CATALOG.get(plan_id)
        if not plan:
            return jsonify({'error': f'Invalid plan: {plan_id}'}), 400
        
        user = mongo_client.get_user_by_email(user_email)
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        from payments.razorpay_service import RazorpayService
        razorpay = RazorpayService.from_env()
        
        if not razorpay.enabled:
            return jsonify({'error': 'Payment service not available'}), 503
        
        # Get price based on billing period
        if billing_period == 'yearly':
            amount = plan.get('price_inr_yearly', plan['price_inr'] * 12)
        else:
            amount = plan['price_inr']
            
        # Generate short unique receipt (max 40 chars for Razorpay)
        receipt = f"pln_{plan_id[:8]}_{uuid.uuid4().hex[:22]}"
        
        order = razorpay.create_order(
            amount_inr=float(amount),
            receipt=receipt,
            notes={
                **notes, 
                'user_id': user_email, 
                'plan_id': plan_id, 
                'plan_name': plan['name'],
                'billing_period': billing_period
            }
        )
        
        logger.info(f"✓ Payment order created: {order.order_id} for {user_email} - Plan: {plan_id} ({billing_period}) (₹{amount})")
        
        return jsonify({
            'success': True,
            'order_id': order.order_id,
            'amount': order.amount,
            'currency': order.currency,
            'key': order.key,
            'plan_name': plan['name'],
            'plan_id': plan_id,
            'billing_period': billing_period
        })
    except Exception as e:
        logger.error(f"Error creating payment order: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/payment/create-token-order', methods=['POST'])
@require_api_key
@require_auth
def create_token_addon_order():
    """Create a Razorpay order for token addon purchase. Token price: ₹20 per token"""
    try:
        data = request.get_json()
        token_amount = data.get('token_amount')
        user_email = data.get('email') or data.get('user_id')
        notes = data.get('notes', {})
        
        if not token_amount or token_amount <= 0:
            return jsonify({'error': 'Valid token amount is required'}), 400
        if not user_email:
            return jsonify({'error': 'User email is required'}), 400
        
        user = mongo_client.get_user_by_email(user_email)
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        from payments.razorpay_service import RazorpayService
        razorpay = RazorpayService.from_env()
        
        if not razorpay.enabled:
            return jsonify({'error': 'Payment service not available'}), 503
        
        amount = token_amount * ADDON_TOKEN_PRICE_INR
        # Generate short unique receipt (max 40 chars for Razorpay)
        receipt = f"tok_{token_amount}_{uuid.uuid4().hex[:22]}"
        
        order = razorpay.create_order(
            amount_inr=float(amount),
            receipt=receipt,
            notes={**notes, 'user_id': user_email, 'token_amount': token_amount, 'type': 'token_addon'}
        )
        
        logger.info(f"✓ Token addon order created: {order.order_id} for {user_email} - {token_amount} tokens (₹{amount})")
        
        return jsonify({
            'success': True,
            'order_id': order.order_id,
            'amount': order.amount,
            'currency': order.currency,
            'key': order.key,
            'token_amount': token_amount,
            'price_per_token': ADDON_TOKEN_PRICE_INR
        })
    except Exception as e:
        logger.error(f"Error creating token addon order: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/payment/verify', methods=['POST'])
@require_api_key
@require_auth
def verify_payment():
    """Verify payment signature and update user's plan or tokens."""
    try:
        data = request.get_json()
        razorpay_order_id = data.get('razorpay_order_id')
        razorpay_payment_id = data.get('razorpay_payment_id')
        razorpay_signature = data.get('razorpay_signature')
        user_email = data.get('email')
        plan_id = data.get('plan_id')
        token_amount = data.get('token_amount')
        
        if not all([razorpay_order_id, razorpay_payment_id, razorpay_signature]):
            return jsonify({'error': 'Payment details are incomplete'}), 400
        if not user_email:
            return jsonify({'error': 'User email is required'}), 400
        
        from payments.razorpay_service import RazorpayService
        razorpay = RazorpayService.from_env()
        
        if not razorpay.enabled:
            return jsonify({'error': 'Payment service not available'}), 503
        
        is_valid = razorpay.verify_payment_signature(
            razorpay_order_id=razorpay_order_id,
            razorpay_payment_id=razorpay_payment_id,
            razorpay_signature=razorpay_signature
        )
        
        if not is_valid:
            logger.error(f"❌ Payment signature verification failed for {user_email}")
            return jsonify({'error': 'Payment verification failed'}), 400
        
        logger.info(f"✓ Payment signature verified: {razorpay_payment_id} for {user_email}")
        
        payment = razorpay.fetch_payment(razorpay_payment_id)
        if not payment:
            logger.error(f"Failed to fetch payment details: {razorpay_payment_id}")
            return jsonify({'error': 'Failed to fetch payment details'}), 500
        
        if plan_id:
            plan = PLAN_CATALOG.get(plan_id)
            if not plan:
                return jsonify({'error': f'Invalid plan: {plan_id}'}), 400
            
            # Get billing period from payment notes (set during order creation)
            billing_period = 'monthly'
            if payment:
                notes = payment.get('notes', {})
                billing_period = notes.get('billing_period', 'monthly')
            
            logger.info(f"🔄 Updating plan for {user_email} to {plan_id} ({billing_period})")
            result = mongo_client.update_user_plan(user_email, plan_id, billing_period)
            if result:
                logger.info(f"✅ User {user_email} upgraded to {plan_id} plan ({billing_period})")
                invalidate_profile_cache(user_email)
                
                # Save payment history
                mongo_client.save_payment_history({
                    "user_email": user_email,
                    "payment_id": razorpay_payment_id,
                    "order_id": razorpay_order_id,
                    "amount": payment.get('amount', 0),
                    "currency": payment.get('currency', 'INR'),
                    "type": "plan_upgrade",
                    "plan_id": plan_id,
                    "plan_name": plan['name'],
                    "billing_period": billing_period,
                    "payment_method": payment.get('method', 'razorpay')
                })
                
                return jsonify({
                    'success': True,
                    'message': f'Payment successful! Upgraded to {plan["name"]} plan',
                    'plan_id': plan_id,
                    'plan_name': plan['name'],
                    'payment_id': razorpay_payment_id,
                    'billing_period': billing_period
                })
            else:
                logger.error(f"Failed to update plan for {user_email}")
                return jsonify({'error': 'Failed to update plan'}), 500
                
        elif token_amount:
            result = mongo_client.credit_tokens(user_email, token_amount, 'addon_purchase')
            if result:
                logger.info(f"✅ Added {token_amount} tokens to {user_email}")
                invalidate_profile_cache(user_email)
                
                # Save payment history
                mongo_client.save_payment_history({
                    "user_email": user_email,
                    "payment_id": razorpay_payment_id,
                    "order_id": razorpay_order_id,
                    "amount": payment.get('amount', 0),
                    "currency": payment.get('currency', 'INR'),
                    "type": "token_addon",
                    "token_amount": token_amount,
                    "payment_method": payment.get('method', 'razorpay')
                })
                
                return jsonify({
                    'success': True,
                    'message': f'Payment successful! Added {token_amount} tokens to your account',
                    'tokens_added': token_amount,
                    'payment_id': razorpay_payment_id
                })
            else:
                logger.error(f"Failed to credit tokens for {user_email}")
                return jsonify({'error': 'Failed to credit tokens'}), 500
        else:
            return jsonify({'error': 'Invalid payment type'}), 400
            
    except Exception as e:
        logger.error(f"Error verifying payment: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/payment/history', methods=['GET'])
@require_api_key
@require_auth
def get_payment_history():
    """Get payment history for user."""
    try:
        user_email = request.args.get('email')
        if not user_email:
            return jsonify({'error': 'User email is required'}), 400
        
        payments = mongo_client.get_payment_history(user_email)
        
        return jsonify({
            'success': True,
            'payments': payments
        })
    except Exception as e:
        logger.error(f"Error fetching payment history: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/payment/invoice/<payment_id>', methods=['GET'])
@require_api_key
@require_auth
def download_invoice(payment_id):
    """Generate and download invoice PDF for a payment."""
    try:
        user_email = request.args.get('email')
        if not user_email:
            return jsonify({'error': 'User email is required'}), 400
        
        # Get payment details from history
        payments = mongo_client.get_payment_history(user_email)
        payment = next((p for p in payments if p.get('payment_id') == payment_id), None)
        
        if not payment:
            return jsonify({'error': 'Payment not found'}), 404
        
        # Get user details
        user = mongo_client.get_user_by_email(user_email)
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        # Generate simple HTML invoice
        amount_inr = payment.get('amount', 0) / 100  # Convert paise to rupees
        invoice_html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Invoice - {payment_id}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }}
        .header {{ text-align: center; border-bottom: 2px solid #4F46E5; padding-bottom: 20px; margin-bottom: 30px; }}
        .company-name {{ font-size: 28px; font-weight: bold; color: #4F46E5; }}
        .invoice-title {{ font-size: 20px; margin-top: 10px; }}
        .section {{ margin: 20px 0; }}
        .section-title {{ font-weight: bold; font-size: 16px; margin-bottom: 10px; color: #333; }}
        .info-row {{ display: flex; justify-content: space-between; padding: 8px 0; border-bottom: 1px solid #eee; }}
        .label {{ font-weight: 600; color: #666; }}
        .value {{ color: #333; }}
        .total {{ font-size: 20px; font-weight: bold; color: #4F46E5; margin-top: 20px; padding-top: 20px; border-top: 2px solid #4F46E5; }}
        .footer {{ text-align: center; margin-top: 40px; padding-top: 20px; border-top: 1px solid #eee; color: #666; font-size: 12px; }}
    </style>
</head>
<body>
    <div class="header">
        <div class="company-name">PII SENTINEL</div>
        <div class="invoice-title">PAYMENT INVOICE</div>
    </div>
    
    <div class="section">
        <div class="section-title">Customer Information</div>
        <div class="info-row">
            <span class="label">Name:</span>
            <span class="value">{user.get('fullName', 'N/A')}</span>
        </div>
        <div class="info-row">
            <span class="label">Email:</span>
            <span class="value">{user_email}</span>
        </div>
        <div class="info-row">
            <span class="label">Date:</span>
            <span class="value">{payment.get('created_at', datetime.utcnow()).strftime('%B %d, %Y %I:%M %p')}</span>
        </div>
    </div>
    
    <div class="section">
        <div class="section-title">Payment Details</div>
        <div class="info-row">
            <span class="label">Payment ID:</span>
            <span class="value">{payment_id}</span>
        </div>
        <div class="info-row">
            <span class="label">Order ID:</span>
            <span class="value">{payment.get('order_id', 'N/A')}</span>
        </div>
        <div class="info-row">
            <span class="label">Payment Method:</span>
            <span class="value">{payment.get('payment_method', 'Razorpay').upper()}</span>
        </div>
        <div class="info-row">
            <span class="label">Status:</span>
            <span class="value">{payment.get('status', 'completed').upper()}</span>
        </div>
    </div>
    
    <div class="section">
        <div class="section-title">Transaction Details</div>
        """
        
        if payment.get('type') == 'plan_upgrade':
            invoice_html += f"""
        <div class="info-row">
            <span class="label">Type:</span>
            <span class="value">Plan Upgrade</span>
        </div>
        <div class="info-row">
            <span class="label">Plan:</span>
            <span class="value">{payment.get('plan_name', 'N/A')}</span>
        </div>
        <div class="info-row">
            <span class="label">Billing Period:</span>
            <span class="value">{payment.get('billing_period', 'monthly').capitalize()}</span>
        </div>
            """
        elif payment.get('type') == 'token_addon':
            invoice_html += f"""
        <div class="info-row">
            <span class="label">Type:</span>
            <span class="value">Token Addon Purchase</span>
        </div>
        <div class="info-row">
            <span class="label">Tokens Added:</span>
            <span class="value">{payment.get('token_amount', 0)} tokens</span>
        </div>
            """
        
        invoice_html += f"""
    </div>
    
    <div class="total">
        <div class="info-row" style="border: none;">
            <span class="label">TOTAL AMOUNT PAID:</span>
            <span class="value">₹{amount_inr:.2f}</span>
        </div>
    </div>
    
    <div class="footer">
        <p>Thank you for your business!</p>
        <p>PII SENTINEL - Secure PII Detection & Protection</p>
        <p>For support, contact: support@piisentinel.com</p>
    </div>
</body>
</html>
        """
        
        # Return HTML as downloadable file
        response = make_response(invoice_html)
        response.headers['Content-Type'] = 'text/html'
        response.headers['Content-Disposition'] = f'attachment; filename=invoice_{payment_id}.html'
        
        return response
        
    except Exception as e:
        logger.error(f"Error generating invoice: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# TEMPORARY ADMIN ENDPOINT - DELETE AFTER USE
@app.route('/api/admin/fix-plan', methods=['POST'])
def admin_fix_plan():
    """TEMPORARY: Manually fix user plan after failed payment update."""
    try:
        data = request.get_json()
        admin_key = data.get('admin_key')
        user_email = data.get('email')
        plan_id = data.get('plan_id', 'professional')
        billing_period = data.get('billing_period', 'monthly')
        
        # Simple admin authentication
        if admin_key != 'fix-payment-issue-2024':
            return jsonify({'error': 'Unauthorized'}), 401
        
        if not user_email:
            return jsonify({'error': 'Email is required'}), 400
        
        logger.info(f"🔧 ADMIN: Manually updating plan for {user_email} to {plan_id}")
        
        # Update the plan
        result = mongo_client.update_user_plan(user_email, plan_id, billing_period)
        
        if result:
            logger.info(f"✅ ADMIN: Successfully updated {user_email} to {plan_id}")
            invalidate_profile_cache(user_email)
            
            # Fetch updated user info
            user = mongo_client.get_user_by_email(user_email)
            
            return jsonify({
                'success': True,
                'message': f'Successfully updated {user_email} to {plan_id} plan',
                'plan_id': user.get('plan_id'),
                'tokens_balance': user.get('tokens_balance'),
                'tokens_total': user.get('tokens_total'),
                'features_enabled': user.get('features_enabled')
            })
        else:
            logger.error(f"❌ ADMIN: Failed to update plan for {user_email}")
            return jsonify({'error': 'Failed to update plan'}), 500
            
    except Exception as e:
        logger.error(f"Error in admin fix-plan: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# ============================================
# ANALYSIS REPORT GENERATION
# ============================================

@app.route('/api/batch/<batch_id>/generate-report', methods=['POST'])
@require_api_key
@require_auth
def generate_analysis_report(batch_id):
    """Generate comprehensive PDF report for batch analysis."""
    try:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib import colors
        from reportlab.lib.units import mm, inch
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image as RLImage
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        from reportlab.pdfgen import canvas as pdf_canvas
        from io import BytesIO
        import tempfile
        import matplotlib
        matplotlib.use('Agg')  # Non-interactive backend
        import matplotlib.pyplot as plt
        import numpy as np
        
        data = request.get_json()
        user_email = data.get('email')
        
        if not user_email:
            return jsonify({'error': 'User email is required'}), 400
        
        # Get batch analysis data
        batch = mongo_client.get_batch(batch_id)
        if not batch:
            return jsonify({'error': 'Batch not found'}), 404
        
        # Get token consumption data
        user = mongo_client.get_user_by_email(user_email)
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        files = batch.get('files', [])
        stats = batch.get('stats', {})
        total_piis = stats.get('piis', 0)
        breakdown = stats.get('breakdown', {})
        
        # Generate charts
        temp_chart_dir = os.path.join(os.getcwd(), 'temp_charts')
        os.makedirs(temp_chart_dir, exist_ok=True)
        
        # Generate pie chart
        pie_chart_path = os.path.join(temp_chart_dir, f'pie_{batch_id}.png')
        sorted_breakdown = sorted(breakdown.items(), key=lambda x: x[1], reverse=True)[:12]
        
        if sorted_breakdown:
            labels = [item[0] for item in sorted_breakdown]
            sizes = [item[1] for item in sorted_breakdown]
            
            fig, ax = plt.subplots(figsize=(8, 6))
            colors_pie = plt.cm.Set3(np.linspace(0, 1, len(labels)))
            wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.1f%%',
                                               colors=colors_pie, startangle=90)
            ax.set_title('PII Type Distribution (Top 12)', fontsize=14, fontweight='bold', pad=20)
            
            # Make percentage text more readable
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')
                autotext.set_fontsize(9)
            
            # Make labels more readable
            for text in texts:
                text.set_fontsize(8)
            
            plt.tight_layout()
            plt.savefig(pie_chart_path, dpi=150, bbox_inches='tight')
            plt.close()
        
        # Generate bar chart
        bar_chart_path = os.path.join(temp_chart_dir, f'bar_{batch_id}.png')
        top_10 = sorted(breakdown.items(), key=lambda x: x[1], reverse=True)[:10]
        
        if top_10:
            labels_bar = [item[0][:15] for item in top_10]  # Truncate long labels
            values_bar = [item[1] for item in top_10]
            
            fig, ax = plt.subplots(figsize=(10, 6))
            bars = ax.barh(labels_bar, values_bar, color='#667eea')
            ax.set_xlabel('Count', fontsize=11, fontweight='bold')
            ax.set_title('Top 10 PII Types (Count)', fontsize=14, fontweight='bold', pad=20)
            ax.invert_yaxis()  # Highest on top
            
            # Add value labels on bars
            for bar in bars:
                width = bar.get_width()
                ax.text(width, bar.get_y() + bar.get_height()/2, f'{int(width)}',
                       ha='left', va='center', fontweight='bold', fontsize=9)
            
            plt.tight_layout()
            plt.savefig(bar_chart_path, dpi=150, bbox_inches='tight')
            plt.close()
        
        # Create PDF buffer
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=15*mm,
            bottomMargin=15*mm,
            leftMargin=20*mm,
            rightMargin=20*mm
        )
        
        # Container for PDF elements
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#4F46E5'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=10,
            spaceBefore=15,
            fontName='Helvetica-Bold'
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#4338CA'),
            spaceAfter=8,
            spaceBefore=10,
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=10,
            spaceAfter=6,
            fontName='Helvetica'
        )
        
        # PAGE 1: Cover Page with Company Info
        story.append(Spacer(1, 50))
        story.append(Paragraph("PII SENTINEL", title_style))
        story.append(Spacer(1, 10))
        story.append(Paragraph("Comprehensive Analysis Report", heading_style))
        story.append(Spacer(1, 30))
        
        # About section
        story.append(Paragraph("<b>About PII Sentinel</b>", subheading_style))
        about_text = """
        PII Sentinel is an enterprise-grade solution for detecting, analyzing, and protecting 
        Personally Identifiable Information (PII) in your documents. Our advanced AI-powered 
        system helps organizations maintain compliance with data protection regulations while 
        securing sensitive information.
        """
        story.append(Paragraph(about_text, body_style))
        story.append(Spacer(1, 15))
        
        story.append(Paragraph("<b>Our Product</b>", subheading_style))
        product_text = """
        • Automated PII Detection across 50+ data types<br/>
        • Real-time Risk Assessment and Compliance Scoring<br/>
        • Advanced Masking and Redaction Capabilities<br/>
        • Comprehensive Analytics and Reporting<br/>
        • Enterprise-grade Security and Encryption
        """
        story.append(Paragraph(product_text, body_style))
        story.append(Spacer(1, 20))
        
        # Report metadata
        report_data = [
            ['Report Generated:', datetime.utcnow().strftime('%B %d, %Y at %I:%M %p UTC')],
            ['Generated By:', user.get('fullName', user_email)],
            ['Batch ID:', batch_id],
            ['Batch Name:', batch.get('name', 'N/A')]
        ]
        report_table = Table(report_data, colWidths=[120, 300])
        report_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, -1), 'Helvetica', 9),
            ('FONT', (0, 0), (0, -1), 'Helvetica-Bold', 9),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#4338CA')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(report_table)
        
        story.append(PageBreak())
        
        # PAGE 2: Table of Contents
        story.append(Paragraph("Table of Contents", heading_style))
        story.append(Spacer(1, 10))
        
        toc_data = [
            ['Section', 'Page'],
            ['1. Executive Summary', '3'],
            ['2. Analysis Statistics', '4'],
            ['3. PII Detection Results', '5'],
            ['4. Risk Assessment', '6'],
            ['5. Token Usage', '7'],
            ['6. Detailed PII Breakdown', '8+']
        ]
        toc_table = Table(toc_data, colWidths=[350, 80])
        toc_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, 0), 'Helvetica-Bold', 11),
            ('FONT', (0, 1), (-1, -1), 'Helvetica', 10),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E7FF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('ALIGN', (1, 0), (1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(toc_table)
        
        story.append(PageBreak())
        
        # PAGE 3+: Analysis Header & Statistics
        story.append(Paragraph("1. Analysis Overview", heading_style))
        
        # Analysis Header Stats (same as frontend)
        header_stats = [
            ['Metric', 'Value'],
            ['Total Files Processed', str(len(files))],
            ['Total PIIs Detected', str(total_piis)],
            ['Unique PII Types', str(len(breakdown))],
            ['Batch Name', batch.get('name', 'N/A')],
            ['Batch Status', batch.get('status', 'N/A').upper()],
            ['Processing Date', batch.get('created_at', datetime.utcnow()).strftime('%B %d, %Y %I:%M %p')],
        ]
        header_table = Table(header_stats, colWidths=[200, 230])
        header_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, 0), 'Helvetica-Bold', 12),
            ('FONT', (0, 1), (-1, -1), 'Helvetica', 11),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E7FF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('FONT', (0, 1), (0, -1), 'Helvetica-Bold', 11),
            ('TEXTCOLOR', (0, 1), (0, -1), colors.HexColor('#4338CA')),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 10),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(header_table)
        story.append(Spacer(1, 20))
        
        # Add Pie Chart
        story.append(Paragraph("2. PII Type Distribution", heading_style))
        if os.path.exists(pie_chart_path):
            pie_img = RLImage(pie_chart_path, width=450, height=338)
            story.append(pie_img)
            story.append(Spacer(1, 15))
        
        story.append(PageBreak())
        
        # Add Bar Chart
        story.append(Paragraph("3. Top PII Types Analysis", heading_style))
        if os.path.exists(bar_chart_path):
            bar_img = RLImage(bar_chart_path, width=480, height=288)
            story.append(bar_img)
            story.append(Spacer(1, 15))
        
        story.append(PageBreak())
        # Executive Summary
        story.append(Paragraph("4. Executive Summary", heading_style))
        
        summary_text = f"""
        This report contains a comprehensive analysis of {len(files)} file(s) processed in batch "{batch.get('name', 'N/A')}". 
        Our system detected a total of <b>{total_piis} PII instances</b> across <b>{len(breakdown)} different PII types</b>.
        """
        story.append(Paragraph(summary_text, body_style))
        story.append(Spacer(1, 15))
        
        # Key metrics
        story.append(Paragraph("5. Analysis Statistics", heading_style))
        
        stats_data = [
            ['Metric', 'Value'],
            ['Total Files Analyzed', str(len(files))],
            ['Total PIIs Detected', str(total_piis)],
            ['Unique PII Types', str(len(breakdown))],
            ['Batch Status', batch.get('status', 'N/A').upper()],
            ['Created Date', batch.get('created_at', datetime.utcnow()).strftime('%B %d, %Y')],
        ]
        stats_table = Table(stats_data, colWidths=[250, 180])
        stats_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, 0), 'Helvetica-Bold', 11),
            ('FONT', (0, 1), (-1, -1), 'Helvetica', 10),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E7FF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('FONT', (0, 1), (0, -1), 'Helvetica-Bold', 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 15))
        
        # PII Type Distribution
        story.append(Paragraph("6. PII Detection Results", heading_style))
        story.append(Paragraph("PII Type Distribution (All Types)", subheading_style))
        
        # Sort breakdown by count - SHOW ALL, not just top 15
        sorted_breakdown = sorted(breakdown.items(), key=lambda x: x[1], reverse=True)
        
        pii_data = [['PII Type', 'Count', '% of Total']]
        for pii_type, count in sorted_breakdown:
            percentage = (count / total_piis * 100) if total_piis > 0 else 0
            pii_data.append([pii_type, str(count), f"{percentage:.1f}%"])
        
        pii_table = Table(pii_data, colWidths=[250, 90, 90])
        pii_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, 0), 'Helvetica-Bold', 11),
            ('FONT', (0, 1), (-1, -1), 'Helvetica', 9),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E7FF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 6),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F9FAFB')]),
        ]))
        story.append(pii_table)
        
        story.append(PageBreak())
        
        # Risk Assessment
        story.append(Paragraph("7. Risk Assessment", heading_style))
        
        # Calculate risk score
        high_risk_types = ['AADHAAR', 'PAN', 'PASSPORT', 'BANK_ACCOUNT', 'CREDIT_CARD', 'PASSWORD', 'API_KEY']
        high_risk_count = sum(count for pii_type, count in breakdown.items() if any(hr in pii_type.upper() for hr in high_risk_types))
        risk_score = min(100, int((high_risk_count / max(total_piis, 1)) * 100) + 20)
        
        risk_level = 'LOW' if risk_score < 40 else 'MEDIUM' if risk_score < 70 else 'HIGH'
        risk_color = colors.green if risk_score < 40 else colors.orange if risk_score < 70 else colors.red
        
        risk_data = [
            ['Risk Metric', 'Value'],
            ['Overall Risk Score', f"{risk_score}/100"],
            ['Risk Level', risk_level],
            ['High-Risk PIIs Detected', str(high_risk_count)],
            ['Sensitive Data Types', str(len([p for p in breakdown.keys() if any(hr in p.upper() for hr in high_risk_types)]))]
        ]
        risk_table = Table(risk_data, colWidths=[250, 180])
        risk_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, 0), 'Helvetica-Bold', 11),
            ('FONT', (0, 1), (-1, -1), 'Helvetica', 10),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E7FF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('FONT', (0, 1), (0, -1), 'Helvetica-Bold', 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(risk_table)
        story.append(Spacer(1, 15))
        
        # Token Usage
        story.append(Paragraph("8. Token Usage", heading_style))
        
        token_data = [
            ['Token Metric', 'Value'],
            ['Current Plan', user.get('plan_id', 'starter').upper()],
            ['Total Tokens', str(user.get('tokens_total', 0)) if user.get('tokens_total') is not None else 'Unlimited'],
            ['Tokens Used', str(user.get('tokens_used', 0))],
            ['Tokens Balance', str(user.get('tokens_balance', 0)) if user.get('tokens_balance') is not None else 'Unlimited'],
            ['Report Generation Cost', '100 tokens']
        ]
        token_table = Table(token_data, colWidths=[250, 180])
        token_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, 0), 'Helvetica-Bold', 11),
            ('FONT', (0, 1), (-1, -1), 'Helvetica', 10),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E7FF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('FONT', (0, 1), (0, -1), 'Helvetica-Bold', 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(token_table)
        
        story.append(PageBreak())
        
        # Detailed PII Breakdown by File - SHOW ALL PIIs, not just 50
        story.append(Paragraph("9. Detailed PII Breakdown", heading_style))
        
        # Show ALL files, not just first 10
        for idx, file in enumerate(files, 1):
            story.append(Paragraph(f"File {idx}: {file.get('filename', 'Unknown')}", subheading_style))
            
            file_piis = file.get('piis', [])
            if not file_piis:
                story.append(Paragraph("No PIIs detected in this file.", body_style))
                story.append(Spacer(1, 10))
                continue
            
            # Create detailed table for this file - SHOW ALL PIIs
            file_pii_data = [['#', 'PII Type', 'Value', 'Confidence']]
            for pii_idx, pii in enumerate(file_piis, 1):
                pii_type = pii.get('type', 'Unknown')
                pii_value = pii.get('value', 'N/A')
                # Mask sensitive values
                if len(str(pii_value)) > 20:
                    pii_value = str(pii_value)[:17] + '...'
                confidence = f"{pii.get('confidence', 0.9) * 100:.0f}%"
                file_pii_data.append([str(pii_idx), pii_type, str(pii_value), confidence])
            
            file_pii_table = Table(file_pii_data, colWidths=[30, 120, 200, 70])
            file_pii_table.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, 0), 'Helvetica-Bold', 9),
                ('FONT', (0, 1), (-1, -1), 'Helvetica', 8),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E7FF')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
                ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                ('ALIGN', (3, 1), (3, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('PADDING', (0, 0), (-1, -1), 4),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F9FAFB')]),
            ]))
            story.append(file_pii_table)
            story.append(Spacer(1, 15))
            
            # Add page break after each file for better readability
            if idx < len(files):
                story.append(PageBreak())
        
        # Add page numbers and branding
        def add_page_decorations(canvas, doc):
            canvas.saveState()
            # Brand name top left
            canvas.setFont('Helvetica-Bold', 10)
            canvas.setFillColor(colors.HexColor('#4F46E5'))
            canvas.drawString(20*mm, A4[1] - 10*mm, "PII SENTINEL")
            
            # Page number bottom center
            canvas.setFont('Helvetica', 9)
            canvas.setFillColor(colors.grey)
            page_num = canvas.getPageNumber()
            canvas.drawCentredString(A4[0] / 2, 10*mm, f"Page {page_num}")
            
            canvas.restoreState()
        
        # Build PDF
        doc.build(story, onFirstPage=add_page_decorations, onLaterPages=add_page_decorations)
        
        # Get PDF bytes
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        # Clean up temporary chart files
        try:
            if os.path.exists(pie_chart_path):
                os.remove(pie_chart_path)
            if os.path.exists(bar_chart_path):
                os.remove(bar_chart_path)
        except Exception as e:
            logger.warning(f"Failed to clean up chart files: {e}")
        
        # Save to temp file
        temp_dir = os.path.join(os.getcwd(), 'temp_reports')
        os.makedirs(temp_dir, exist_ok=True)
        
        report_filename = f"PII_Sentinel_Report_{batch_id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.pdf"
        report_path = os.path.join(temp_dir, report_filename)
        
        with open(report_path, 'wb') as f:
            f.write(pdf_bytes)
        
        # Deduct tokens for report generation (100 tokens)
        debit_result = mongo_client.debit_tokens(
            user_email,
            100,
            'generate_analysis_report',
            {'batch_id': batch_id, 'report_filename': report_filename}
        )
        
        if not debit_result.get('success'):
            # Clean up file if token deduction failed
            if os.path.exists(report_path):
                os.remove(report_path)
            return jsonify({'error': debit_result.get('error', 'Insufficient tokens')}), 400
        
        # Invalidate cache after token deduction
        invalidate_profile_cache(user_email)
        
        logger.info(f"✅ Generated analysis report for batch {batch_id}: {report_filename}")
        
        return jsonify({
            'success': True,
            'filename': report_filename,
            'download_url': f'/api/batch/{batch_id}/download-report?filename={report_filename}',
            'tokens_used': 100,
            'tokens_remaining': debit_result.get('balance')
        })
        
    except Exception as e:
        logger.error(f"Error generating analysis report: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/batch/<batch_id>/download-report', methods=['GET'])
@require_api_key
def download_analysis_report(batch_id):
    """Download generated analysis report."""
    try:
        filename = request.args.get('filename')
        if not filename:
            return jsonify({'error': 'Filename is required'}), 400
        
        report_path = os.path.join(os.getcwd(), 'temp_reports', filename)
        
        if not os.path.exists(report_path):
            return jsonify({'error': 'Report not found'}), 404
        
        return send_file(
            report_path,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logger.error(f"Error downloading analysis report: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/batch/<batch_id>/share-report', methods=['POST'])
@require_api_key
@require_auth
def share_analysis_report(batch_id):
    """Generate shareable link for analysis report."""
    try:
        data = request.get_json()
        filename = data.get('filename')
        share_method = data.get('method')  # 'email', 'whatsapp', 'link'
        recipient = data.get('recipient')  # email address or phone number
        
        if not filename:
            return jsonify({'error': 'Filename is required'}), 400
        
        report_path = os.path.join(os.getcwd(), 'temp_reports', filename)
        
        if not os.path.exists(report_path):
            return jsonify({'error': 'Report not found'}), 404
        
        # Generate shareable link (valid for 7 days)
        share_token = str(uuid.uuid4())
        expiry = datetime.utcnow() + timedelta(days=7)
        
        # Store share token in database
        mongo_client.db['ReportShares'].insert_one({
            'share_token': share_token,
            'batch_id': batch_id,
            'filename': filename,
            'created_at': datetime.utcnow(),
            'expires_at': expiry,
            'shared_by': data.get('email'),
            'share_method': share_method,
            'recipient': recipient
        })
        
        share_url = f"{os.getenv('FRONTEND_URL', 'http://localhost:3000')}/shared-report/{share_token}"
        
        response_data = {
            'success': True,
            'share_url': share_url,
            'expires_at': expiry.isoformat(),
            'share_method': share_method
        }
        
        # Generate method-specific sharing content
        if share_method == 'whatsapp':
            whatsapp_text = f"View PII Sentinel Analysis Report: {share_url}"
            response_data['whatsapp_url'] = f"https://wa.me/?text={requests.utils.quote(whatsapp_text)}"
        
        elif share_method == 'email':
            response_data['email_subject'] = "PII Sentinel Analysis Report"
            response_data['email_body'] = f"Please find the PII Sentinel analysis report here: {share_url}\n\nThis link expires on {expiry.strftime('%B %d, %Y')}."
        
        logger.info(f"✅ Generated share link for report {filename}: {share_token}")
        
        return jsonify(response_data)
        
    except Exception as e:
        logger.error(f"Error sharing analysis report: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    port = int(os.getenv('FLASK_PORT', 5000))
    host = os.getenv('FLASK_HOST', '0.0.0.0')
    
    logger.info(f"Starting Flask server on {host}:{port}")
    app.run(host=host, port=port, debug=False, threaded=True)

