"""
DOCX-Specific PII Detection Module
Extracts text from DOCX files and detects PII with document structure awareness
Uses advanced PII detector with 40+ patterns
"""

import logging
from typing import List, Dict
from docx import Document
from pii_detector import pii_detector  # Use advanced detector with ALL patterns

logger = logging.getLogger(__name__)


class DOCXPIIDetector:
    """DOCX-specific PII detection (OPTIMIZED)"""
    
    def __init__(self):
        # No need to initialize - just extract and return text
        pass
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            str: Extracted text
        """
        try:
            doc = Document(docx_path)
            text = ""
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                text += "\n--- TABLE ---\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text += row_text + "\n"
            
            return text
        
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            return ""
    
    def detect_pii_in_docx(self, docx_path: str) -> Dict:
        """
        Detect PII in DOCX file using ADVANCED detector (40+ patterns)
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            dict: Detection results
        """
        try:
            # Extract text
            text = self.extract_text_from_docx(docx_path)
            
            if not text:
                return {
                    'success': False,
                    'error': 'Failed to extract text from DOCX',
                    'detections': []
                }
            
            # Use ADVANCED detector with all patterns
            logger.info(f"Scanning DOCX with advanced detector (40+ patterns)...")
            detections = pii_detector.scan_text(text)
            
            # Format detections
            formatted_detections = []
            for detection in detections:
                formatted_detections.append({
                    'type': detection.get('type', 'UNKNOWN'),
                    'value': detection.get('match', detection.get('value', '')),
                    'confidence': detection.get('confidence', 0.5),
                    'start': detection.get('start', 0),
                    'end': detection.get('end', 0)
                })
            
            logger.info(f"✓ Found {len(formatted_detections)} PII instances in DOCX")
            
            # Categorize detections
            categorized = self._categorize_detections(formatted_detections)
            
            # Document structure analysis
            doc_structure = self._analyze_document_structure(docx_path)
            
            return {
                'success': True,
                'file': docx_path,
                'format': 'DOCX',
                'total_pii_found': len(formatted_detections),
                'detections': formatted_detections,
                'categorized': categorized,
                'extracted_text_length': len(text),
                'document_structure': doc_structure
            }
        
        except Exception as e:
            logger.error(f"Error detecting PII in DOCX: {e}", exc_info=True)
            return {
                'success': False,
                'error': str(e),
                'detections': []
            }
    
    def detect_pii_in_docx_by_element(self, docx_path: str) -> Dict:
        """
        Detect PII in DOCX with element-level granularity (paragraphs, tables)
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            dict: Detection results organized by document elements
        """
        try:
            doc = Document(docx_path)
            element_detections = {}
            all_detections = []
            
            element_num = 1
            
            # Process paragraphs
            for para_num, para in enumerate(doc.paragraphs, 1):
                if para.text.strip():
                    detections = self.pii_detector.detect(para.text)
                    
                    if detections:
                        key = f'para_{para_num}'
                        element_detections[key] = {
                            'type': 'paragraph',
                            'element_number': para_num,
                            'content': para.text[:100] + '...' if len(para.text) > 100 else para.text,
                            'pii_count': len(detections),
                            'detections': detections,
                            'summary': self.pii_detector.get_pii_summary(detections)
                        }
                        all_detections.extend(detections)
            
            # Process tables
            for table_num, table in enumerate(doc.tables, 1):
                for row_num, row in enumerate(table.rows, 1):
                    row_text = " | ".join(cell.text for cell in row.cells)
                    
                    if row_text.strip():
                        detections = self.pii_detector.detect(row_text)
                        
                        if detections:
                            key = f'table_{table_num}_row_{row_num}'
                            element_detections[key] = {
                                'type': 'table_row',
                                'table_number': table_num,
                                'row_number': row_num,
                                'content': row_text[:100] + '...' if len(row_text) > 100 else row_text,
                                'pii_count': len(detections),
                                'detections': detections,
                                'summary': self.pii_detector.get_pii_summary(detections)
                            }
                            all_detections.extend(detections)
            
            return {
                'success': True,
                'file': docx_path,
                'format': 'DOCX',
                'total_pii_found': len(all_detections),
                'total_elements_with_pii': len(element_detections),
                'element_detections': element_detections,
                'categorized': self._categorize_detections(all_detections)
            }
        
        except Exception as e:
            logger.error(f"Error detecting PII by element: {e}", exc_info=True)
            return {
                'success': False,
                'error': str(e)
            }
    
    @staticmethod
    def _analyze_document_structure(docx_path: str) -> Dict:
        """Analyze DOCX document structure"""
        try:
            doc = Document(docx_path)
            
            return {
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'total_sections': len(doc.sections),
                'has_headers': any(section.header.paragraphs for section in doc.sections),
                'has_footers': any(section.footer.paragraphs for section in doc.sections)
            }
        except:
            return {}
    
    @staticmethod
    def _categorize_detections(detections: List[Dict]) -> Dict:
        """Categorize detections by type and category"""
        categorized = {}
        
        for detection in detections:
            category = detection['category']
            pii_type = detection['type']
            
            if category not in categorized:
                categorized[category] = {}
            
            if pii_type not in categorized[category]:
                categorized[category][pii_type] = []
            
            categorized[category][pii_type].append(detection)
        
        return categorized


# Export detector instance
docx_detector = DOCXPIIDetector()

