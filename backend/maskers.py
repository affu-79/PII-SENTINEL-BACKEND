"""
Masking module for PII redaction.
Supports hash (AES-GCM encryption) and blur masking.
"""
import os
import re
import cv2
import numpy as np
from cryptography.hazmat.primitives.ciphers.aead import AESGCM
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.backends import default_backend
import base64
import logging
from typing import List, Dict, Tuple, Optional, Any
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from PIL import Image as PILImage
import io

logger = logging.getLogger(__name__)


class Masker:
    """PII Masking utilities."""
    
    def __init__(self):
        self.backend = default_backend()
    
    def derive_key(self, password: str, salt: bytes, use_sha512: bool = False) -> bytes:
        """Derive AES key from password using PBKDF2 with SHA256 or SHA512."""
        algorithm = hashes.SHA512() if use_sha512 else hashes.SHA256()
        iterations = 150000 if use_sha512 else 100000  # More iterations for SHA512
        
        kdf = PBKDF2HMAC(
            algorithm=algorithm,
            length=32,  # 256-bit key for AES-256
            salt=salt,
            iterations=iterations,
            backend=self.backend
        )
        return kdf.derive(password.encode('utf-8'))
    
    def hash_mask(self, value: str, password: str) -> Dict[str, str]:
        """
        Hash mask a PII value using AES-GCM encryption.
        Returns dict with masked_value and metadata for decryption.
        """
        # Generate salt and IV
        salt = os.urandom(16)
        iv = os.urandom(12)  # 96-bit IV for GCM
        
        # Derive key
        key = self.derive_key(password, salt)
        
        # Encrypt
        aesgcm = AESGCM(key)
        ciphertext = aesgcm.encrypt(iv, value.encode('utf-8'), None)
        
        # Encode for storage
        masked_value = base64.b64encode(ciphertext).decode('utf-8')
        salt_b64 = base64.b64encode(salt).decode('utf-8')
        iv_b64 = base64.b64encode(iv).decode('utf-8')
        
        return {
            'masked_value': masked_value,
            'hash_meta': {
                'salt': salt_b64,
                'iv': iv_b64,
                'algorithm': 'AES-GCM-256'
            }
        }
    
    def decrypt_hash(self, masked_value: str, hash_meta: Dict[str, str], password: str) -> str:
        """Decrypt a hash-masked value."""
        try:
            salt = base64.b64decode(hash_meta['salt'])
            iv = base64.b64decode(hash_meta['iv'])
            ciphertext = base64.b64decode(masked_value)
            
            key = self.derive_key(password, salt)
            aesgcm = AESGCM(key)
            plaintext = aesgcm.decrypt(iv, ciphertext, None)
            
            return plaintext.decode('utf-8')
        except Exception as e:
            logger.error(f"Decryption failed: {e}")
            raise
    
    def blur_region(self, image: np.ndarray, bbox: Tuple[int, int, int, int], blur_strength: int = 15) -> np.ndarray:
        """Apply Gaussian blur to a region in an image."""
        x1, y1, x2, y2 = bbox
        # Ensure coordinates are within image bounds
        h, w = image.shape[:2]
        x1, y1 = max(0, x1), max(0, y1)
        x2, y2 = min(w, x2), min(h, y2)
        
        if x2 <= x1 or y2 <= y1:
            return image
        
        # Extract region
        region = image[y1:y2, x1:x2]
        
        # Apply blur (use odd kernel size)
        if blur_strength % 2 == 0:
            blur_strength += 1
        blurred_region = cv2.GaussianBlur(region, (blur_strength, blur_strength), 0)
        
        # Replace region
        masked_image = image.copy()
        masked_image[y1:y2, x1:x2] = blurred_region
        
        return masked_image
    
    def mask_image(self, image_path: str, pii_results: List[Dict], output_path: str, mask_type: str = "blur") -> str:
        """Mask PIIs in an image file."""
        image = cv2.imread(image_path)
        if image is None:
            raise ValueError(f"Could not read image: {image_path}")
        
        if mask_type == "blur":
            for pii in pii_results:
                if 'bbox' in pii:
                    bbox = pii['bbox']
                    image = self.blur_region(image, bbox)
        else:
            # For hash, we still blur visually but store hash in metadata
            for pii in pii_results:
                if 'bbox' in pii:
                    bbox = pii['bbox']
                    image = self.blur_region(image, bbox)
        
        cv2.imwrite(output_path, image)
        return output_path
    
    def _get_pii_variations(self, pii: Dict) -> List[str]:
        """Get all possible variations of a PII value for masking."""
        value = pii.get('value', '').strip()
        normalized = pii.get('normalized', '').strip()
        
        variations = []
        
        # Add original value
        if value:
            variations.append(value)
        
        # Add normalized value if different
        if normalized and normalized != value:
            variations.append(normalized)
        
        # For numeric PIIs, try variations with/without spaces and dashes
        pii_type = pii.get('type', '').upper()
        if pii_type in ('AADHAAR', 'PHONE', 'BANK_ACCOUNT', 'CARD_NUMBER', 'IMEI'):
            # Try without spaces/dashes
            no_spaces = re.sub(r'[\s\-]', '', value)
            if no_spaces and no_spaces != value and no_spaces not in variations:
                variations.append(no_spaces)
            
            # Try with spaces (for Aadhaar: XXXX XXXX XXXX)
            if pii_type == 'AADHAAR' and len(no_spaces) == 12:
                spaced = f"{no_spaces[:4]} {no_spaces[4:8]} {no_spaces[8:]}"
                if spaced not in variations:
                    variations.append(spaced)
            
            # Try with dashes
            if len(no_spaces) >= 10:
                dashed = '-'.join([no_spaces[i:i+4] for i in range(0, len(no_spaces), 4)])
                if dashed not in variations:
                    variations.append(dashed)
        
        # For email/upi, try case variations
        if pii_type in ('EMAIL', 'UPI'):
            variations.append(value.lower())
            variations.append(value.upper())
        
        return variations
    
    def mask_pdf(self, pdf_path: str, pii_results: List[Dict], output_path: str, mask_type: str = "blur", password: Optional[str] = None) -> str:
        """Mask PIIs in a PDF file."""
        doc = fitz.open(pdf_path)
        
        # Group PIIs by page
        piis_by_page = {}
        for pii in pii_results:
            page_num = pii.get('page', 0)
            if page_num not in piis_by_page:
                piis_by_page[page_num] = []
            piis_by_page[page_num].append(pii)
        
        # Process each page
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_piis = piis_by_page.get(page_num, [])
            
            if not page_piis:
                continue
            
            # Collect all PII text instances first (before any modifications)
            all_pii_instances = []
            for pii in page_piis:
                # Get all variations of this PII
                variations = self._get_pii_variations(pii)
                
                for pii_value in variations:
                    if not pii_value or len(pii_value) < 3:  # Skip very short values
                        continue
                    
                    try:
                        # Search for all instances of this PII value on the page
                        # Use case-insensitive search for better matching
                        text_instances = page.search_for(pii_value, flags=fitz.TEXT_DEHYPHENATE)
                        for inst in text_instances:
                            all_pii_instances.append(inst)
                    except Exception as e:
                        logger.debug(f"Error searching for PII '{pii_value}': {e}")
                        # Fallback: try exact match
                        try:
                            text_instances = page.search_for(pii_value)
                            for inst in text_instances:
                                all_pii_instances.append(inst)
                        except:
                            pass
            
            # Remove duplicate/overlapping instances
            unique_instances = []
            for inst in all_pii_instances:
                rect = fitz.Rect(inst)
                # Check if this instance overlaps significantly with any existing one
                is_duplicate = False
                for existing in unique_instances:
                    existing_rect = fitz.Rect(existing)
                    if rect.intersects(existing_rect):
                        # If overlap is > 50%, consider it duplicate
                        intersection = rect & existing_rect
                        if intersection.get_area() > min(rect.get_area(), existing_rect.get_area()) * 0.5:
                            is_duplicate = True
                            break
                if not is_duplicate:
                    unique_instances.append(inst)
            
            # Method 1: Redact PII instances
            for inst in unique_instances:
                rect = fitz.Rect(inst)
                # Expand slightly to ensure full coverage
                rect.x0 -= 2
                rect.y0 -= 2
                rect.x1 += 2
                rect.y1 += 2
                page.add_redact_annot(rect, fill=(0, 0, 0))  # Black box
            
            # Apply all redactions
            page.apply_redactions()
            
            # Method 2: Also use image-based blurring for visual effect
            # Convert page to image (after redactions, so we see the black boxes)
            pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
            img_data = pix.tobytes("png")
            img = PILImage.open(io.BytesIO(img_data))
            img_array = np.array(img)
            
            # Convert RGB to BGR for OpenCV
            if len(img_array.shape) == 3:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
            
            # Apply additional blur to PII regions (using saved instances)
            for inst in unique_instances:
                # Scale bbox from 72 DPI to 300 DPI
                scale = 300 / 72
                scaled_bbox = (
                    int(inst.x0 * scale),
                    int(inst.y0 * scale),
                    int(inst.x1 * scale),
                    int(inst.y1 * scale)
                )
                # Expand bbox slightly for better coverage
                padding = 10
                scaled_bbox = (
                    max(0, scaled_bbox[0] - padding),
                    max(0, scaled_bbox[1] - padding),
                    min(img_array.shape[1], scaled_bbox[2] + padding),
                    min(img_array.shape[0], scaled_bbox[3] + padding)
                )
                img_array = self.blur_region(img_array, scaled_bbox)
            
            # Convert back to RGB
            if len(img_array.shape) == 3:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_BGR2RGB)
            
            # Convert back to PIL Image
            masked_img = PILImage.fromarray(img_array)
            
            # Replace page with masked image
            rect = page.rect
            img_bytes = io.BytesIO()
            masked_img.save(img_bytes, format='PNG')
            img_bytes.seek(0)
            
            # Clear page and insert masked image
            page.clean_contents()
            page.insert_image(rect, stream=img_bytes.read())
        
        doc.save(output_path)
        doc.close()
        return output_path
    
    def mask_docx(self, docx_path: str, pii_results: List[Dict], output_path: str, mask_type: str = "blur", password: Optional[str] = None) -> str:
        """Mask PIIs in a DOCX file."""
        doc = Document(docx_path)
        
        # Collect all PII variations for efficient masking
        all_pii_values = []
        for pii in pii_results:
            variations = self._get_pii_variations(pii)
            for val in variations:
                if val and len(val) >= 3:
                    all_pii_values.append((val, pii))
        
        def mask_text_in_content(text: str) -> str:
            """Mask all PIIs in text content."""
            masked_text = text
            for pii_value, pii in all_pii_values:
                if pii_value in masked_text:
                    # Use regex for case-insensitive matching where appropriate
                    pii_type = pii.get('type', '').upper()
                    if pii_type in ('EMAIL', 'UPI'):
                        # Case-insensitive for emails/UPI
                        pattern = re.compile(re.escape(pii_value), re.IGNORECASE)
                        if mask_type == "hash" and password:
                            masked = self.hash_mask(pii_value, password)['masked_value']
                        else:
                            masked = '█' * len(pii_value)
                        masked_text = pattern.sub(masked, masked_text)
                    else:
                        # Exact match for other types
                        if mask_type == "hash" and password:
                            masked = self.hash_mask(pii_value, password)['masked_value']
                        else:
                            masked = '█' * len(pii_value)
                        masked_text = masked_text.replace(pii_value, masked)
            return masked_text
        
        # Mask in paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                paragraph.text = mask_text_in_content(paragraph.text)
        
        # Mask in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        cell.text = mask_text_in_content(cell.text)
        
        # Mask in headers and footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text:
                        paragraph.text = mask_text_in_content(paragraph.text)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text:
                        paragraph.text = mask_text_in_content(paragraph.text)
        
        doc.save(output_path)
        return output_path
    
    def mask_text_file(self, text_path: str, pii_results: List[Dict], output_path: str, mask_type: str = "blur", password: Optional[str] = None) -> Dict[str, Any]:
        """
        Mask PIIs in a text file.
        Returns dict with 'output_path' and 'hash_meta_map' (if hash masking).
        """
        # Try to read with UTF-8, fallback to latin-1 if needed
        try:
            with open(text_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
        except:
            with open(text_path, 'r', encoding='latin-1', errors='replace') as f:
                content = f.read()
        
        # Store hash_meta mapping for decryption (created during masking)
        hash_meta_map = {}
        
        # Collect all PII variations
        all_pii_values = []
        for pii in pii_results:
            variations = self._get_pii_variations(pii)
            for val in variations:
                if val and len(val) >= 3:
                    all_pii_values.append((val, pii))
        
        # Sort by length (longest first) to avoid partial matches
        all_pii_values.sort(key=lambda x: len(x[0]), reverse=True)
        
        # Mask all PIIs
        for pii_value, pii in all_pii_values:
            if pii_value in content:
                pii_type = pii.get('type', '').upper()
                
                # Use regex for better matching
                if pii_type in ('EMAIL', 'UPI', 'USERNAME'):
                    # Case-insensitive for these types
                    pattern = re.compile(re.escape(pii_value), re.IGNORECASE)
                    if mask_type == "hash" and password:
                        hash_result = self.hash_mask(pii_value, password)
                        masked = hash_result['masked_value']
                        # Store hash_meta for this specific encryption
                        hash_meta_map[masked] = {
                            'hash_meta': hash_result['hash_meta'],
                            'original_value': pii_value,
                            'pii_type': pii.get('type', ''),
                            'page': pii.get('page', 0)
                        }
                    else:
                        masked = '█' * len(pii_value)
                    content = pattern.sub(masked, content)
                else:
                    # Exact match for other types
                    if mask_type == "hash" and password:
                        hash_result = self.hash_mask(pii_value, password)
                        masked = hash_result['masked_value']
                        # Store hash_meta for this specific encryption
                        hash_meta_map[masked] = {
                            'hash_meta': hash_result['hash_meta'],
                            'original_value': pii_value,
                            'pii_type': pii.get('type', ''),
                            'page': pii.get('page', 0)
                        }
                    else:
                        masked = '█' * len(pii_value)
                    content = content.replace(pii_value, masked)
        
        with open(output_path, 'w', encoding='utf-8', errors='replace') as f:
            f.write(content)
        
        return {
            'output_path': output_path,
            'hash_meta_map': hash_meta_map if mask_type == 'hash' and password else {}
        }


# Global instance
masker = Masker()

