"""
Worker functions for processing files.
This module contains the actual processing logic that runs in parallel.
"""
import os
import logging
import time
from typing import Dict, Any, List, Tuple, Optional
import cv2
import numpy as np
from docx import Document
import pandas as pd
from functools import lru_cache
from ocr_engine import get_ocr_engine
from pii_detector_label_based import label_based_detector
from utils import (
    is_pdf_file, is_docx_file, is_image_file, is_text_file,
    sanitize_filename, get_timestamp, ensure_dir
)

logger = logging.getLogger(__name__)

# DETECTION STRATEGY:
# - PDF, TXT, DOCX → LABEL-BASED detection (explicit labels only)
# - CSV → ADVANCED detector (generic patterns)
# - Images → ADVANCED detector (generic patterns)

# Try to import advanced detector for CSV and images
try:
    from pii_detector_advanced import advanced_pii_detector
    ADVANCED_DETECTOR_AVAILABLE = True
    logger.info("✅ Advanced detector available for CSV/Images")
except ImportError:
    ADVANCED_DETECTOR_AVAILABLE = False
    logger.warning("⚠️ Advanced detector not available for CSV/Images")

logger.info("✅ Label-Based PII Detector ready for PDF/TXT/DOCX")


# Cache OCR engine instance per worker thread
_ocr_engine_cache = None

def _get_cached_ocr_engine():
    """Get cached OCR engine instance."""
    global _ocr_engine_cache
    if _ocr_engine_cache is None:
        _ocr_engine_cache = get_ocr_engine()
    return _ocr_engine_cache

def extract_text_from_file(file_path: str, file_type: str) -> Tuple[str, List[Tuple], int, Optional[List[Dict[str, Any]]]]:
    """
    Extract text from a file.
    Returns: (text, bboxes_list, page_count, image_pii_matches)
    """
    ocr_engine = _get_cached_ocr_engine()
    text_parts = []
    all_bboxes = []
    page_count = 0
    image_pii_matches: Optional[List[Dict[str, Any]]] = None
    
    if is_pdf_file(file_path):
        # OPTIMIZATION: Try text extraction first (much faster than OCR)
        # Only use OCR if text extraction fails or returns minimal text
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            page_count = len(doc)
            extracted_text_parts = []
            
            # Extract text directly from PDF (fast - milliseconds vs seconds for OCR)
            for page_num in range(page_count):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    extracted_text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            
            doc.close()
            
            # If we got substantial text, use it (skip slow OCR)
            if extracted_text_parts and len(' '.join(extracted_text_parts)) > 50:
                text_parts = extracted_text_parts
                logger.debug(f"Using fast text extraction for PDF: {file_path}")
            else:
                # Fallback to OCR if text extraction failed or returned minimal text
                logger.debug(f"Text extraction returned minimal text, using OCR for: {file_path}")
                images = ocr_engine.pdf_to_images(file_path)
                page_count = len(images)
                
                # Process pages in parallel batches for speed
                for image, page_num in images:
                    text, bboxes = ocr_engine.extract_text(image, preprocess=False)  # Skip preprocessing for speed
                    if text:
                        text_parts.append(f"[Page {page_num + 1}]\n{text}")
                        all_bboxes.extend([(bbox, page_num) for bbox in bboxes])
                    del image  # Free memory immediately
        except Exception as e:
            logger.warning(f"Fast text extraction failed for {file_path}, using OCR: {e}")
            # Fallback to OCR
            images = ocr_engine.pdf_to_images(file_path)
            page_count = len(images)
            for image, page_num in images:
                text, bboxes = ocr_engine.extract_text(image, preprocess=False)
            if text:
                text_parts.append(f"[Page {page_num + 1}]\n{text}")
                all_bboxes.extend([(bbox, page_num) for bbox in bboxes])
            del image
    
    elif is_image_file(file_path):
        # Use NEW Image OCR Pipeline for images (PaddleOCR + Tesseract fallback)
        try:
            from image_ocr_pipeline import get_pipeline
            from pii_detector import PIIDetector as PaddlePIIDetector
            
            # Initialize the image OCR pipeline
            pii_detector_instance = PaddlePIIDetector()
            pipeline = get_pipeline(pii_detector_instance)
            
            # Read image file as bytes
            with open(file_path, 'rb') as img_file:
                image_bytes = img_file.read()
            
            filename = os.path.basename(file_path)
            
            # Process image with the new pipeline
            result = pipeline.process_single_image(image_bytes, filename)
            
            # Extract text and PIIs
            text_parts.append(result.ocr_result.full_text)
            
            # Store detailed PII matches (with bounding boxes)
            image_pii_matches = [match.to_dict() for match in result.pii_matches]
            
            # Convert PII matches to bboxes format (if available)
            for pii_match in result.pii_matches:
                if pii_match.bbox:
                    # Convert to tuple format: (x, y, w, h)
                    bbox_tuple = (
                        pii_match.bbox['x'],
                        pii_match.bbox['y'],
                        pii_match.bbox['width'],
                        pii_match.bbox['height']
                    )
                    all_bboxes.append((bbox_tuple, 0))
            
            page_count = 1
            
            logger.info(f"✅ Processed image {filename} with PaddleOCR pipeline: {result.total_piis} PIIs found")
            
        except ImportError as e:
            # Fallback to old OCR engine if image_ocr_pipeline is not available
            logger.warning(f"Image OCR pipeline not available, using fallback: {e}")
            image = cv2.imread(file_path, cv2.IMREAD_COLOR)
            if image is not None:
                text, bboxes = ocr_engine.extract_text(image)
                text_parts.append(text)
                all_bboxes = [(bbox, 0) for bbox in bboxes]
                page_count = 1
                del image  # Free memory
        except Exception as e:
            # Fallback to old OCR engine on any error
            logger.error(f"Error processing image with new pipeline: {e}, using fallback")
            image = cv2.imread(file_path, cv2.IMREAD_COLOR)
            if image is not None:
                text, bboxes = ocr_engine.extract_text(image)
                text_parts.append(text)
                all_bboxes = [(bbox, 0) for bbox in bboxes]
                page_count = 1
                del image  # Free memory
    
    elif is_docx_file(file_path):
        # Extract text from DOCX with optimized parsing
        try:
            doc = Document(file_path)
            paragraphs = []
            # Use list comprehension for faster processing
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text_parts.append("\n".join(paragraphs))
            page_count = 1  # DOCX doesn't have explicit pages
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
    
    elif is_text_file(file_path):
        # Read text file with optimized encoding handling
        try:
            if file_path.endswith('.csv'):
                # For CSV files, extract all cell values for PII detection
                # This ensures PIIs in CSV cells are properly detected
                # Pandas read_csv (pandas>=2.0 removed errors= parameter)
                # Try UTF-8 first, then fallback to latin-1 if decoding fails
                try:
                    df = pd.read_csv(file_path, low_memory=False, encoding='utf-8')
                except UnicodeDecodeError:
                    logger.warning(f"UTF-8 decode failed for CSV {file_path}, trying latin-1")
                    df = pd.read_csv(file_path, low_memory=False, encoding='latin-1')
                except TypeError as e:
                    # pandas < 1.3 may not support low_memory keyword set to False?
                    logger.warning(f"read_csv TypeError for {file_path}: {e}, retrying with defaults")
                    df = pd.read_csv(file_path)
                except Exception as e:
                    logger.error(f"Error reading CSV {file_path} with pandas: {e}", exc_info=True)
                    # Fallback to raw text reading
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                    page_count = 1
                else:
                    # Extract all cell values as text (better for PII detection)
                    # Convert all cells to string and join with spaces/newlines
                    cell_values = []
                    for col in df.columns:
                        cell_values.extend(df[col].astype(str).tolist())
                    
                    # Also keep the structured format for context
                    text = "\n".join([
                        " ".join([str(val) for val in row if pd.notna(val)]) 
                        for _, row in df.iterrows()
                    ])
                    
                    # Add all cell values as a flat list for comprehensive PII detection
                    text += "\n" + " ".join([str(val) for val in cell_values if pd.notna(val) and str(val).strip()])
                    
                    logger.debug(f"Extracted {len(cell_values)} cell values from CSV {file_path}")
            else:
                # Use buffered reading for large files
                with open(file_path, 'r', encoding='utf-8', errors='ignore', buffering=8192) as f:
                    text = f.read()
            text_parts.append(text)
            page_count = 1
        except Exception as e:
            logger.error(f"Error reading text file: {e}", exc_info=True)
    
    full_text = "\n\n".join(text_parts)
    return full_text, all_bboxes, page_count, image_pii_matches


def process_file(file_info: Dict[str, Any]) -> Dict[str, Any]:
    """
    Process a single file: extract text, detect PIIs, return results.
    
    Args:
        file_info: Dict with 'filepath', 'filename', 'batch_id', etc.
    
    Returns:
        Dict with processing results
    """
    start_time = time.time()
    filepath = file_info['filepath']
    filename = file_info['filename']
    batch_id = file_info.get('batch_id', 'unknown')
    
    logger.info(f"Processing file: {filename}")
    
    try:
        # Extract text
        text, bboxes_list, page_count, image_pii_matches = extract_text_from_file(
            filepath,
            file_info.get('file_type', '')
        )
        
        # Reduce logging overhead for speed (only log if needed)
        if len(text) > 0:
            logger.debug(f"Extracted {len(text)} chars from {filename}")
        
        if not text.strip():
            logger.warning(f"No text extracted from {filename}")
            return {
                'filename': filename,
                'success': False,
                'error': 'No text extracted',
                'piis': [],
                'page_count': page_count
            }
        
        # Detect PIIs using appropriate detector based on file type
        all_piis = []
        
        # Choose detector based on file type
        file_ext = file_info.get('file_type', '').lower()
        
        # DETECTION STRATEGY:
        # - PDF (.pdf), TXT (.txt), DOCX (.docx, .doc) → LABEL-BASED (explicit labels)
        # - CSV, Images, Others → ADVANCED detector (generic patterns)
        
        # Use image pipeline PIIs if available
        if image_pii_matches:
            piis = image_pii_matches
            logger.info(f"🖼️ Using {len(piis)} PIIs from image OCR pipeline for {filename}")
        else:
            if file_ext in ['.pdf', '.txt', '.docx', '.doc']:
                selected_detector = label_based_detector
                detector_name = "LABEL-BASED"
                logger.info(f"Using LABEL-BASED detector for {filename} ({file_ext})")
            elif file_ext == '.csv' or is_image_file(filename):
                if ADVANCED_DETECTOR_AVAILABLE:
                    selected_detector = advanced_pii_detector
                    detector_name = "ADVANCED"
                    logger.info(f"Using ADVANCED detector for {filename} ({file_ext})")
                else:
                    selected_detector = label_based_detector
                    detector_name = "LABEL-BASED (fallback)"
                    logger.warning(f"Advanced detector not available, using LABEL-BASED for {filename}")
            else:
                selected_detector = label_based_detector
                detector_name = "LABEL-BASED (default)"
                logger.info(f"Using LABEL-BASED detector for {filename} ({file_ext})")
            
            # Fast PII detection (reduced logging for speed)
            try:
                logger.info(f"Detecting PIIs in {filename}, text length: {len(text)}, detector: {detector_name}")
                piis = selected_detector.detect_pii(text, page_num=0)
                
                # Ensure piis is a list
                if not isinstance(piis, list):
                    logger.warning(f"PII detection returned non-list type: {type(piis)}, converting to list")
                    piis = []
                
                logger.info(f"✓ Detected {len(piis)} PIIs in {filename} using {detector_name}")
                if len(piis) > 0:
                    sample_types = [p.get('type', 'unknown') for p in piis[:5] if isinstance(p, dict)]
                    logger.info(f"  Sample PII types: {sample_types}")
            except Exception as e:
                logger.error(f"✗ Error during PII detection for {filename}: {e}", exc_info=True)
                piis = []
        
        # Apply smart deduplication (for images especially)
        if is_image_file(filename):
            try:
                # Use context-aware PII filtering for images
                from context_aware_pii_filter import get_context_aware_filter
                
                context_filter = get_context_aware_filter()
                
                # Filter PIIs based on document context
                filtered_piis, doc_type = context_filter.filter_by_context(piis, text)
                
                logger.info(f"📊 Context-aware filtering: {len(piis)} → {len(filtered_piis)} PIIs for {doc_type} document")
                
                # Use filtered list for display
                piis = filtered_piis
                
            except ImportError as e:
                logger.warning(f"Context-aware filter not available: {e}")
                # Fallback to basic deduplication
                try:
                    from pii_deduplicator import smart_pii_deduplication
                    
                    display_piis, masking_piis = smart_pii_deduplication(
                        piis,
                        confidence_threshold=0.7,
                        remove_substrings=True
                    )
                    
                    logger.info(f"📊 Deduplication: {len(piis)} → {len(display_piis)} unique PIIs (keeping {len(masking_piis)} for masking)")
                    piis = display_piis
                    
                except ImportError as e2:
                    logger.warning(f"PII deduplicator not available: {e2}")
            except Exception as e:
                logger.error(f"Error during context-aware filtering: {e}")
        
        # Process and format PIIs properly
        for pii in piis:
            # Ensure pii is a dictionary
            if not isinstance(pii, dict):
                logger.warning(f"Invalid PII format (not a dict): {type(pii)}, skipping")
                continue
            
            # Extract required fields with defaults
            pii_type = pii.get('type', 'unknown')
            pii_value = pii.get('value') or pii.get('match', '')
            
            if not pii_type or not pii_value:
                logger.warning(f"Incomplete PII data: type={pii_type}, value={pii_value}, skipping")
                continue
            
            pii_result = {
                'type': str(pii_type),
                'value': str(pii_value),
                'normalized': pii.get('normalized', str(pii_value)),
                'confidence': float(pii.get('confidence', 0.8)),
                'page': int(pii.get('page', 0)),
                'start': int(pii.get('start', 0)),
                'end': int(pii.get('end', 0))
            }
            
            # CRITICAL FIX: Preserve bbox from original PII if it exists (for images)
            if 'bbox' in pii and pii['bbox'] is not None:
                # Image PIIs already have bbox from OCR pipeline - DON'T OVERWRITE
                pii_result['bbox'] = pii['bbox']
                logger.debug(f"   Preserved bbox from image OCR: {pii['bbox']}")
            elif bboxes_list and pii.get('start') is not None:
                # For non-image files, try to find matching bbox from text extraction
                if len(bboxes_list) > 0:
                    bbox, page = bboxes_list[0]
                    pii_result['bbox'] = bbox
                    pii_result['page'] = page
            
            # Preserve additional metadata from context-aware filtering
            if 'occurrence_count' in pii:
                pii_result['occurrence_count'] = pii['occurrence_count']
            if 'all_instances' in pii:
                pii_result['all_instances'] = pii['all_instances']
            if 'all_bboxes' in pii:
                pii_result['all_bboxes'] = pii['all_bboxes']
            if 'document_type' in pii:
                pii_result['document_type'] = pii['document_type']
            
            all_piis.append(pii_result)
        
        # For multi-page PDFs, detect PIIs per page
        if page_count > 1:
            # Re-run detection per page if we have page-separated text
            # This is a simplified version
            pass
        
        processing_time = time.time() - start_time
        
        # Determine which pipeline was used
        pipeline_used = "UNKNOWN"
        document_type = "GENERIC"
        if is_image_file(filename):
            pipeline_used = "image_ocr_pipeline (PaddleOCR + Tesseract)"
            # Document type was determined during context-aware filtering
            # Extract from piis metadata if available
            if piis and len(piis) > 0:
                first_pii = piis[0]
                if 'document_type' in first_pii:
                    document_type = first_pii['document_type']
        elif file_ext in ['.pdf', '.txt', '.docx', '.doc']:
            pipeline_used = "text_extraction + label_based_detector"
        elif file_ext == '.csv':
            pipeline_used = "pandas + advanced_detector" if ADVANCED_DETECTOR_AVAILABLE else "pandas + label_based_detector"
        
        result = {
            'filename': filename,
            'success': True,
            'piis': all_piis,
            'pii_count': len(all_piis),
            'page_count': page_count,
            'text_length': len(text),
            'timestamp': get_timestamp(),
            'processing_time': processing_time,  # Add for ETA calculation
            'detector_used': detector_name,
            'pipeline_used': pipeline_used,  # NEW: Pipeline information for browser console
            'document_type': document_type  # NEW: Document type classification
        }
        
        logger.info(f"✓ Processed {filename}: {len(all_piis)} PIIs detected ({document_type}) using {pipeline_used}, time: {processing_time:.2f}s")
        if len(all_piis) > 0:
            logger.info(f"  PII types found: {list(set([p.get('type', 'unknown') for p in all_piis]))}")
        return result
        
    except Exception as e:
        processing_time = time.time() - start_time
        logger.error(f"Error processing {filename}: {e} (time: {processing_time:.2f}s)", exc_info=True)
        return {
            'filename': filename,
            'success': False,
            'error': str(e),
            'piis': [],
            'page_count': 0,
            'processing_time': processing_time
        }

